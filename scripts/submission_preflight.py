#!/usr/bin/env python
from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

import numpy as np

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
INDEPENDENT_MANUSCRIPT_DOCX = (
    "geology_led_sequential_gaussian_simulation_for_uncertainty_analysis_of_a_"
    "stratiform_graphite_deposit_in_tanzania.docx"
)
CLEAN_REQUIRED_FILES = [
    "Manuscript.docx", "Cover_Letter.docx",
    "Fig1.tif", "Fig2.tif", "Fig3.tif", "Fig4.tif",
    "Fig5.tif", "Fig6.tif", "Fig7.tif",
    "ESM_1.pdf", "ESM_2.xlsx",
]
ESSENTIAL_S2_FILES = {
    "variogram_model.json",
    "validation_metrics.json",
    "cutoff_occupancy_uncertainty.csv",
    "sgs_meta.json",
}


def fail(msg: str) -> None:
    print(f"FAIL: {msg}")
    raise SystemExit(1)


def require(path: Path) -> None:
    if not path.exists():
        fail(f"Missing required artifact: {path}")


def _docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _docx_paragraph_lines(path: Path) -> list[str]:
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]


def _extract_abstract_and_keywords(manuscript_docx: Path) -> tuple[str, list[str]]:
    lines = _docx_paragraph_lines(manuscript_docx)
    abstract_lines: list[str] = []
    keywords: list[str] = []
    in_abstract = False
    for line in lines:
        low = line.lower().strip()
        if low == "abstract":
            in_abstract = True
            continue
        if in_abstract and (re.match(r"^\d+(\.\d+)*\s", line) or low.startswith("keywords:")):
            in_abstract = False
        if in_abstract:
            abstract_lines.append(line.strip())
        if low.startswith("keywords:"):
            raw = line.split(":", 1)[1].strip()
            keywords = [k.strip() for k in re.split(r"[;,]", raw) if k.strip()]
    return " ".join(abstract_lines).strip(), keywords


def _count_abstract_words(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9%\-]+", text))


def _docx_has_line_numbering(path: Path) -> bool:
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml")
    return b"w:lnNumType" in xml


def _docx_is_single_column(path: Path) -> bool:
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml")
    root = ET.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for cols in root.findall(".//w:sectPr/w:cols", ns):
        num = cols.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num")
        if num and str(num).strip().isdigit() and int(num) > 1:
            return False
    return True


def _docx_media_count(path: Path) -> int:
    with zipfile.ZipFile(path) as zf:
        return len([name for name in zf.namelist() if name.startswith("word/media/") and not name.endswith("/")])


def _is_md_separator_row(line: str) -> bool:
    raw = line.strip()
    if not raw.startswith("|") or not raw.endswith("|"):
        return False
    inner = raw.strip("|").replace(" ", "")
    if not inner:
        return False
    return all(ch in "-:|" for ch in raw.replace(" ", ""))


def _find_body_markdown_table_before_tables(md_path: Path) -> int | None:
    if not md_path.exists():
        return None
    lines = md_path.read_text(encoding="utf-8").splitlines()
    tables_idx = -1
    for i, line in enumerate(lines):
        if re.match(r"^\s*##\s*TABLES\s*$", line, flags=re.IGNORECASE):
            tables_idx = i
            break
    if tables_idx == -1:
        return None

    in_fence = False
    for i in range(0, tables_idx):
        stripped = lines[i].strip()
        if stripped.startswith("```"):
            in_fence = not in_fence
            continue
        if in_fence or not stripped.startswith("|"):
            continue
        j = i + 1
        while j < tables_idx and not lines[j].strip():
            j += 1
        if j < tables_idx and _is_md_separator_row(lines[j]):
            return i + 1
    return None


def _extract_supplement_claims(md_path: Path) -> set[str]:
    if not md_path.exists():
        return set()
    text = md_path.read_text(encoding="utf-8")
    claimed = set()
    for m in re.finditer(r"`?supplement/([A-Za-z0-9_.-]+)`?", text):
        claimed.add(m.group(1))
    return claimed


def _extract_highlight_bullets_from_docx(path: Path) -> list[str]:
    lines = _docx_paragraph_lines(path)
    bullets: list[str] = []
    for line in lines:
        normalized = line.lstrip("•*- ").strip()
        if normalized:
            bullets.append(normalized)
    return bullets


def _check_figure_resolution(fig_path: Path) -> tuple[tuple[int, int], tuple[float, float], str]:
    try:
        from PIL import Image  # type: ignore
    except Exception as exc:
        fail(f"Pillow unavailable for figure resolution checks: {exc}")
    with Image.open(fig_path) as im:
        size = im.size
        dpi = im.info.get("dpi", (0.0, 0.0))
        compression = str(im.info.get("compression", "")).strip().lower()
    try:
        dpi_tuple = (float(dpi[0]), float(dpi[1]))  # type: ignore[index]
    except Exception:
        dpi_tuple = (0.0, 0.0)
    return size, dpi_tuple, compression


def _figure_minimums(name: str) -> tuple[int, int]:
    return 1000, 5000


def _docx_review_markup_counts(path: Path) -> dict[str, int | bool]:
    counts: dict[str, int | bool] = {
        "comments_part": False,
        "w:ins": 0,
        "w:del": 0,
        "w:moveFrom": 0,
        "w:moveTo": 0,
        "w:commentRangeStart": 0,
        "w:commentReference": 0,
        "trackRevisions": 0,
    }
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        counts["comments_part"] = any(name.lower().startswith("word/comments") for name in names)
        for name in names:
            lower = name.lower()
            if not (lower.startswith("word/") and lower.endswith(".xml")):
                continue
            text = zf.read(name).decode("utf-8", errors="ignore")
            for key in ["w:ins", "w:del", "w:moveFrom", "w:moveTo", "w:commentRangeStart", "w:commentReference"]:
                counts[key] = int(counts[key]) + len(
                    re.findall(r"<" + re.escape(key) + r"(?=[\s>/])", text)
                )
            counts["trackRevisions"] = int(counts["trackRevisions"]) + len(
                re.findall(r"<w:trackRevisions(?=[\s>/])", text)
            )
    return counts


def _fail_on_docx_review_markup(path: Path) -> None:
    counts = _docx_review_markup_counts(path)
    if counts["comments_part"] or any(int(counts[k]) > 0 for k in counts if k != "comments_part"):
        fail(f"{path.name} contains comments or tracked-change markup: {counts}")


def _fail_on_s2_path_leaks(sub_dir: Path) -> None:
    leak_re = re.compile(r"C:\\|Users[\\/]|OneDrive[\\/]|Desktop[\\/]|Tanga_New", flags=re.I)
    with zipfile.ZipFile(sub_dir / "Supplementary_Data_S2.zip") as zf:
        for name in zf.namelist():
            if name.endswith("/") or not name.lower().endswith((".json", ".csv", ".txt", ".md", ".xml")):
                continue
            text = zf.read(name).decode("utf-8", errors="ignore")
            if leak_re.search(text):
                fail(f"Supplementary_Data_S2.zip contains local path leakage in {name}")


def check_clean_jaes_package(sub_dir: Path) -> None:
    required = CLEAN_REQUIRED_FILES
    figure_files = [name for name in CLEAN_REQUIRED_FILES if name.lower().endswith(".tif")]
    for name in required:
        require(sub_dir / name)

    # Ignore transient Microsoft Office owner files. They can appear when a
    # generated DOCX is open locally, but they are not part of the upload set.
    present = sorted(item.name for item in sub_dir.iterdir() if item.is_file() and not item.name.startswith("~$"))
    extras = sorted(set(present) - set(required))
    if extras:
        fail(f"Clean package contains extra files beyond required-only set: {extras}")
    if len(present) != len(required):
        fail(f"Clean package file count mismatch: expected {len(required)}, found {len(present)}")

    forbidden_names = {
        "paper.md",
        "paper_body.md",
        "tables_final.md",
        "figure_captions_final.md",
        "SOURCE_OF_TRUTH.md",
        "source_of_truth.submission.json",
        "SUBMISSION_CHECKLIST.md",
        "CONTRADICTION_FIX_REPORT.md",
        "preflight_report.json",
        "submission_package_final_clean.zip",
        "sudip_manuscript.docx",
        "paper.docx",
        "Table.docx",
        "Fig.docx",
    }
    present_forbidden = sorted(name for name in forbidden_names if (sub_dir / name).exists())
    if present_forbidden:
        fail(f"Clean package contains internal or duplicate files: {present_forbidden}")
    for item in sub_dir.iterdir():
        if item.name.startswith("~$"):
            continue
        if item.suffix.lower() in {".md", ".json"}:
            fail(f"Clean package contains non-submission source file: {item.name}")

    manuscript = _docx_text(sub_dir / "04_Manuscript.docx")
    cover = _docx_text(sub_dir / "07_Cover_Letter.docx")
    title = _docx_text(sub_dir / "01_Title_Page.docx")
    combined = "\n".join([manuscript, cover, title])
    required_tokens_any = [
        ["Abstract"],
        ["Keywords:"],
        ["no specific external grant"],
        ["belongs to the project data holder and is subject to confidentiality restrictions"],
        ["CRediT"],
    ]
    for token_group in required_tokens_any:
        if not any(tok.lower() in combined.lower() for tok in token_group):
            fail(f"Clean package missing required JAES/front-matter token: {' or '.join(token_group)}")
    forbidden_text = [
        "SOURCE_OF_TRUTH",
        "source_of_truth",
        "CONTRADICTION",
        "reviewer-first",
        "First SGS uncertainty " "workflow",
        "Name - Institution - Email",
        "Best method from baseline panel",
        "anonymized composite table",
        "Tanga graphite project",
        "risked_tonnage.csv",
        "Screening-cutoff uncertainty volume is reported without resource claims",
        "TDM006",
        "TDM007",
        "TDM008",
        "60.9%",
        "65.5%",
        "44.5%",
        "Geology-Conditioned Uncertainty in a Stratiform Flake-Graphite System of the Tanzanian Mozambique Belt",
    ]
    for token in forbidden_text:
        if token.lower() in combined.lower():
            fail(f"Clean package contains internal/risky wording: {token}")
    if re.search(r"(?<!\\f)rac\{", manuscript):
        fail("Clean manuscript contains a malformed LaTeX fraction command")
    if "Corresponding Author" in combined:
        fail("Clean package uses `Corresponding Author`; use `Corresponding author`")

    with zipfile.ZipFile(sub_dir / "Supplementary_Data_S2.zip") as zf:
        names = {Path(n).name for n in zf.namelist() if not n.endswith("/")}
        if names != ESSENTIAL_S2_FILES:
            fail(
                "Supplementary_Data_S2.zip must contain exactly "
                f"{sorted(ESSENTIAL_S2_FILES)}, found {sorted(names)}"
            )
        try:
            validation_text = zf.read("validation_metrics.json").decode("utf-8")
            validation_payload = json.loads(validation_text)
            sgs_meta_payload = json.loads(zf.read("sgs_meta.json").decode("utf-8"))
        except Exception as exc:
            fail(f"S2 JSON is not readable: {exc}")
        simulation_meta = sgs_meta_payload.get("config", {}).get("simulation", {})
        if simulation_meta.get("kriging_type") != "SK_style_effective":
            fail("S2 sgs_meta.json must report the implemented SK-style estimator as primary")
        if "postrun_review_pack" in sgs_meta_payload:
            fail("S2 sgs_meta.json still contains stale internal postrun-review metadata")
        categorical_meta = sgs_meta_payload.get("categorical_domain_simulation", {})
        if categorical_meta.get("method") != "fixed_local_probability_sampling":
            fail("S2 sgs_meta.json does not disclose the categorical probability-sampling method")
        if categorical_meta.get("search_radii_m") != [250.0, 200.0, 20.0]:
            fail("S2 categorical-domain search radii do not match the canonical run")
        if int(categorical_meta.get("max_neighbors", 0) or 0) != 20:
            fail("S2 categorical-domain maximum-neighbour count is not 20")
        if float(categorical_meta.get("global_class_prior_weight", 0.0) or 0.0) != 2.0:
            fail("S2 categorical-domain prior weight is not 2.0")
        if "weathering_upgrade" in validation_text.lower() or "weathering upgrade" in validation_text.lower():
            fail("S2 validation_metrics.json still exposes weathering-upgrade wording")
        scalar = validation_payload.get("geology_uncertainty_scalar_summaries", {})
        required_scalar_keys = {
            "signed_graphitic_host_profile",
            "weathering_stratified_unsigned_contact_distance",
            "graphitic_weathering_comparison",
            "domain_uncertainty",
            "thickness_geometry",
            "support_ladder",
            "vertical_continuity",
            "simulation_support_validation",
            "physical_domain_audit",
        }
        missing_scalar = sorted(k for k in required_scalar_keys if k not in scalar)
        if missing_scalar:
            fail(f"S2 validation_metrics.json missing compact scalar summaries: {missing_scalar}")
        if any(key in scalar for key in ("contact_analysis", "weathering_summary", "weathering_relative_mean_contrast")):
            fail("S2 validation_metrics.json still contains a superseded contact/weathering summary")
        signed_scalar = scalar.get("signed_graphitic_host_profile", {})
        if (signed_scalar.get("n_composites"), signed_scalar.get("contact_count"), signed_scalar.get("contact_holes")) != (711, 134, 42):
            fail("S2 signed graphitic-host counts do not match the manuscript source of truth")
        unsigned_scalar = scalar.get("weathering_stratified_unsigned_contact_distance", {})
        if (unsigned_scalar.get("n_composites"), unsigned_scalar.get("n_holes")) != (546, 64):
            fail("S2 unsigned weathering-stratified contact counts are inconsistent")
        weathering_scalar = scalar.get("graphitic_weathering_comparison", {})
        if (weathering_scalar.get("fresh_n"), weathering_scalar.get("weathered_n")) != (3566, 382):
            fail("S2 graphitic weathering counts do not match the manuscript")
        examples = validation_payload.get("calculation_examples", {})
        if not {"normalised_shannon_entropy", "graphitic_probability", "thickness_aperture"}.issubset(examples):
            fail("S2 validation_metrics.json lacks the required synthetic calculation examples")
        cv_baseline = validation_payload.get("withheld_composite_validation_baseline", {})
        cv_families = cv_baseline.get("fold_families", []) if isinstance(cv_baseline, dict) else []
        cv_family_names = {
            str(row.get("validation_family"))
            for row in cv_families
            if isinstance(row, dict)
        }
        required_cv_families = {"blocked_500", "leave_hole", "leave_section_100m"}
        missing_cv = sorted(required_cv_families - cv_family_names)
        if missing_cv:
            fail(f"S2 validation_metrics.json missing withheld-composite validation baseline families: {missing_cv}")
        required_diagnostics = {
            "variogram_reproduction",
            "realisation_count_normalised_sensitivity",
            "spatial_overlap_bootstrap",
            "signed_graphitic_host_contact",
            "ensemble_convergence",
            "support_aligned_mean_decomposition",
            "categorical_domain_grouped_validation",
            "no_domain_pilot_realisation_bootstrap",
            "directional_swath_curves",
        }
        missing_diagnostics = sorted(key for key in required_diagnostics if key not in validation_payload)
        if missing_diagnostics:
            fail(f"S2 validation_metrics.json missing requested diagnostics: {missing_diagnostics}")
        decomposition = validation_payload.get("support_aligned_mean_decomposition", {})
        classes = decomposition.get("classes", [])
        if len(classes) != 3:
            fail("S2 support-aligned mean decomposition must contain three classes")
        if abs(float(decomposition.get("fraction_sum_pct", 0.0)) - 100.0) > 1e-6:
            fail("S2 support-aligned mean fractions do not sum to 100%")
        if float(decomposition.get("reconstruction_error_tgc_pct", 1.0)) > 0.001:
            fail("S2 support-aligned weighted means do not reconstruct the whole-grid mean")
        categorical_validation = validation_payload.get("categorical_domain_grouped_validation", {})
        if not categorical_validation.get("zero_hole_leakage") or int(categorical_validation.get("hole_overlap_count", 1)) != 0:
            fail("S2 grouped categorical validation has drillhole leakage")
        if int(categorical_validation.get("n_splits", 0)) != 5:
            fail("S2 grouped categorical validation must contain five folds")
        confusion = categorical_validation.get("confusion_matrix", {})
        if confusion.get("labels") != ["fresh_graphitic", "weathered_graphitic", "host_waste"]:
            fail("S2 categorical confusion matrix labels are incomplete or out of order")
        if np.asarray(confusion.get("counts", [])).shape != (3, 3):
            fail("S2 categorical confusion matrix must be 3 x 3")
        search_support = categorical_validation.get("search_support", {})
        within_support = search_support.get("within_support", {}) if isinstance(search_support, dict) else {}
        outside_support = search_support.get("outside_support", {}) if isinstance(search_support, dict) else {}
        if int(within_support.get("n", 0)) + int(outside_support.get("n", 0)) != int(categorical_validation.get("n_composites", -1)):
            fail("S2 categorical search-support decomposition does not reconstruct the withheld population")
        graphitic_validation = categorical_validation.get("graphitic_vs_host", {})
        if len(graphitic_validation.get("calibration_by_probability_decile", [])) != 10:
            fail("S2 categorical reliability data must contain ten fixed probability bins")
        nested_platt = graphitic_validation.get("nested_platt_recalibration_sensitivity", {})
        if nested_platt.get("applied_to_canonical_domain_realisations") is not False:
            fail("S2 must state that nested Platt recalibration was not applied to canonical realisations")
        entropy_ranking = categorical_validation.get("entropy_error_ranking", {}).get("within_search_support", {})
        if entropy_ranking.get("entropy_error_roc_auc") is None:
            fail("S2 lacks within-search entropy error-ranking evidence")
        null_bootstrap = validation_payload.get("no_domain_pilot_realisation_bootstrap", {})
        if int(null_bootstrap.get("n_bootstrap", 0)) != 200:
            fail("S2 no-domain pilot bootstrap must record 200 resamples")
        if int(null_bootstrap.get("independent_seed_families_completed", 0)) != 1:
            fail("S2 no-domain pilot bootstrap must disclose the single independent seed family")
        null_intervals = null_bootstrap.get("bootstrap_5_50_95", {})
        if not {"hist_overlap", "qq_rmse", "swath_corr_x", "swath_corr_y", "swath_corr_z"}.issubset(null_intervals):
            fail("S2 no-domain pilot bootstrap intervals are incomplete")
        swaths = validation_payload.get("directional_swath_curves", {}).get("curves", {})
        if not {"along_strike", "down_dip", "normal_to_plane"}.issubset(swaths):
            fail("S2 directional swath curves omit a canonical geological axis")

        convergence = validation_payload.get("ensemble_convergence", {})
        if convergence.get("checkpoints") != [5, 10, 20, 30, 50, 75, 100]:
            fail("S2 ensemble_convergence does not contain all required checkpoints")
        if int(convergence.get("random_subsets_per_checkpoint", 0) or 0) != 200:
            fail("S2 ensemble_convergence must record 200 random subsets per checkpoint")
        if int(convergence.get("seed", 0) or 0) != 20260706:
            fail("S2 ensemble_convergence seed is not 20260706")
        if "acceptance_passed" not in convergence or not convergence.get("acceptance_gates"):
            fail("S2 ensemble_convergence missing acceptance result or gate details")
        contact_profile = validation_payload.get("signed_graphitic_host_contact", {})
        if len(contact_profile.get("bin_rows", [])) != 6:
            fail("S2 signed_graphitic_host_contact must contain all six signed-distance bins")
        variogram_reproduction = validation_payload.get("variogram_reproduction", {})
        direction_curves = variogram_reproduction.get("direction_curves", {})
        required_directions = {"along_strike", "down_dip", "normal_to_plane"}
        if not required_directions.issubset(direction_curves):
            fail("S2 variogram_reproduction.direction_curves missing a canonical geological direction")
        for direction in required_directions:
            curve = direction_curves.get(direction, {})
            if not curve.get("lag_m") or not curve.get("simulation_p50_gamma"):
                fail(f"S2 variogram curve for {direction} is incomplete")
        nested_diagnostics = validation_payload.get("validation_gap_summaries", {})
        if not required_diagnostics.issubset(nested_diagnostics):
            fail("S2 validation_gap_summaries is incomplete")
        try:
            cutoff_header = zf.read("cutoff_occupancy_uncertainty.csv").decode("utf-8", errors="replace").splitlines()[0]
        except Exception as exc:
            fail(f"cutoff_occupancy_uncertainty.csv in S2 is not readable: {exc}")
        forbidden_cols = ["tonnage", "contained_p50", "grade_p50"]
        for col in forbidden_cols:
            if col in cutoff_header.lower():
                fail(f"S2 cutoff occupancy CSV still exposes resource-style column name: {col}")
    _fail_on_s2_path_leaks(sub_dir)
    # Guard against manuscript over-claiming supplemental evidence not actually packaged.
    claimed_supp_files = _extract_supplement_claims(ROOT / "manuscript.md") | _extract_supplement_claims(ROOT / "tables.md")
    missing_claims = sorted(name for name in claimed_supp_files if name not in ESSENTIAL_S2_FILES)
    if missing_claims:
        fail(
            "Manuscript/tables claim supplement files not present in S2: "
            f"{missing_claims}"
        )

    body_table_line = _find_body_markdown_table_before_tables(ROOT / "manuscript.md")
    if body_table_line is not None:
        fail(
            "Body markdown table detected before '## TABLES' in manuscript.md "
            f"(line {body_table_line}). Move body tables to the end-table section."
        )

    manuscript_docx = sub_dir / "04_Manuscript.docx"
    highlights_docx = sub_dir / "02_Highlights.docx"
    for docx_path in sub_dir.glob("*.docx"):
        if docx_path.name.startswith("~$"):
            continue
        _fail_on_docx_review_markup(docx_path)

    abstract, keywords = _extract_abstract_and_keywords(manuscript_docx)
    if not abstract:
        fail("Manuscript abstract not detected")
    abstract_words = _count_abstract_words(abstract)
    if abstract_words > 250:
        fail(f"Abstract exceeds 250 words ({abstract_words})")
    if not (1 <= len(keywords) <= 7):
        fail(f"Keywords count must be 1-7, found {len(keywords)}")

    bullets = _extract_highlight_bullets_from_docx(highlights_docx)
    if not (3 <= len(bullets) <= 5):
        fail(f"Highlights count must be 3-5, found {len(bullets)}")
    over = [b for b in bullets if len(b) > 85]
    if over:
        fail("Highlights contain entries longer than 85 characters")

    if not _docx_has_line_numbering(manuscript_docx):
        fail("Manuscript does not appear to use continuous line numbering")
    if not _docx_is_single_column(manuscript_docx):
        fail("Manuscript appears to use multi-column layout")
    media_count = _docx_media_count(manuscript_docx)
    if media_count < 7:
        fail(f"Manuscript does not embed all end-of-manuscript figures (found {media_count}, expected >=7)")

    table_titles = re.findall(r"\bTable\s+([0-9]+)\.", manuscript)
    if sorted(set(int(n) for n in table_titles if n.isdigit()))[:5] != [1, 2, 3, 4, 5]:
        fail("Manuscript does not expose the required Table 1-5 sequence")
    if any(int(n) > 5 for n in table_titles if n.isdigit()):
        fail("Manuscript contains more than the allowed five main tables")

    # Email and institutional affiliation are mandatory in the package text artifacts.
    if not re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", "\n".join([title, cover])):
        fail("Corresponding author email not detected in title/cover documents")
    if not re.search(r"(?:\+\d{1,3}[-\s]?)?\d[\d\s().-]{7,}\d", "\n".join([title, cover])):
        fail("Corresponding author phone not detected in title/cover documents")
    if "Affiliations:" not in title and "Affiliation:" not in title:
        fail("Title page missing corresponding-author postal affiliation block")
    if re.search(r"\+91\s*90000\s*00000", "\n".join([title, cover])):
        fail("Placeholder phone number detected in title/cover documents")

    manuscript_lines = _docx_paragraph_lines(manuscript_docx)
    manuscript_text_full = "\n".join(manuscript_lines)
    lower_manuscript_text = manuscript_text_full.lower()
    if "the supplement provides reviewer-auditable output tables and run metadata" not in lower_manuscript_text:
        fail("Manuscript must state S2 is audit-level output tables/run metadata, not full workflow regeneration")
    if "supplementary figure" in lower_manuscript_text or "supplementary figures" in lower_manuscript_text:
        fail("Manuscript references supplementary figures not present in required-only S2")
    if "supplementary table" in lower_manuscript_text or "supplementary tables" in lower_manuscript_text:
        fail("Manuscript references supplementary tables not present in required-only S2")
    if "supplementary extension pack" in lower_manuscript_text or "supplementary run artifacts" in lower_manuscript_text:
        fail("Manuscript references non-uploaded supplementary run artifacts")
    if "emitted contact tables and figures" in lower_manuscript_text:
        fail("Manuscript claims contact tables/figures outside the required-only S2 upload")
    if re.search(r"equations?[^.\n]{0,120}Table\s+8", manuscript_text_full, flags=re.IGNORECASE):
        fail("Manuscript points equations to Table 8; equations are Table 10")
    if re.search(r"Table\s+10\s+summarizes\s+that\s+chain", manuscript_text_full, flags=re.IGNORECASE):
        fail("Manuscript points geology-to-model chain to Table 10; chain is Table 12")
    caps_text = (ROOT / "figure_captions.md").read_text(encoding="utf-8", errors="ignore").lower()
    if "supplementary figure" in caps_text:
        fail("figure_captions.md contains supplementary figure captions not included in required-only package")
    if "supplementary table" in caps_text:
        fail("figure_captions.md contains supplementary table references not included in required-only package")
    figure_caption_nums = [
        int(n)
        for n in re.findall(r"(?im)^\s*Figure\s+([0-9]+)\.\s+", manuscript_text_full)
        if n.isdigit()
    ]
    if figure_caption_nums != [1, 2, 3, 4, 5, 6, 7]:
        fail(f"Manuscript must contain exactly Figure captions 1-7 in order, found {figure_caption_nums}")
    fig6_caption_match = re.search(
        r"(?ims)^\s*Figure\s+6\.\s+(.*?)(?=^\s*Figure\s+7\.)",
        manuscript_text_full,
    )
    if not fig6_caption_match:
        fail("Figure 6 caption not detected in clean manuscript")
    fig6_caption = re.sub(r"\s+", " ", fig6_caption_match.group(1)).lower()
    for token in [
        "plan-plus-section localisation",
        "section is at northing",
        "b) p(tgc > 3%)",
        "c) matched p90-p10",
        "d-f) fixed-index realisations 1, 50 and 100",
        "panels b-f use 4x vertical exaggeration",
        "plus or minus 75 m",
    ]:
        if token not in fig6_caption:
            fail(f"Figure 6 caption missing required story/panel token: {token}")
    source_manuscript = (ROOT / "manuscript.md").read_text(encoding="utf-8", errors="ignore")
    if "n labels report bin support" in caps_text:
        fail("Figure 7 caption still describes overlapping pointwise n labels")
    if "aligned bars beneath each swath report composite support" not in lower_manuscript_text:
        fail("Figure 7 caption does not document the separated sample-support strips")
    if "special thanks to the sakariya geology team" in source_manuscript.lower():
        fail("Manuscript acknowledgements retain the named Sakariya geology-team sentence")
    plain_inline_patterns = [
        r"where\s+Z_comp\s+is",
        r"L_i\s+is\s+interval\s+length",
        r"where\s+p_k\(u\)\s+is",
        r"where\s+T_G\^?\{?\(r\)\}?\s+is",
    ]
    for pattern in plain_inline_patterns:
        if re.search(pattern, source_manuscript, flags=re.IGNORECASE):
            fail(f"manuscript.md contains an untypeset inline formula reference matching /{pattern}/")
    for phrase in [
        "The evidence is not an overall model-ranking result",
        "Geological domaining improves the SGS data-analysis framework",
    ]:
        if phrase in source_manuscript and phrase not in manuscript_text_full:
            fail(f"Clean manuscript DOCX appears stale or missing source phrase: {phrase}")

    def _has_numbered_heading(text_blob: str, phrase: str) -> bool:
        return bool(re.search(rf"(?m)^(?:#+\s*)?3\.\d+[^\n]*{re.escape(phrase)}", text_blob))

    for label, text_blob in (("clean manuscript", manuscript_text_full), ("manuscript.md", source_manuscript)):
        if re.search(r"(?m)^(?:#+\s*)?3\.\d+[^\n]*quality assurance and quality control \(QA/QC\)", text_blob):
            fail(f"{label} QA/QC heading must use title capitalization")
        if re.search(r"(?m)^(?:#+\s*)?3\.\d+[^\n]*normal-score transformation \(NST\)", text_blob):
            fail(f"{label} NST heading must use title capitalization")
        if not _has_numbered_heading(text_blob, "Quality Assurance and Quality Control (QA/QC)"):
            fail(f"{label} missing professional QA/QC heading")
        if not _has_numbered_heading(text_blob, "Normal-Score Transformation (NST)"):
            fail(f"{label} missing professional NST heading")

    if "Best method from baseline panel" in (ROOT / "tables.md").read_text(encoding="utf-8", errors="ignore"):
        fail("tables.md still contains placeholder baseline-method text")
    if "Corresponding Author" in source_manuscript:
        fail("manuscript.md still uses `Corresponding Author`; use `Corresponding author`")
    if "+91-02717-404800" in source_manuscript:
        fail("manuscript.md still uses old phone hyphenation")
    if "supplement/cutoff_occupancy_uncertainty.csv" in source_manuscript:
        fail("manuscript.md still cites the old supplement subfolder path for cutoff occupancy")
    if "weathering upgrade" in (ROOT / "tables.md").read_text(encoding="utf-8", errors="ignore").lower():
        fail("tables.md still contains misleading weathering-upgrade row")
    if "physical-domain audit" not in source_manuscript.lower():
        fail("manuscript.md must disclose the physical-domain audit for unbounded SGS output")
    ai_idx = -1
    ref_idx = -1
    fig_caps_idx = -1
    figs_idx = -1
    for i, ln in enumerate(manuscript_lines):
        low = ln.lower()
        if "declaration of generative ai" in low and ai_idx == -1:
            ai_idx = i
        if "references" == low.strip() or re.match(r"^\s*\d+\.\s+references\s*$", low):
            ref_idx = i
        if low.strip() == "figure captions" or low.strip() == "figures captions":
            fig_caps_idx = i
        if low.strip() == "figures":
            figs_idx = i
    if ai_idx == -1:
        fail("Manuscript missing Declaration of Generative AI section")
    if ref_idx != -1 and ai_idx > ref_idx:
        fail("Declaration of Generative AI section must appear before References")
    if fig_caps_idx == -1 or figs_idx == -1:
        fail("Manuscript must include FIGURE CAPTIONS and FIGURES end sections")
    if figs_idx <= fig_caps_idx:
        fail("FIGURES section must appear after FIGURE CAPTIONS section")
    if re.search(r"(?im)^\s*Caption:\s*Figure\s+\d+", manuscript_text_full):
        fail("Duplicate figure captions detected in FIGURES section (remove repeated caption text)")

    if "first deposit-scale stochastic uncertainty study" in manuscript_text_full.lower():
        fail("Novelty wording is overstated; replace absolute first-claim phrasing")
    if ("screening-stage" not in manuscript_text_full.lower()) and ("screening level" not in manuscript_text_full.lower()):
        fail("Manuscript must clearly frame outputs as screening-stage/screening-level diagnostics")
    if re.search(r"`\s*Z_comp|`\s*w_i|`\s*p_i|`\s*P\(Z", manuscript_text_full):
        fail("Equation block still uses code-style backticks; equations must be editable publication text")

    table_caps: list[tuple[int, int]] = []
    for idx, ln in enumerate(manuscript_lines):
        m = re.match(r"^Table\s+(\d+)\.\s+", ln.strip(), flags=re.IGNORECASE)
        if m:
            table_caps.append((idx, int(m.group(1))))
    if not table_caps:
        fail("No table captions detected in manuscript")
    nums = [n for _, n in table_caps]
    expected = list(range(1, len(nums) + 1))
    if nums != expected:
        fail(f"Table numbering must be continuous without reset. Found {nums}, expected {expected}")
    if ref_idx != -1 and table_caps[0][0] <= ref_idx:
        fail("Tables must be placed after References in the manuscript end matter")
    if nums != [1, 2, 3, 4, 5]:
        fail(f"Expected exactly five end tables numbered 1-5, found {nums}")

    if ref_idx != -1:
        end_idx = len(manuscript_lines)
        for i, ln in enumerate(manuscript_lines):
            if ln.strip().lower() == "tables":
                end_idx = i
                break
        ref_lines = manuscript_lines[ref_idx + 1 : end_idx]
        for ln in ref_lines:
            core = ln.split("http", 1)[0]
            if re.search(r"\.\s+[A-Z][A-Za-z'`-]+,\s+[A-Z].*?\d{4}\.", core):
                fail("Reference list appears to contain merged entries; split into one reference per paragraph")

    for fig in figure_files:
        fig_path = sub_dir / fig
        size, dpi, compression = _check_figure_resolution(fig_path)
        if compression not in {"tiff_lzw", "tiff_adobe_deflate", "tiff_deflate"}:
            fail(f"{fig} must use lossless LZW or Deflate TIFF compression (detected {compression or 'unknown'})")
        min_dpi, min_width = _figure_minimums(fig)
        if min(dpi) < min_dpi:
            fail(f"{fig} DPI below {min_dpi} (detected {dpi[0]:.1f} x {dpi[1]:.1f})")
        if size[0] < min_width:
            fail(f"{fig} width below {min_width} px minimum guidance (detected width {size[0]})")
        width_mm = float(size[0]) / max(float(dpi[0]), 1e-9) * 25.4
        if abs(width_mm - 183.0) > 0.10:
            fail(
                f"{fig} physical width must be 183 mm (detected {width_mm:.3f} mm)"
            )
        if fig == "Fig06_TGC_Uncertainty_Spread_Map.tif" and size[0] <= size[1]:
            fail(
                "Fig06_TGC_Uncertainty_Spread_Map.tif must be the final landscape "
                "multi-panel spread/occupancy figure, not the older single-map portrait output"
            )


def load_truth(sub_dir: Path) -> dict:
    truth_path = sub_dir / "source_of_truth.submission.json"
    require(truth_path)
    return json.loads(truth_path.read_text(encoding="utf-8"))


def is_independent_profile(sub_dir: Path, truth: dict) -> bool:
    if str(truth.get("package_profile", "")).strip().lower() == "independent":
        return True
    return "independent" in sub_dir.name.lower()


def check_common_files(sub_dir: Path) -> tuple[Path, Path, Path, Path, Path]:
    paper = sub_dir / "paper.md"
    body = sub_dir / "paper_body.md"
    tables = sub_dir / "tables_final.md"
    captions = sub_dir / "figure_captions_final.md"
    paper_docx = sub_dir / "paper.docx"
    for path in [paper, body, tables, captions, paper_docx]:
        require(path)
    require(sub_dir / "references.bib")
    return paper, body, tables, captions, paper_docx


def check_independent_manuscript_docx(sub_dir: Path) -> Path:
    manuscript_docx = sub_dir / INDEPENDENT_MANUSCRIPT_DOCX
    require(manuscript_docx)
    return manuscript_docx


def check_refs_and_figures(sub_dir: Path, captions: Path, independent: bool) -> None:
    fig_refs = re.findall(r"\*\*Figure\s+([0-9]+[A-Z]?)\.", captions.read_text(encoding="utf-8"))
    if not fig_refs:
        fail("No figure captions found")
    figure_map = (
        {
            "1": "figure_1_regional_geology_map.png",
            "2": "variogram.png",
            "3": "histogram_validation.png",
            "4": "qq_plot.png",
            "5A": "swath_x.png",
            "5B": "swath_y.png",
            "5C": "swath_z.png",
            "6": "tonnage_risk_curve.png",
        }
        if independent
        else {
            "1": "figure_1_regional_geology_map.png",
            "2": "structural_anisotropy_prior.png",
            "3": "drill_sections_lithology_tgc.png",
            "4": "contact_weathering_tgc.png",
            "5": "spatial_uncertainty_products.png",
            "6": "tgc_uncertainty_spread_map.png",
            "7": "model_validation_limits.png",
        }
    )
    figs_dir = sub_dir / "figures"
    missing: list[str] = []
    for fid in sorted(set(fig_refs)):
        image_name = figure_map.get(fid)
        if not image_name:
            continue
        if not (figs_dir / image_name).exists():
            missing.append(f"Figure {fid} -> {image_name}")
    if missing:
        fail(f"Missing figure image files for captioned figures: {missing}")


def check_reproducibility_markers(sub_dir: Path, paper: Path, body: Path, independent: bool) -> None:
    source_md = sub_dir / "SOURCE_OF_TRUTH.md"
    checklist_md = sub_dir / "SUBMISSION_CHECKLIST.md"
    truth_json = sub_dir / "source_of_truth.submission.json"
    for path in [source_md, checklist_md, truth_json]:
        require(path)

    source_txt = source_md.read_text(encoding="utf-8")
    checklist_txt = checklist_md.read_text(encoding="utf-8")
    manuscript_txt = paper.read_text(encoding="utf-8") + "\n" + body.read_text(encoding="utf-8")
    merged = "\n".join([source_txt, checklist_txt, manuscript_txt])

    required_tokens = [
        "python -m src.run_all",
        "config/main_config.yaml",
        "sgs_checkpoint_state.json",
        "sgs_reals_checkpoint.npy",
        "sgs_reals_ns_checkpoint.npy",
        "python scripts/build_submission_package.py",
        "--run-dir",
        "--strict",
        "source_of_truth.submission.json",
        "Science-run regeneration command:",
        "Package regeneration command:",
    ]
    if independent:
        required_tokens.append("--independent")

    for token in required_tokens:
        if token not in merged:
            fail(f"Reproducibility marker missing from package docs: {token}")


def check_writing_implementation_plan(sub_dir: Path, paper: Path, body: Path) -> None:
    source_md = sub_dir / "SOURCE_OF_TRUTH.md"
    checklist_md = sub_dir / "SUBMISSION_CHECKLIST.md"
    merged = "\n".join(
        [
            source_md.read_text(encoding="utf-8"),
            checklist_md.read_text(encoding="utf-8"),
            paper.read_text(encoding="utf-8"),
            body.read_text(encoding="utf-8"),
        ]
    )
    required_tokens = [
        "Package Evidence Scope",
        "Completed evidence:",
        "Run-backed evidence tracks:",
        "Literature-Aligned Scope",
        "Package guard:",
        "Do not cite missing files, unrun sensitivity cases, or internal drafting notes as manuscript evidence.",
    ]
    for token in required_tokens:
        if token not in merged:
            fail(f"Package evidence-scope marker missing from package content: {token}")

    forbidden_tokens = [
        "Post-Finish Writing Implementation Plan",
        "Post-Run Manuscript Fill Template",
        "NotebookLM assistance rule",
        "Reviewer point-by-point closure matrix",
        "Author Implementation Instructions",
        "Implementation Protocol",
    ]
    for token in forbidden_tokens:
        if token.lower() in merged.lower():
            fail(f"Internal writing-plan wording present in package content: {token}")


def check_legacy_truth_numbers(truth: dict, paper: Path, body: Path, tables: Path, captions: Path) -> None:
    txt = (
        paper.read_text(encoding="utf-8")
        + "\n"
        + body.read_text(encoding="utf-8")
        + "\n"
        + tables.read_text(encoding="utf-8")
        + "\n"
        + captions.read_text(encoding="utf-8")
    )
    forbidden_patterns = [
        r"Post-Run Writing Implementation Plan",
        r"NotebookLM assistance rule",
        r"Reviewer point-by-point closure matrix",
        r"Author Implementation Instructions",
        r"Implementation Protocol",
        r"\bNot evaluated\b",
        r"cross_validation_(?:300|600|blocked_300)\.json",
        r"\bFigure\s+(?:8|9|10|11|12)\b",
        r"100\s*x\s*100\s*x\s*10",
        r"\b400[- ]realization",
        r"reviewer-first",
        r"_archive_",
    ]
    for pat in forbidden_patterns:
        if re.search(pat, txt, flags=re.IGNORECASE):
            fail(f"Stale or internal wording present in package content: /{pat}/")
    dims = truth["grid"]["dims"]
    support = truth["grid"]["cell_size_m"]
    risk = truth["risk_3pct"]
    metrics = truth["validation_metrics"]
    volume_rows = [
        f"| P10 uncertainty volume above screening cutoff (Mt-equivalent) | {risk['tonnage_mt']['p10']:.2f} |",
        f"| P50 uncertainty volume above screening cutoff (Mt-equivalent) | {risk['tonnage_mt']['p50']:.2f} |",
        f"| P90 uncertainty volume above screening cutoff (Mt-equivalent) | {risk['tonnage_mt']['p90']:.2f} |",
    ]
    swath_triplet = (
        f"{metrics['swath_corr_x']:.4f} / "
        f"{metrics['swath_corr_y']:.4f} / "
        f"{metrics['swath_corr_z']:.4f}"
    )
    expected_any = [
        [str(truth["simulation"]["n_real"])],
        [
            f"{support[0]:.0f} m x {support[1]:.0f} m x {support[2]:.0f} m",
            f"{support[0]:.0f} m × {support[1]:.0f} m × {support[2]:.0f} m",
        ],
        [f"{risk['density_t_per_m3']:.2f}"],
        ["full block volume x density", "full block volume × density"],
        [volume_rows[0]],
        [volume_rows[1]],
        [volume_rows[2]],
        [f"{risk['grade_pct']['p50']:.2f}"],
        [f"{metrics['hist_overlap']:.4f}"],
        [f"{metrics['qq_rmse']:.4f}"],
        [swath_triplet],
    ]
    for choices in expected_any:
        if not any(token in txt for token in choices):
            fail(f"Expected truth-aligned token not found: {choices[0]}")


def check_independent_content(truth: dict, paper: Path, body: Path, tables: Path, captions: Path) -> None:
    paper_txt = paper.read_text(encoding="utf-8")
    body_txt = body.read_text(encoding="utf-8")
    tables_txt = tables.read_text(encoding="utf-8")
    captions_txt = captions.read_text(encoding="utf-8")
    merged = "\n".join([paper_txt, body_txt, tables_txt, captions_txt])

    forbidden_patterns = [
        r"\breviewer\b",
        r"\bnotebooklm\b",
        r"reviewer-aligned",
        r"previous\s+paper",
        r"previous\s+draft",
        r"updated\s+document",
        r"revised\s+workflow",
        r"submission_ready/",
        r"\.tmp_build_",
    ]
    for pat in forbidden_patterns:
        if re.search(pat, merged, flags=re.IGNORECASE):
            fail(f"Forbidden standalone-package wording present: /{pat}/")

    sim_support = truth["grid"]["simulation_support_m"]
    rep_support = truth["grid"]["reporting_support_m"]
    sim_token = f"{sim_support[0]:.0f} m x {sim_support[1]:.0f} m x {sim_support[2]:.0f} m"
    rep_token = f"{rep_support[0]:.0f} m x {rep_support[1]:.0f} m x {rep_support[2]:.0f} m"
    required_tokens = [
        sim_token,
        rep_token,
        "unit total sill",
        "one nugget interpretation",
        "fixed neighborhood",
        "Top-cut | Not applied",
        "where Z_comp is",
        "vertical continuity remains the weakest",
        "Screening-Cutoff Uncertainty Envelope",
    ]
    for token in required_tokens:
        if token not in merged:
            fail(f"Independent package missing required token: {token}")

    if sim_support == rep_support:
        fail("Independent profile must use finer simulation support than reporting support")

    row = truth["risk_3pct"]["tonnage_mt"]
    triplet = f"{row['p10']:.2f}/{row['p50']:.2f}/{row['p90']:.2f}"
    if triplet not in paper_txt and triplet not in body_txt:
        fail(f"3% screening-cutoff triplet missing from manuscript: {triplet}")

    if re.search(r"\d+\.\d{4,}", tables_txt):
        fail("Independent tables contain excessive decimal precision (>3 decimals)")

    if re.search(r"\|\s*0\.0\s*\|", tables_txt) or re.search(r"\|\s*2\.0\s*\|", tables_txt):
        fail("Independent tables include cutoff rows outside the practical decision range")


def check_zip(sub_dir: Path, zip_path: Path, independent: bool) -> None:
    require(zip_path)
    with zipfile.ZipFile(zip_path) as zf:
        names = set(zf.namelist())
    required = {
        "source_of_truth.submission.json",
        "references.bib",
        "paper.md",
        "paper_body.md",
        "paper.docx",
        "tables_final.md",
        "figure_captions_final.md",
    }
    if independent:
        required.add(INDEPENDENT_MANUSCRIPT_DOCX)
        required.update(
            {
                "figures/figure_1_regional_geology_map.png",
                "figures/variogram.png",
                "figures/histogram_validation.png",
                "figures/qq_plot.png",
                "figures/spatial_uncertainty_products.png",
                "figures/model_validation_limits.png",
                "supplement/cutoff_occupancy_uncertainty.csv",
                "supplement/validation_metrics.json",
                "supplement/variogram_model.json",
                "supplement/sgs_meta.json",
            }
        )
    else:
        required.update({"figures/figure_1_regional_geology_map.png"})
    missing = sorted(required - names)
    if missing:
        fail(f"Zip missing required entries: {missing}")



MME_TITLE = "Geological Support and Reporting-Envelope Effects on Grade Uncertainty in a Tanzanian Stratiform Graphite System"
MME_JOURNAL = "Mining, Metallurgy & Exploration"
MME_COLLECTION = "Industrial Minerals: Geology, Extraction and Use"
MME_ORCID = "0009-0001-5030-7524"


def _write_mme_report(issues: list[str]) -> None:
    path = ROOT / "build" / "submission_work" / "preflight_report.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "status": "pass" if not issues else "fail",
        "package_profile": "mme",
        "run_dir": str(ROOT / "output" / "a3_geology_aligned_250_200_20_nr100"),
        "checks": {
            "missing_required_items": [],
            "content_issues": issues,
            "visual_issues": [],
            "final_submission_issues": [],
        },
    }
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def _xlsx_text_and_sheets(path: Path) -> tuple[str, set[str]]:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, read_only=True, data_only=False)
        parts: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        parts.append(str(cell.value))
        sheets = set(wb.sheetnames)
        wb.close()
        return "\n".join(parts), sheets
    except Exception as exc:
        return f"XLSX_READ_ERROR: {exc}", set()


def _pdf_text(path: Path) -> str:
    try:
        from PyPDF2 import PdfReader
        return "\n".join((page.extract_text() or "") for page in PdfReader(str(path)).pages)
    except Exception:
        try:
            from pypdf import PdfReader
            return "\n".join((page.extract_text() or "") for page in PdfReader(str(path)).pages)
        except Exception:
            return ""


def check_mme_package(sub_dir: Path) -> None:
    issues: list[str] = []
    actual = {p.name for p in sub_dir.iterdir() if p.is_file()}
    expected = set(CLEAN_REQUIRED_FILES)
    for name in sorted(expected - actual): issues.append(f"missing required upload file: {name}")
    for name in sorted(actual - expected): issues.append(f"unexpected upload file: {name}")
    if expected - actual:
        _write_mme_report(issues)
        fail("; ".join(issues))

    manuscript = sub_dir / "Manuscript.docx"
    cover = sub_dir / "Cover_Letter.docx"
    text = _docx_text(manuscript)
    cover_text = _docx_text(cover)
    if MME_TITLE not in text: issues.append("working title missing from manuscript")
    if MME_TITLE not in cover_text: issues.append("working title missing from cover letter")
    if MME_ORCID not in text: issues.append("ORCID missing from manuscript title page")
    if MME_ORCID not in cover_text: issues.append("ORCID missing from cover letter")
    for token in ("Sudipta Chanda", "sudipta.chanda@sakariya.in", "Sakariya Mines and Minerals"):
        if token not in text: issues.append(f"title-page author information missing: {token}")
    abstract, keywords = _extract_abstract_and_keywords(manuscript)
    n_words = _count_abstract_words(abstract)
    if not 150 <= n_words <= 250: issues.append(f"abstract has {n_words} words; MME requires 150-250")
    if not 4 <= len(keywords) <= 6: issues.append(f"keyword count is {len(keywords)}; MME requires 4-6")

    required_sections = ["Introduction", "Geological Setting", "Data and Methods", "Results", "Discussion",
                         "Conclusions", "Statements and Declarations", "Competing Interests", "Funding", "Ethics Approval", "Consent to Participate", "Consent for Publication",
                         "Author Contributions", "Data Availability", "Acknowledgements", "References"]
    for heading in required_sections:
        if heading.lower() not in text.lower(): issues.append(f"required manuscript section missing: {heading}")
    if "generative ai" not in text.lower() or "methods" not in text.lower():
        issues.append("generative-AI assistance is not documented in Methods")
    if "Online Resource 1" not in text or "Online Resource 2" not in text:
        issues.append("both Online Resources must be cited and captioned in the manuscript")
    doc = Document(manuscript)
    caption_paragraphs = [p for p in doc.paragraphs if re.match(r"^Fig\.\s*\d+\s+", p.text.strip())]
    caption_numbers = [int(re.match(r"^Fig\.\s*(\d+)", p.text.strip()).group(1)) for p in caption_paragraphs]
    if caption_numbers != list(range(1, 8)):
        issues.append(f"manuscript figure captions are incomplete or out of order: {caption_numbers}")
    for paragraph in caption_paragraphs:
        if paragraph.text.rstrip().endswith("."):
            issues.append(f"Springer caption has terminal punctuation: {paragraph.text[:24]}")
        if paragraph.style is not None and paragraph.style.name != "Caption":
            issues.append(f"figure caption does not use Caption style: {paragraph.text[:24]}")
    practical_markers = (
        "Implications for Graphite Exploration and Resource Evaluation",
        "Practical Decision-Use Matrix for Graphite Exploration and Resource Evaluation",
        "Archive-derived lode envelope", "Envelope-weighted P50 TGC", "P90-P10 TGC spread",
        "P(TGC > 3%)", "Raw categorical frequencies and entropy", "Geology-blind null sensitivity",
    )
    for marker in practical_markers:
        if marker not in text:
            issues.append(f"practical mining-value content missing: {marker}")
    if _docx_media_count(manuscript) != 7: issues.append("Manuscript.docx must embed exactly seven figures")

    doc = Document(manuscript)
    limitation_indices = [
        index for index, paragraph in enumerate(doc.paragraphs)
        if "limitations and future validation" in paragraph.text.lower()
    ]
    if len(limitation_indices) != 1:
        issues.append(f"manuscript must contain exactly one limitations section; found {len(limitation_indices)}")
    else:
        limitation_paragraphs: list[str] = []
        for paragraph in doc.paragraphs[limitation_indices[0] + 1:]:
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.lower().startswith("heading"):
                break
            if paragraph.text.strip():
                limitation_paragraphs.append(paragraph.text.strip())
        if len(limitation_paragraphs) != 1:
            issues.append(f"limitations section must contain one paragraph; found {len(limitation_paragraphs)}")
    if "principal remaining scientific test" in text.lower():
        issues.append("limitations are repeated in Conclusions instead of being consolidated in Discussion")
    null_design_markers = (
        "Geology-Blind Composite Null Sensitivity",
        "composite configuration sensitivity",
        "150 x 150 x 150 m",
        "105/15/195 degree axis labels",
        "enabled vertical trend",
    )
    for marker in null_design_markers:
        if marker not in text:
            issues.append(f"composite-null design disclosure missing: {marker}")
    heading_levels = set()
    for para in doc.paragraphs:
        name = para.style.name if para.style is not None else ""
        match = re.match(r"Heading\s+(\d+)", name, flags=re.I)
        if match: heading_levels.add(int(match.group(1)))
    if any(level > 3 for level in heading_levels): issues.append("manuscript uses more than three displayed heading levels")
    normal = doc.styles["Normal"]
    if normal.font.size is not None and abs(normal.font.size.pt - 10.0) > 0.2:
        issues.append(f"Normal style is {normal.font.size.pt:.1f} pt; MME requests 10-point plain text")
    if not _docx_is_single_column(manuscript): issues.append("manuscript is not single-column")

    lower = (text + "\n" + cover_text).lower()
    for residue in ("journal of african earth sciences", "jaes", "elsevier", "highlights.docx",
                    "supplementary_data_s2", "supplementary data s2", " in s2", "public s2",
                    "submission_final_jaes"):
        if residue in lower: issues.append(f"JAES/Elsevier residue present: {residue}")
    for leak in ("c:\\users", "onedrive", "desktop\\tanga_new", "/users/"):
        if leak in lower: issues.append(f"local path leakage present: {leak}")

    if re.search(r"\([A-Z][A-Za-z-]+(?:\s+et al\.)?,?\s+(?:19|20)\d{2}[a-z]?\)", text) or re.search(r"\b[A-Z][A-Za-z-]+(?:\s+et al\.|\s+and\s+[A-Z][A-Za-z-]+)\s+\((?:19|20)\d{2}\)", text):
        issues.append("author-date citation remains; MME requires numbered square brackets")
    if chr(0xC2) + chr(0xB0) in text or "\ufffd" in text:
        issues.append("encoding artefact remains in manuscript text")
    abstract_pos = text.lower().find("abstract")
    for token in ("Acknowledgements:", "Author contributions:"):
        token_pos = text.find(token)
        if token_pos < 0 or (abstract_pos >= 0 and token_pos > abstract_pos):
            issues.append(f"title-page metadata missing before Abstract: {token}")
    body_text = text[:text.lower().rfind("references")] if "references" in text.lower() else text
    cited: set[int] = set()
    for group in re.findall(r"\[([0-9,;\-\s]+)\]", body_text):
        for part in re.split(r"[,;]", group):
            part = part.strip()
            if re.fullmatch(r"\d+", part):
                cited.add(int(part))
            elif re.fullmatch(r"\d+\s*-\s*\d+", part):
                left, right = (int(x.strip()) for x in part.split("-", 1))
                cited.update(range(min(left, right), max(left, right) + 1))
    ref_lines = [line for line in _docx_paragraph_lines(manuscript) if re.match(r"^\[\d+\]\s+", line)]
    ref_numbers = {int(re.match(r"^\[(\d+)\]", line).group(1)) for line in ref_lines}
    if not cited: issues.append("no numbered in-text citations detected")
    if ref_numbers and ref_numbers != set(range(1, max(ref_numbers)+1)): issues.append("reference list numbering is not consecutive")
    if ref_numbers and not cited.issubset(ref_numbers): issues.append("one or more numbered citations have no reference-list entry")
    if ref_numbers and not ref_numbers.issubset(cited): issues.append("one or more reference-list entries are uncited")
    if ref_lines and any(not re.search(r"\((?:19|20)\d{2}\)", line) for line in ref_lines):
        issues.append("one or more references do not use Springer year-in-parentheses formatting")
    if ref_lines and any(re.search(r",\s*(?:19|20)\d{2}\.", line) for line in ref_lines):
        issues.append("legacy author-comma-year reference formatting remains")
    if ref_lines and any("https://doi.org/" not in line for line in ref_lines if "doi" in line.lower()):
        issues.append("a DOI-bearing reference is not formatted as a full https://doi.org/ link")

    try:
        from PIL import Image
        for idx in range(1,8):
            path = sub_dir / f"Fig{idx}.tif"
            with Image.open(path) as image:
                dpi = image.info.get("dpi", (0,0)); dx = float(dpi[0] or 0); dy = float(dpi[1] or 0)
                width_mm = image.width / dx * 25.4 if dx else 0
                height_mm = image.height / dy * 25.4 if dy else 999
                compression = str(image.info.get("compression", "")).lower()
                if image.mode != "RGB": issues.append(f"Fig{idx}.tif is {image.mode}, not RGB")
                if dx < 600 or dy < 600: issues.append(f"Fig{idx}.tif resolution below 600 dpi")
                if width_mm > 174.2 or height_mm > 234.2: issues.append(f"Fig{idx}.tif exceeds 174 x 234 mm ({width_mm:.1f} x {height_mm:.1f})")
                if compression not in {"tiff_lzw","tiff_adobe_deflate","tiff_deflate","lzw","deflate"}:
                    issues.append(f"Fig{idx}.tif does not use verified lossless LZW/Deflate compression")
    except Exception as exc:
        issues.append(f"TIFF inspection failed: {exc}")

    pdf_text = _pdf_text(sub_dir / "ESM_1.pdf")
    xlsx_text, sheets = _xlsx_text_and_sheets(sub_dir / "ESM_2.xlsx")
    pdf_tokens = (MME_TITLE, MME_JOURNAL, "Sudipta Chanda", MME_ORCID)
    xlsx_tokens = (MME_TITLE, MME_JOURNAL, "Sudipta Chanda", MME_ORCID)
    pdf_compact = re.sub(r"[^a-z0-9]+", "", pdf_text.lower())
    for token in pdf_tokens:
        token_compact = re.sub(r"[^a-z0-9]+", "", token.lower())
        if token_compact not in pdf_compact: issues.append(f"ESM_1.pdf missing metadata text: {token}")
    for token in xlsx_tokens:
        if token not in xlsx_text: issues.append(f"ESM_2.xlsx missing metadata text: {token}")
    esm_lower = (pdf_text + "\n" + xlsx_text).lower()
    for residue in ("journal of african earth sciences", "jaes", "elsevier", "supplementary_data_s2"):
        if residue in esm_lower:
            issues.append(f"JAES/Elsevier residue present in Online Resources: {residue}")
    required_sheets = {"README","Run Metadata","Validation Metrics","Variogram Models","Convergence",
                       "Reporting Envelope","Support Decomposition","Categorical Validation","Contact Statistics",
                       "Occupancy Diagnostics","Repeated Null Summary","Repeated Null Seeds"}
    if not required_sheets.issubset(sheets): issues.append(f"ESM_2.xlsx missing sheets: {sorted(required_sheets-sheets)}")

    summary_path = ROOT / "build" / "factorial_validation" / "five_seed_summary.json"
    metrics_path = ROOT / "build" / "factorial_validation" / "five_seed_metrics.csv"
    if not summary_path.exists() or not metrics_path.exists():
        issues.append("five independent no-domain seed families are not complete")
    else:
        try:
            import csv
            import math
            rows = list(csv.DictReader(metrics_path.open("r", encoding="utf-8-sig", newline="")))
            expected_seeds = {9101, 9201, 9301, 9401, 9501}
            required_metrics = {"mean_sim", "hist_overlap", "qq_rmse", "swath_corr_x", "swath_corr_y", "swath_corr_z"}
            if len(rows) != 5:
                issues.append(f"repeated-null metrics contain {len(rows)} rows; expected 5")
            columns = set(rows[0]) if rows else set()
            missing_columns = ({"seed"} | required_metrics) - columns
            if missing_columns:
                issues.append(f"repeated-null metrics missing columns: {sorted(missing_columns)}")
            seeds = {int(float(row["seed"])) for row in rows if row.get("seed")}
            if seeds != expected_seeds:
                issues.append(f"repeated-null seed set is incomplete: {sorted(seeds)}")
            for row in rows:
                for metric in required_metrics:
                    if metric not in row or not math.isfinite(float(row[metric])):
                        issues.append(f"repeated-null seed {row.get('seed','?')} has invalid {metric}")
            summary = json.loads(summary_path.read_text(encoding="utf-8"))
            if set(int(seed) for seed in summary.get("seeds", [])) != expected_seeds:
                issues.append("five_seed_summary.json does not list the required seed set")
            summary_metrics = summary.get("metrics", {}) or {}
            for metric in required_metrics:
                metric_summary = summary_metrics.get(metric, {}) or {}
                for stat_name in ("median", "min", "max", "std"):
                    value = metric_summary.get(stat_name)
                    if value is None or not math.isfinite(float(value)):
                        issues.append(f"five_seed_summary.json missing finite {metric}.{stat_name}")
            completed_markers = (
                "Five independent no-domain isotropic families were completed with seeds",
                "Repetition across all five seeds establishes the robustness of this global-fit behaviour",
            )
            for marker in completed_markers:
                if marker not in text:
                    issues.append(f"completed five-seed evidence is not integrated in manuscript: {marker}")
            for stale in ("one independent 20-realisation family", "should be completed before directional results"):
                if stale in text:
                    issues.append(f"stale pending null-campaign wording remains after completion: {stale}")
            for token in ("Independent no-domain seed metrics", "Median", "Minimum", "Maximum", "SD"):
                if token not in pdf_text:
                    issues.append(f"ESM_1.pdf missing completed seed-table element: {token}")
            for seed in sorted(expected_seeds):
                if str(seed) not in xlsx_text:
                    issues.append(f"ESM_2.xlsx missing repeated-null seed: {seed}")
        except Exception as exc:
            issues.append(f"repeated-null metrics could not be verified: {exc}")

    _write_mme_report(issues)
    if issues: fail("; ".join(issues))
    print("OK: MME submission preflight passed")

def main() -> None:
    parser = argparse.ArgumentParser(description="Run strict checks on a built submission package.")
    parser.add_argument("--sub-dir", default="submission", help="Submission directory to validate")
    args = parser.parse_args()

    sub_dir = Path(args.sub_dir)
    if not sub_dir.is_absolute():
        sub_dir = ROOT / sub_dir

    if (sub_dir / "Manuscript.docx").exists():
        check_mme_package(sub_dir)
        return

    if (sub_dir / "04_Manuscript.docx").exists():
        check_clean_jaes_package(sub_dir)
        print("OK: legacy clean JAES submission preflight passed")
        return

    truth = load_truth(sub_dir)
    paper, body, tables, captions, _paper_docx = check_common_files(sub_dir)
    independent = is_independent_profile(sub_dir, truth)

    if independent:
        zip_path = ROOT / "submission_package_independent_clean.zip"
        check_independent_manuscript_docx(sub_dir)
        check_independent_content(truth, paper, body, tables, captions)
    else:
        zip_path = ROOT / "submission_package_final_clean.zip"
        check_legacy_truth_numbers(truth, paper, body, tables, captions)

    check_reproducibility_markers(sub_dir, paper, body, independent)
    check_writing_implementation_plan(sub_dir, paper, body)
    check_refs_and_figures(sub_dir, captions, independent)
    check_zip(sub_dir, zip_path, independent)
    print("OK: submission preflight passed")


if __name__ == "__main__":
    main()
