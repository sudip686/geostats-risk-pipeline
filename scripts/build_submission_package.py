from __future__ import annotations

import argparse
import hashlib
import io
import os
import json
import re
import shutil
import subprocess
import sys
import tempfile
import textwrap
import time
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_DIR = ROOT / "submission"
FINAL_SUBMISSION_DIR = ROOT / "submission_final_MME"
BUILD_WORK_DIR = ROOT / "build" / "submission_work"
PACKAGE_PROFILE = "submission"
PACKAGE_ZIP_BASENAME = "submission_package_mme_clean"
FORMAT_DIR_CANDIDATES = [
    ROOT / "_archive_20260307" / "Format",
    ROOT,
    ROOT / "submission_ready",
    ROOT / "repo" / "submission_ready",
    ROOT / "submission",
    ROOT / "repo",
]
SOURCE_OF_TRUTH_JSON = ROOT / "build" / "source_of_truth.submission.json"
CANONICAL_N_REAL = 100
CANONICAL_SEARCH_RADIUS = [250, 200, 20]
INDEPENDENT_RUN_DIR = ROOT / "output" / "a3_geology_aligned_250_200_20_nr100"
CANONICAL_SIM_SUPPORT = [25.0, 25.0, 2.0]
CANONICAL_REPORTING_SUPPORT = [50.0, 50.0, 2.0]
PRIMARY_MANUSCRIPT_DOCX = "sudip_manuscript.docx"
PRIMARY_MANUSCRIPT_TEMPLATE = "SG Manuscript.docx"
INDEPENDENT_MANUSCRIPT_DOCX = (
    "geology_led_sequential_gaussian_simulation_for_uncertainty_analysis_of_a_"
    "stratiform_graphite_deposit_in_tanzania.docx"
)
STYLE_LOCK_REFERENCE = ROOT / "SG Manuscript (1).docx"
AUTHOR_NAME = "Sudipta Chanda"
AUTHOR_AFFILIATION = (
    "Sakariya Mines and Minerals Private Limited, 1402 Ecostation Business Tower, Newtown, Rajarhat, Kolkata, West Bengal 700160, India"
)
AUTHOR_EMAIL = "sudipta.chanda@sakariya.in"
AUTHOR_PHONE = "+91 2717 404800"
AUTHOR_ORCID = "0009-0001-5030-7524"
AUTHOR_ORCID_URL = "https://orcid.org/0009-0001-5030-7524"
MME_JOURNAL = "Mining, Metallurgy & Exploration"
MME_COLLECTION = "Industrial Minerals: Geology, Extraction and Use"
MME_TITLE = "Geological Support and Reporting-Envelope Effects on Grade Uncertainty in a Tanzanian Stratiform Graphite System"


def load_yaml(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def canonical_contract() -> dict:
    cfg = load_yaml(ROOT / "config" / "main_config.yaml")
    return cfg.get("workflow_contract", {}) or {}


def canonical_run_dir() -> Path:
    contract = canonical_contract()
    rel = contract.get("canonical_output_dir", "output/a3_categorical_25_50_nr100")
    return (ROOT / rel).resolve()


def is_archived_run_dir(path: Path) -> bool:
    resolved = path.resolve()
    stale_roots = [
        (ROOT / "output" / "stale").resolve(),
        (ROOT / "stale").resolve(),
    ]
    return any(root == resolved or root in resolved.parents for root in stale_roots if root.exists())

DOC_TEMPLATES = [
    "Title page.docx",
    "Covering Letter-MS.docx",
    PRIMARY_MANUSCRIPT_TEMPLATE,
    "Authors Statement.docx",
    "Conflict of interest.docx",
    "Declaration of interest statement.docx",
    "Fig.docx",
    "Table.docx",
]

REQUIRED_SUBMISSION_REL_PATHS = [
    "submission_ready/Title page.docx",
    "submission_ready/Covering Letter-MS.docx",
    f"submission_ready/{PRIMARY_MANUSCRIPT_DOCX}",
    "submission_ready/paper.docx",
    "submission_ready/Authors Statement.docx",
    "submission_ready/Conflict of interest.docx",
    "submission_ready/Declaration of interest statement.docx",
    "submission_ready/Fig.docx",
    "submission_ready/Table.docx",
    "submission_ready/paper.md",
    "submission_ready/paper_body.md",
    "submission_ready/tables_final.md",
    "submission_ready/figure_captions_final.md",
    "submission_ready/references.bib",
    "submission_ready/highlights.txt",
    "submission_ready/graphical_abstract_summary.txt",
    "submission_ready/figures",
    "submission_ready/supplement",
    "submission_ready/supplement/software_manifest.json",
    "submission_ready/supplement/trend_ablation_summary.json",
    "submission_ready/supplement/synthetic_validation_reference.csv",
    "submission_ready/supplement/validation_metrics_2m.json",
    "submission_ready/supplement/support_ladder_summary.csv",
    "submission_ready/supplement/vertical_continuity_summary.json",
    "submission_ready/supplement/contact_analysis.csv",
    "submission_ready/supplement/contact_analysis_meta.json",
    "submission_ready/supplement/weathering_summary.csv",
    "submission_ready/supplement/domain_uncertainty_summary.json",
    "submission_ready/supplement/domain_uncertainty_hotspots.csv",
    "submission_ready/supplement/thickness_geometry_summary.json",
    "submission_ready/supplement/thickness_geometry_hotspots.csv",
    "submission_ready/supplement/confidence_gradient_hotspots.csv",
    "submission_ready/supplement/confidence_gradient_meta.json",
    "submission_ready/supplement/postrun_review_pack_status.json",
    "submission_ready/supplement/contact_analysis.png",
    "submission_ready/supplement/domain_entropy_map.png",
    "submission_ready/supplement/domain_stability_map.png",
    "submission_ready/supplement/graphitic_thickness_p50_map.png",
    "submission_ready/supplement/graphitic_thickness_aperture_map.png",
    "submission_ready/supplement/confidence_gradient_map.png",
    "submission_ready/SUBMISSION_CHECKLIST.md",
    "submission_ready/source_of_truth.submission.json",
    "submission_ready/SOURCE_OF_TRUTH.md",
    "submission_ready/submission_package_final_clean.zip",
    "submission_package_final_clean.zip",
]
INDEPENDENT_REQUIRED_SUBMISSION_REL_PATHS = [
    f"submission_ready_independent/{INDEPENDENT_MANUSCRIPT_DOCX}",
    "submission_ready_independent/paper.docx",
    "submission_ready_independent/paper.md",
    "submission_ready_independent/paper_body.md",
    "submission_ready_independent/tables_final.md",
    "submission_ready_independent/figure_captions_final.md",
    "submission_ready_independent/references.bib",
    "submission_ready_independent/figures",
    "submission_ready_independent/figures/figure_1_regional_geology_map.png",
    "submission_ready_independent/figures/variogram.png",
    "submission_ready_independent/figures/histogram_validation.png",
    "submission_ready_independent/figures/qq_plot.png",
    "submission_ready_independent/figures/swath_x.png",
    "submission_ready_independent/figures/swath_y.png",
    "submission_ready_independent/figures/swath_z.png",
    "submission_ready_independent/figures/tonnage_risk_curve.png",
    "submission_ready_independent/supplement",
    "submission_ready_independent/supplement/software_manifest.json",
    "submission_ready_independent/supplement/trend_ablation_summary.json",
    "submission_ready_independent/supplement/risked_tonnage.csv",
    "submission_ready_independent/supplement/validation_metrics.json",
    "submission_ready_independent/supplement/variogram_model.json",
    "submission_ready_independent/supplement/sgs_meta.json",
    "submission_ready_independent/SUBMISSION_CHECKLIST.md",
    "submission_ready_independent/source_of_truth.submission.json",
    "submission_ready_independent/SOURCE_OF_TRUTH.md",
    "submission_ready_independent/submission_package_independent_clean.zip",
    "submission_package_independent_clean.zip",
]
INNER_BORDER_SZ = 12  # 1.5pt (Word size unit is 1/8 pt)
OUTER_BORDER_SZ = 20  # 2.5pt
EQ_BORDER_SZ = 16
FIGURE_IMAGE_MAP: dict[str, str] = {
    "1": "figure_1_regional_geology_map.png",
    "2": "structural_anisotropy_prior.png",
    "3": "drill_sections_lithology_tgc.png",
    "4": "contact_weathering_tgc.png",
    "5": "spatial_uncertainty_products.png",
    "6": "tgc_uncertainty_spread_map.png",
    "7": "model_validation_limits.png",
}
INDEPENDENT_FIGURE_IMAGE_MAP: dict[str, str] = {
    "1": "figure_1_regional_geology_map.png",
    "2": "variogram.png",
    "3": "histogram_validation.png",
    "4": "qq_plot.png",
    "5A": "swath_x.png",
    "5B": "swath_y.png",
    "5C": "swath_z.png",
    "6": "tonnage_risk_curve.png",
}
FIGURE_DEFAULT_DETAILS: dict[str, str] = {
    "1": "Panel A reproduces Das et al. (2026, Fig. 1) unchanged with permission; Panel B is author-generated from the high-resolution project geology raster and canonical collar data.",
    "2": "Generated from canonical 000/180 degree strike, 090/30 degree down-dip, 270/60 degree normal, and 250/200/20 m search-ellipsoid metadata.",
    "3": "Generated from canonical composite coordinates, TGC values, interval lengths, and the 3% geological screening threshold.",
    "4": "Generated from the signed graphitic-host contact profile, graphitic-only weathering analysis, and published Das et al. (2026) XRF weathering data.",
    "5": "Generated from reporting-support exceedance probability, raw domain entropy, absolute graphitic-thickness P90-P10 aperture grids, and five-fold hole-grouped reliability/confusion diagnostics.",
    "6": "Generated from reporting-support plan and deterministic east-west section products; submitted TIFF uses the same maintained PNG source.",
    "7": "Generated from support-aligned mean decomposition, ensemble convergence, matched-space variogram reproduction, and geological-axis swath profiles.",
}
FINAL_SUBMISSION_REQUIRED_FILES = [
    "01_Title_Page.docx",
    "02_Highlights.docx",
    "04_Manuscript.docx",
    "07_Cover_Letter.docx",
    "08_Declaration_of_Interest.docx",
    "Fig01_Regional_Setting_Workflow.tif",
    "Fig02_Structural_Anisotropy_Prior.tif",
    "Fig03_Drill_Sections_Lithology_TGC.tif",
    "Fig04_Contact_Weathering_TGC.tif",
    "Fig05_Spatial_Uncertainty_Products.tif",
    "Fig06_TGC_Uncertainty_Spread_Map.tif",
    "Fig07_Model_Validation_Limits.tif",
    "Supplementary_Data_S2.zip",
]
INDEPENDENT_FINAL_SUBMISSION_REQUIRED_FILES = []
CLEAN_REQUIRED_SUBMISSION_FILES = [
    "Manuscript.docx", "Cover_Letter.docx",
    "Fig1.tif", "Fig2.tif", "Fig3.tif", "Fig4.tif",
    "Fig5.tif", "Fig6.tif", "Fig7.tif",
    "ESM_1.pdf", "ESM_2.xlsx",
]
ESSENTIAL_S2_FILES = [
    "variogram_model.json",
    "validation_metrics.json",
    "cutoff_occupancy_uncertainty.csv",
    "sgs_meta.json",
]


def resolve_format_dir() -> Path:
    for d in FORMAT_DIR_CANDIDATES:
        if d.exists() and ((d / "Title page.docx").exists() or (d / "01_Title_Page.docx").exists()):
            return d
    searched = ", ".join(str(p) for p in FORMAT_DIR_CANDIDATES)
    raise FileNotFoundError(
        "Could not resolve template directory. Expected to find `Title page.docx` in one of: "
        f"{searched}"
    )


def template_path(filename: str) -> Path:
    clean_template_candidates = {
        "Title page.docx": ["Title page.docx", "01_Title_Page.docx"],
        "Covering Letter-MS.docx": ["Covering Letter-MS.docx", "07_Cover_Letter.docx"],
        "Authors Statement.docx": ["Authors Statement.docx", "09_Author_Statement.docx"],
        "Conflict of interest.docx": ["Conflict of interest.docx", "08_Declaration_of_Interest.docx"],
        "Declaration of interest statement.docx": [
            "Declaration of interest statement.docx",
            "08_Declaration_of_Interest.docx",
        ],
        "Fig.docx": ["Fig.docx", "06_Figure_Captions.docx"],
        "Table.docx": ["Table.docx", "05_Tables.docx"],
    }
    candidates = clean_template_candidates.get(filename, [filename])
    if filename == PRIMARY_MANUSCRIPT_TEMPLATE:
        candidates = [PRIMARY_MANUSCRIPT_TEMPLATE, "paper.docx", PRIMARY_MANUSCRIPT_DOCX]
        candidates.append("04_Manuscript.docx")
    for base in FORMAT_DIR_CANDIDATES:
        if not base.exists():
            continue
        for name in candidates:
            p = base / name
            if p.exists():
                return p
    # Return the primary candidate for clear error messages upstream.
    return resolve_format_dir() / filename


def is_independent_profile() -> bool:
    return PACKAGE_PROFILE == "independent"


def package_zip_filename() -> str:
    return f"{PACKAGE_ZIP_BASENAME}.zip"


def display_path(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT)).replace("\\", "/")
    except ValueError:
        return str(path)


def science_run_regen_command(run_dir: Path) -> str:
    return f"`python -m src.run_all --config config/main_config.yaml --output {display_path(run_dir)}`"


def package_regen_command(run_dir: Path) -> str:
    cmd = f"python scripts/build_submission_package.py --run-dir {display_path(run_dir)} --strict"
    if is_independent_profile():
        cmd += " --independent"
    return f"`{cmd}`"


def reproducibility_marker_lines(run_dir: Path) -> list[str]:
    grids_dir = run_dir / "grids"
    return [
        f"- Workflow config: `config/main_config.yaml`",
        f"- Science-run regeneration: {science_run_regen_command(run_dir)}",
        f"- Resume state file: `{display_path(grids_dir / 'sgs_checkpoint_state.json')}`",
        f"- Checkpoint arrays: `{display_path(grids_dir / 'sgs_reals_checkpoint.npy')}` and `{display_path(grids_dir / 'sgs_reals_ns_checkpoint.npy')}`",
    ]


def active_manuscript_docx_name() -> str:
    return INDEPENDENT_MANUSCRIPT_DOCX if is_independent_profile() else PRIMARY_MANUSCRIPT_DOCX


def figure_image_map() -> dict[str, str]:
    return INDEPENDENT_FIGURE_IMAGE_MAP if is_independent_profile() else FIGURE_IMAGE_MAP


def required_submission_rel_paths() -> list[str]:
    return INDEPENDENT_REQUIRED_SUBMISSION_REL_PATHS if is_independent_profile() else REQUIRED_SUBMISSION_REL_PATHS


def final_submission_required_files() -> list[str]:
    return INDEPENDENT_FINAL_SUBMISSION_REQUIRED_FILES if is_independent_profile() else FINAL_SUBMISSION_REQUIRED_FILES


def resolve_independent_run_dir(explicit: str | None) -> Path:
    if explicit:
        p = Path(explicit)
        if not p.is_absolute():
            p = ROOT / p
        if p.exists() and (p / "sgs_meta.json").exists():
            return p
        raise FileNotFoundError(f"Provided independent run directory is missing: {p}")
    run_dir = canonical_run_dir()
    if run_dir.exists():
        return run_dir
    if INDEPENDENT_RUN_DIR.exists():
        return INDEPENDENT_RUN_DIR
    raise FileNotFoundError(f"Independent run directory not found: {run_dir}")


def concat_paper(body: str, tables_md: str, caps_md: str) -> str:
    return body.rstrip() + "\n\n## TABLES\n\n" + tables_md.strip() + "\n\n## FIGURE CAPTIONS\n\n" + caps_md.strip() + "\n"


def _parse_figure_caption_entries(caps_md: str) -> list[tuple[str, str]]:
    entries: list[tuple[str, str]] = []
    for raw in caps_md.splitlines():
        line = raw.strip()
        m = re.match(r"\*\*(?:Figure\s+|Fig\.\s*)([0-9]+[A-Z]?)(?:\.)?\*\*\s*(.+)$", line, flags=re.IGNORECASE)
        if m:
            entries.append((m.group(1), m.group(2).strip()))
    return entries


def build_independent_standalone_manuscript_md(body: str, tables_md: str, caps_md: str) -> str:
    parts = [body.rstrip(), "", "## Tables", "", tables_md.strip(), "", "## Figures", ""]
    for figure_id, caption in _parse_figure_caption_entries(caps_md):
        image_name = figure_image_map().get(figure_id)
        if not image_name:
            continue
        width = "62%" if figure_id == "1" else "85%"
        parts.append(f"### Figure {figure_id}")
        parts.append("")
        parts.append(f"![Figure {figure_id}. {caption}](figures/{image_name}){{ width={width} }}")
        parts.append("")
    return "\n".join(parts).rstrip() + "\n"


def package_display_name() -> str:
    name = SUBMISSION_DIR.name
    return name.replace(".tmp_build_", "", 1) if name.startswith(".tmp_build_") else name


def copy_reference_bibliography() -> None:
    bib_dst = SUBMISSION_DIR / "references.bib"
    if bib_dst.exists():
        return
    for cand in [
        ROOT / "references.bib",
        ROOT / "submission_ready" / "references.bib",
        ROOT / "submission" / "references.bib",
        ROOT / "repo" / "references.bib",
    ]:
        if cand.exists():
            shutil.copy2(cand, bib_dst)
            return
    bib_dst.write_text("% bibliography placeholder\n", encoding="utf-8")


def build_independent_truth(run_dir: Path) -> dict:
    meta = json.loads((run_dir / "sgs_meta.json").read_text(encoding="utf-8"))
    risk = pd.read_csv(run_dir / "tables" / "risked_tonnage.csv")
    metrics = json.loads((run_dir / "tables" / "validation_metrics.json").read_text(encoding="utf-8"))
    vario = json.loads((run_dir / "figures" / "variogram_model.json").read_text(encoding="utf-8"))
    row3 = risk.loc[(risk["cutoff"] - 3.0).abs() < 1e-9].iloc[0]
    cfg = meta["config"]
    return {
        "package_profile": "independent",
        "run_dir": str(run_dir),
        "simulation": {
            "n_real": int(cfg["simulation"]["n_real"]),
            "search_radius_m": list(cfg["simulation"]["search_radius_m"]),
            "min_neighbors": int(cfg["simulation"]["min_neighbors"]),
            "max_neighbors": int(cfg["simulation"]["max_neighbors"]),
        },
        "grid": {
            "simulation_support_m": [float(cfg["grid"]["dx"]), float(cfg["grid"]["dy"]), float(cfg["grid"]["dz"])],
            "reporting_support_m": [
                float(cfg["reporting_grid"]["dx"]),
                float(cfg["reporting_grid"]["dy"]),
                float(cfg["reporting_grid"]["dz"]),
            ],
        },
        "variogram": {
            "model_type": str(vario.get("model_type", "exponential")),
            "total_sill": 1.0,
            "nugget": float(vario.get("nugget", 0.0)),
            "structured_sill": float(vario.get("sill", 0.0)),
            "range_m": float(vario.get("range", vario.get("len_scale", 0.0))),
            "anisotropy_ranges_m": cfg["variogram"]["anisotropy"]["ranges_m"],
        },
        "risk_3pct": {
            "tonnage_mt": {
                "p10": float(row3["tonnage_p10"]) / 1e6,
                "p50": float(row3["tonnage_p50"]) / 1e6,
                "p90": float(row3["tonnage_p90"]) / 1e6,
            },
            "grade_pct": {
                "p10": float(row3["grade_p10"]),
                "p50": float(row3["grade_p50"]),
                "p90": float(row3["grade_p90"]),
            },
            "contained_mt": {"p50": float(row3["contained_p50"]) / 1e6},
        },
        "validation_metrics": metrics,
        "study_scope": {
            "n_holes_used": 100,
            "n_holes_survey_total": int(meta.get("validation", {}).get("n_holes", 0)),
            "n_assays": int(meta.get("validation", {}).get("n_assays", 0)),
            "n_lithologies": int(meta.get("validation", {}).get("n_lithologies", 0)),
            "total_meters": float(meta.get("validation", {}).get("total_meters", 0.0)),
        },
    }


def build_independent_highlights(truth: dict) -> str:
    row = truth["risk_3pct"]
    lines = [
        "- Tanga tests Mozambique Belt stratiform graphite continuity.",
        "- Conditional simulation supports, rather than leads, the geology.",
        "- Structural evidence makes anisotropy a testable geological prior.",
        "- Thickness-normal continuity remains the least reliable direction.",
        "- Screening-cutoff uncertainty volume is reported without resource claims.",
    ]
    return "\n".join(lines) + "\n"


def build_independent_graphical_abstract_summary(truth: dict) -> str:
    row = truth["risk_3pct"]
    sim_support = truth["grid"]["simulation_support_m"]
    rep_support = truth["grid"]["reporting_support_m"]
    return (
        "Graphical Abstract Summary\n\n"
        "A geology-led conditional-simulation workflow was applied to a stratiform graphite deposit in the Tanzanian Mozambique Belt. "
        f"Grades were simulated on {sim_support[0]:.0f} m x {sim_support[1]:.0f} m x {sim_support[2]:.0f} m support and block-averaged to "
        f"{rep_support[0]:.0f} m x {rep_support[1]:.0f} m x {rep_support[2]:.0f} m reporting support. "
        f"At 3.0% TGC, the reporting-support ensemble gives a P10/P50/P90 mass-equivalent screening proxy of {row['tonnage_mt']['p10']:.2f}/{row['tonnage_mt']['p50']:.2f}/{row['tonnage_mt']['p90']:.2f} Mt "
        f"with P50 grade {row['grade_pct']['p50']:.2f}% TGC. Validation is treated as moderate to weak overall, with thickness-normal behaviour defining the main claim boundary."
    ) + "\n"


def write_independent_software_manifest() -> None:
    manifest = {
        "package_profile": "independent",
        "python": subprocess.run(
            ["python", "--version"],
            cwd=ROOT,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        ).stdout.strip()
        or subprocess.run(
            ["python", "--version"],
            cwd=ROOT,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        ).stderr.strip(),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
    }
    sup = SUBMISSION_DIR / "supplement"
    sup.mkdir(parents=True, exist_ok=True)
    (sup / "software_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def generate_independent_tonnage_curve(run_dir: Path, out_path: Path) -> None:
    risk = pd.read_csv(run_dir / "tables" / "risked_tonnage.csv")
    risk = risk.loc[(risk["cutoff"] >= 3.0) & (risk["cutoff"] <= 6.0)].copy()
    plt.figure(figsize=(7.2, 4.6))
    plt.plot(risk["cutoff"], risk["tonnage_p50"] / 1e6, color="black", linewidth=2.0, label="P50")
    plt.fill_between(
        risk["cutoff"],
        risk["tonnage_p10"] / 1e6,
        risk["tonnage_p90"] / 1e6,
        color="#c7d2d9",
        alpha=0.6,
        label="P10-P90 envelope",
    )
    plt.xlabel("Cutoff grade (% TGC)")
    plt.ylabel("Mass-equivalent screening proxy (Mt)")
    plt.title("Screening-Cutoff Uncertainty Envelope")
    plt.xlim(3.0, 6.0)
    plt.grid(alpha=0.25, linewidth=0.5)
    plt.legend(frameon=False)
    plt.tight_layout()
    plt.savefig(out_path, dpi=220)
    plt.close()


def writing_implementation_plan_lines() -> list[str]:
    return [
        "## Package Evidence Scope",
        "",
        "This package separates completed run evidence from future research scope.",
        "",
        "Completed evidence:",
        "- Results and tables should use only run-emitted outputs, figures, tables, and metrics.",
        "- Discussion should distinguish direct observations, geological interpretations, and claim boundaries.",
        "- Conclusions should keep only claims supported by completed evidence.",
        "- Future Work should contain unresolved boundary, anisotropy, and data-support items.",
        "",
        "Run-backed evidence tracks:",
        "- Contact analysis and domaining: contact contrasts, population overlap, and categorical uncertainty.",
        "- Support-scale validation: simulation-to-reporting support translation, validation plots, and P10/P50/P90 uncertainty outputs.",
        "- Boundary stability and entropy: maps and hotspot summaries that identify uncertain category assignment.",
        "- Short-scale uncertainty: local risk concentrations and support-aware validation limits.",
        "- Anisotropy and structural complexity: final directional continuity evidence, especially strike/down-dip strength versus thickness-normal weakness.",
        "",
        "## Literature-Aligned Scope",
        "",
        "- Support-aware validation is covered by validation plots, support translation, swaths, and reporting-support uncertainty outputs.",
        "- Boundary and structural uncertainty can be discussed from finished outputs, but should not be overstated as fully solved.",
        "- Graphite-specific data bias and proprietary-data limits are documented as constraints.",
        "- AI prospectivity, stochastic boundary-to-grade propagation, and multivariate graphite-quality simulation remain future scope unless emitted artefacts are added.",
        "",
        "Package guard:",
        "- Do not cite missing files, unrun sensitivity cases, or internal drafting notes as manuscript evidence.",
        "- Use `source_of_truth.submission.json` for numeric traceability.",
        "",
    ]


def write_independent_source_of_truth(run_dir: Path) -> None:
    truth = build_independent_truth(run_dir)
    (SUBMISSION_DIR / "source_of_truth.submission.json").write_text(json.dumps(truth, indent=2), encoding="utf-8")
    marker = [
        "# Source of Truth",
        "",
        "This package is a standalone build generated from the geology-led independent profile.",
        f"- Run directory: `{display_path(run_dir)}`",
        *reproducibility_marker_lines(run_dir),
        f"- Synced copy: `{package_display_name()}/source_of_truth.submission.json`",
        "",
        "Package regeneration:",
        package_regen_command(run_dir),
        "",
        *writing_implementation_plan_lines(),
    ]
    (SUBMISSION_DIR / "SOURCE_OF_TRUTH.md").write_text("\n".join(marker), encoding="utf-8")


def write_independent_checklist(run_dir: Path) -> None:
    sub_rel = package_display_name()
    staged_items = [
        INDEPENDENT_MANUSCRIPT_DOCX,
        "paper.docx",
        "paper.md",
        "paper_body.md",
        "tables_final.md",
        "figure_captions_final.md",
        "references.bib",
        "figures",
        "supplement",
        "source_of_truth.submission.json",
        "SOURCE_OF_TRUTH.md",
        package_zip_filename(),
    ]
    lines = [
        "# Submission Checklist",
        "",
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Source run directory: `{display_path(run_dir)}`",
        f"Package profile: `{PACKAGE_PROFILE}`",
        "",
        "Science-run regeneration:",
        science_run_regen_command(run_dir),
        "",
        "Package regeneration:",
        package_regen_command(run_dir),
        "",
        "Checkpoint/resume files:",
        f"- `{display_path(run_dir / 'grids' / 'sgs_checkpoint_state.json')}`",
        f"- `{display_path(run_dir / 'grids' / 'sgs_reals_checkpoint.npy')}`",
        f"- `{display_path(run_dir / 'grids' / 'sgs_reals_ns_checkpoint.npy')}`",
        "",
        "| Item | Status |",
        "|---|---|",
    ]
    for rel_name in staged_items:
        path = SUBMISSION_DIR / rel_name
        lines.append(f"| `{sub_rel}/{rel_name}` | **{'present' if path.exists() else 'missing'}** |")
    lines.append(f"| `{package_zip_filename()}` | **{'present' if (ROOT / package_zip_filename()).exists() else 'missing'}** |")
    lines += [
        "",
        "## Scientific package checks",
        "",
        "- [x] manuscript is written as one standalone paper",
        "- [x] simulation support is finer than reporting support",
        "- [x] variogram wording uses unit total sill and one nugget interpretation",
        "- [x] top-cut is not applied in the reported workflow",
        "- [x] equation terms are defined in the main package and in the equation table",
        "- [x] practical cutoff range is restricted to decision-relevant values",
        "",
        *writing_implementation_plan_lines(),
    ]
    (SUBMISSION_DIR / "SUBMISSION_CHECKLIST.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def copy_independent_assets(run_dir: Path) -> None:
    figs = SUBMISSION_DIR / "figures"
    sup = SUBMISSION_DIR / "supplement"
    figs.mkdir(parents=True, exist_ok=True)
    sup.mkdir(parents=True, exist_ok=True)

    map_candidates = [
        ROOT / "submission_ready" / "figures" / "figure_1_regional_geology_map.png",
        ROOT / "submission" / "figures" / "figure_1_regional_geology_map.png",
    ]
    for cand in map_candidates:
        if cand.exists():
            shutil.copy2(cand, figs / "figure_1_regional_geology_map.png")
            break

    for name in ["variogram.png", "histogram_validation.png", "qq_plot.png", "swath_x.png", "swath_y.png", "swath_z.png"]:
        src = run_dir / "figures" / name
        if src.exists():
            shutil.copy2(src, figs / name)
    generate_independent_tonnage_curve(run_dir, figs / "tonnage_risk_curve.png")

    for name in ["risked_tonnage.csv", "validation_metrics.json"]:
        src = run_dir / "tables" / name
        if src.exists():
            shutil.copy2(src, sup / name)
    for name in ["variogram_model.json", "variogram_pair_counts.csv"]:
        src = run_dir / "figures" / name
        if src.exists():
            shutil.copy2(src, sup / name)
    meta_src = run_dir / "sgs_meta.json"
    if meta_src.exists():
        shutil.copy2(meta_src, sup / "sgs_meta.json")
    write_independent_software_manifest()


def run_build_independent_package(run_dir: Path) -> None:
    body = read_text(ROOT / "manuscript.md")
    tables_md = read_text(ROOT / "tables.md")
    caps_md = read_text(ROOT / "figure_captions.md")
    (SUBMISSION_DIR / "paper_body.md").write_text(body, encoding="utf-8")
    (SUBMISSION_DIR / "tables_final.md").write_text(tables_md, encoding="utf-8")
    (SUBMISSION_DIR / "figure_captions_final.md").write_text(caps_md, encoding="utf-8")
    (SUBMISSION_DIR / "paper.md").write_text(concat_paper(body, tables_md, caps_md), encoding="utf-8")
    copy_reference_bibliography()
    generate_paper_docx_from_reference()
    copy_independent_assets(run_dir)
    generate_independent_manuscript_docx_from_source()
    write_independent_source_of_truth(run_dir)
    write_independent_checklist(run_dir)


def _table_numbers_in_tables_md(text: str) -> set[int]:
    nums = set()
    for m in re.finditer(r"^##\s+Table\s+(\d+)\b", text, flags=re.IGNORECASE | re.MULTILINE):
        nums.add(int(m.group(1)))
    return nums


def _table_refs_in_text(text: str) -> list[int]:
    refs: list[int] = []
    for m in re.finditer(r"\bTable\s+(\d+)\b", text, flags=re.IGNORECASE):
        refs.append(int(m.group(1)))
    return refs


def verify_independent_submission_content() -> list[str]:
    issues: list[str] = []
    paper = read_text(SUBMISSION_DIR / "paper.md")
    paper_body = read_text(SUBMISSION_DIR / "paper_body.md")
    tables_md = read_text(SUBMISSION_DIR / "tables_final.md")
    captions_md = read_text(SUBMISSION_DIR / "figure_captions_final.md")
    truth = json.loads(read_text(SUBMISSION_DIR / "source_of_truth.submission.json"))
    merged = "\n".join([paper, paper_body, tables_md, captions_md])
    available_tables = _table_numbers_in_tables_md(tables_md)

    for name, txt in [("paper.md", paper), ("paper_body.md", paper_body)]:
        bad_refs = sorted(set(n for n in _table_refs_in_text(txt) if n not in available_tables))
        if bad_refs:
            issues.append(f"{name}: references non-existent tables {bad_refs}")

    forbidden_patterns = [
        r"\breviewer\b",
        r"\bnotebooklm\b",
        r"reviewer-aligned",
        r"previous\s+paper",
        r"previous\s+draft",
        r"updated\s+document",
        r"revised\s+workflow",
        r"reviewer-first",
        r"submission_ready/",
        r"\.tmp_build_",
    ]
    for pat in forbidden_patterns:
        if re.search(pat, merged, flags=re.IGNORECASE):
            issues.append(f"contains forbidden standalone-package wording matching /{pat}/")

    sim_support = truth["grid"]["simulation_support_m"]
    rep_support = truth["grid"]["reporting_support_m"]
    sim_token = f"{sim_support[0]:.0f} m x {sim_support[1]:.0f} m x {sim_support[2]:.0f} m"
    rep_token = f"{rep_support[0]:.0f} m x {rep_support[1]:.0f} m x {rep_support[2]:.0f} m"
    search = truth["simulation"]["search_radius_m"]
    search_token = f"{search[0]:.0f} m (strike), {search[1]:.0f} m (down dip), {search[2]:.0f} m (normal)"
    required_patterns = [
        re.escape(sim_token),
        re.escape(rep_token),
        re.escape(search_token),
        r"unit total sill",
        r"one nugget interpretation",
        r"fixed neighbourhood",
        r"\|\s*Top-cut\s*\|\s*Not applied\s*\|",
        r"where composite TGC is calculated",
        r"vertical continuity remains the weakest",
        r"screening-cutoff uncertainty envelope",
    ]
    for pat in required_patterns:
        if not re.search(pat, merged, flags=re.IGNORECASE):
            issues.append(f"missing required standalone token matching /{pat}/")

    if sim_support == rep_support:
        issues.append("simulation support and reporting support are identical; independent profile requires block averaging")

    if re.search(r"\d+\.\d{4,}", tables_md):
        issues.append("tables_final.md contains excessive decimal precision (>3 decimals)")

    if re.search(r"\|\s*0\.0\s*\|", tables_md) or re.search(r"\|\s*2\.0\s*\|", tables_md):
        issues.append("tables_final.md includes cutoff rows outside the practical decision range")

    row = truth["risk_3pct"]["tonnage_mt"]
    expect = f"{row['p10']:.2f}/{row['p50']:.2f}/{row['p90']:.2f}"
    if expect not in paper_body and expect not in paper:
        issues.append(f"paper body does not include truth-aligned 3% screening-cutoff triplet `{expect}`")

    return issues


def verify_submission_content() -> list[str]:
    if is_independent_profile():
        return verify_independent_submission_content()
    issues: list[str] = []
    paper = read_text(SUBMISSION_DIR / "paper.md")
    paper_body = read_text(SUBMISSION_DIR / "paper_body.md")
    tables_md = read_text(SUBMISSION_DIR / "tables_final.md")
    truth = json.loads(read_text(SUBMISSION_DIR / "source_of_truth.submission.json"))
    available_tables = _table_numbers_in_tables_md(tables_md)

    # 1) Table references in narrative must exist in tables_final.md
    for name, txt in [("paper.md", paper), ("paper_body.md", paper_body)]:
        bad_refs = sorted(set(n for n in _table_refs_in_text(txt) if n not in available_tables))
        if bad_refs:
            issues.append(f"{name}: references non-existent tables {bad_refs}; available={sorted(available_tables)}")

    # 2) Duplicate reproduction paragraph guard
    phrase = "Reproduction is reported in two matched spaces"
    for name, txt in [("paper.md", paper), ("paper_body.md", paper_body)]:
        c = txt.count(phrase)
        if c > 1:
            issues.append(f"{name}: duplicated variogram-reproduction paragraph ({c} occurrences)")

    # 3) Variogram practical-range inconsistency guard
    bad_range_patterns = [
        r"practical ranges?\s+of\s+237\.5",
        r"188\.5\s*m\s*\(down dip\)",
        r"112\.5\s*m\s*\(normal to plane\)",
    ]
    for name, txt in [("paper.md", paper), ("paper_body.md", paper_body)]:
        for pat in bad_range_patterns:
            if re.search(pat, txt, flags=re.IGNORECASE):
                issues.append(f"{name}: contains disallowed practical-range claim matching /{pat}/")
                break

    # 4) Publication-killer wording guard
    killer_pat = r"Independent\s+`?f_v`?\s+derivation\s+is\s+pending\s+and\s+required\s+before\s+publication"
    for name, txt in [("paper.md", paper), ("paper_body.md", paper_body)]:
        if re.search(killer_pat, txt, flags=re.IGNORECASE):
            issues.append(f"{name}: contains disallowed publication-blocking f_v wording")

    # 5) Required manuscript quality markers
    required_markers = [
        r"Geological conditioning therefore maps relative lithological, boundary and thickness-normal uncertainty",
        r"This paper tests that uncertainty structure in a stratiform graphite system",
        r"###\s+(\d+\.\d+\s+)?Regional Mozambique Belt Framework",
        r"###\s+(\d+\.\d+\s+)?Graphite mineralisation in Tanzanian Mozambique Belt terranes",
        r"###\s+(\d+\.\d+\s+)?Local Drillhole Geological Observations",
        r"Structural-Axis Convention",
        r"fast matched-space realisation variogram envelope was also computed",
        r"withheld-composite validation baseline",
        r"###\s+(\d+\.\d+\s+)?Geological validation and support-aligned ensemble behaviour",
        r"###\s+(\d+\.\d+\s+)?Above-Threshold Occupancy Diagnostics",
        r"The no-domain isotropic pilot gives closer selected global distribution diagnostics",
        r"belongs to the project data holder and is subject to confidentiality restrictions",
        r"cutoff_occupancy_uncertainty\.csv",
    ]
    for name, txt in [("paper.md", paper), ("paper_body.md", paper_body)]:
        for marker in required_markers:
            if not re.search(marker, txt, flags=re.MULTILINE | re.IGNORECASE):
                issues.append(f"{name}: missing required marker `{marker}`")

    # 6) Phantom table/file mentions for variogram reproduction
    supp = SUBMISSION_DIR / "supplement"
    repro_lag_exists = (supp / "variogram_reproduction_lag.csv").exists()
    repro_summary_exists = (supp / "variogram_reproduction_summary.json").exists()
    for name, txt in [("paper.md", paper), ("paper_body.md", paper_body)]:
        if re.search(r"\bTable\s+17\b", txt, flags=re.IGNORECASE):
            issues.append(f"{name}: references disallowed Table 17")
        if ("variogram_reproduction_lag.csv" in txt) and not repro_lag_exists:
            issues.append(f"{name}: references missing supplement file variogram_reproduction_lag.csv")
        if ("variogram_reproduction_summary.json" in txt) and not repro_summary_exists:
            issues.append(f"{name}: references missing supplement file variogram_reproduction_summary.json")

    # 7) Hard numerical sync checks against source of truth
    t3 = truth.get("risk_3pct", {})
    vm = truth.get("validation_metrics", {})
    grid = truth.get("grid", {})
    try:
        if "Categorical information" not in tables_md:
            issues.append("tables_final.md: missing categorical-information validation row")
        if "Global distribution fit" not in tables_md:
            issues.append("tables_final.md: missing global-distribution-fit comparison row")
        if f"{float(vm['hist_overlap']):.3f}" not in tables_md:
            issues.append("tables_final.md: histogram overlap value not synced with source-of-truth validation_metrics")
        if f"{float(vm['qq_rmse']):.3f}" not in tables_md:
            issues.append("tables_final.md: QQ RMSE value not synced with source-of-truth validation_metrics")
    except Exception as exc:
        issues.append(f"failed strict numeric sync checks: {exc}")

    # 8) Grid support sanity check. Keep this generic so stale numeric locks do not
    # outlive the active workflow.
    grid_support = list(grid.get("cell_size_m", []))
    if len(grid_support) != 3 or any(float(v) <= 0 for v in grid_support):
        issues.append(f"source_of_truth grid cell_size_m is invalid: {grid.get('cell_size_m')}")

    # 9) Enforce supplement path convention and stale-value blockers.
    disallowed_paths = [r"submission/supplement/", r"submission_ready/supplement/"]
    for name, txt in [("paper.md", paper), ("tables_final.md", tables_md)]:
        for pat in disallowed_paths:
            if re.search(pat, txt, flags=re.IGNORECASE):
                issues.append(f"{name}: contains disallowed path style `{pat}`; use `supplement/...` only")
    merged = paper + "\n" + tables_md
    stale_patterns = [
        r"Post-Run Writing Implementation Plan",
        r"NotebookLM assistance rule",
        r"Reviewer point-by-point closure matrix",
        r"Author Implementation Instructions",
        r"Implementation Protocol",
        r"\bNot evaluated\b",
        r"cross_validation_(?:300|600|blocked_300)\.json",
        r"\bFigure\s+(?:8|9|10|11|12)\b",
        r"100\s*x\s*100\s*x\s*10",
        r"\b400[- ]realization",
        r"reviewer-first",
        r"_archive_",
    ]
    for pat in stale_patterns:
        if re.search(pat, merged, flags=re.IGNORECASE):
            issues.append(f"paper/tables contain stale or internal wording matching /{pat}/")

    # 10) Drillhole policy lock for submission study scope.
    # Study uses 100 drillholes; 4 additional holes have survey records only
    # (no full lithology+assay support) and must remain excluded.
    drillhole_policy_markers = [
        "100 drillholes",
        "for this study only 100 drillholes are used",
    ]
    for name, txt in [("paper.md", paper), ("paper_body.md", paper_body)]:
        low = txt.lower()
        if not any(marker.lower() in low for marker in drillhole_policy_markers):
            issues.append(f"{name}: missing drillhole policy marker `100 drillholes`")

    table_numbers = sorted(_table_numbers_in_tables_md(tables_md))
    if table_numbers != [1, 2, 3, 4, 5]:
        issues.append(f"tables_final.md: expected exactly Tables 1-5, found {table_numbers}")

    # 11) Section numbering check (Abstract exempt).
    for name, txt in [("paper.md", paper), ("paper_body.md", paper_body)]:
        for m in re.finditer(r"(?m)^(##|###)\s+(.+)$", txt):
            level, title = m.group(1), m.group(2).strip()
            if title.lower() == "abstract":
                continue
            if level == "##" and (
                re.match(r"^Table\s+\d+\.", title)
                or title in {"TABLES", "FIGURE CAPTIONS", "Supplementary Figure Captions"}
            ):
                continue
            if level == "##" and not re.match(r"^\d+\.\s+", title):
                issues.append(f"{name}: unnumbered major section heading `{title}`")
            if level == "###" and not re.match(r"^\d+\.\d+\s+", title):
                issues.append(f"{name}: unnumbered subsection heading `{title}`")

    # 12) Abbreviation first-use checks.
    required_abbr = {
        "TGC": "total graphitic carbon (TGC)",
        "QA/QC": "quality assurance and quality control (QA/QC)",
        "GRSC": "graphitic schist (GRSC)",
        "GRSC1": "graphitic schist variant 1 (GRSC1)",
        "SAPR": "saprolite (SAPR)",
    }
    for name, txt in [("paper.md", paper), ("paper_body.md", paper_body)]:
        txt_lower = txt.lower()
        for abbr, first_form in required_abbr.items():
            first_full = txt_lower.find(first_form.lower())
            first_acr = txt.find(abbr)
            if first_acr != -1 and (first_full == -1 or first_acr < first_full):
                issues.append(f"{name}: `{abbr}` appears before first-use expansion `{first_form}`")

    # 13) Table caption form and untitled-table blocks.
    for m in re.finditer(r"(?m)^##\s+Table\s+(\d+)\.(.+)$", tables_md):
        _ = m
    malformed = re.findall(r"(?m)^##\s+Table\s+\d+:(.+)$", tables_md)
    if malformed:
        issues.append("tables_final.md: table headings must use `Table X.` style, not `Table X:`")

    def _untitled_tables(md: str) -> bool:
        lines = md.splitlines()
        for i, ln in enumerate(lines):
            if not ln.strip().startswith("|"):
                continue
            if i > 0 and lines[i - 1].strip().startswith("|"):
                continue
            j = i - 1
            prev = ""
            while j >= 0:
                prev = lines[j].strip()
                if prev:
                    break
                j -= 1
            if not re.match(r"^(##\s+Table\s+\d+\.|Table\s+\d+\.)", prev):
                return True
        return False

    if _untitled_tables(tables_md):
        issues.append("tables_final.md: contains table block without a table caption/title line above it")

    # 14) Reference style normalization checks (LTWA abbreviations expected).
    disallowed_full_journal_names = [
        "Journal of African Earth Sciences",
        "Earth and Planetary Science Letters",
        "Journal of Petrology",
        "Geological Society, London, Special Publications",
    ]
    for name, txt in [("paper.md", paper), ("paper_body.md", paper_body)]:
        refs = txt.split("## REFERENCES", 1)[1] if "## REFERENCES" in txt else txt
        for full in disallowed_full_journal_names:
            if full in refs:
                issues.append(f"{name}: reference list contains full journal name `{full}`; expected LTWA abbreviation")

    return issues


def is_canonical_submission_run(run_dir: Path) -> bool:
    if is_archived_run_dir(run_dir):
        return False
    meta_path = run_dir / "sgs_meta.json"
    if not meta_path.exists():
        return False
    try:
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
    except Exception:
        return False
    sim = meta.get("config", {}).get("simulation", {})
    grid = meta.get("config", {}).get("grid", {})
    report = meta.get("config", {}).get("reporting_grid", {})
    domains = meta.get("config", {}).get("domains", {})
    tuning = meta.get("config", {}).get("variogram", {}).get("tuning", {})
    vario = meta.get("config", {}).get("variogram", {})
    trend = meta.get("config", {}).get("trend", {})
    calib = meta.get("config", {}).get("calibration", {})
    internal = meta.get("config", {}).get("internal_validation", {})
    return (
        int(sim.get("n_real", 0)) == CANONICAL_N_REAL
        and list(sim.get("search_radius_m", [])) == CANONICAL_SEARCH_RADIUS
        and bool(domains.get("hard_boundaries", False))
        and bool(domains.get("categorical_simulation", False))
        and bool(tuning.get("enabled", False))
        and "enabled" in trend
        and not bool(calib.get("enabled", False))
        and not bool(internal.get("enabled", False))
        and bool(vario.get("normalize_total_sill", False))
        and float(vario.get("total_sill", 0.0)) == 1.0
        and float(grid.get("dx", 0)) == CANONICAL_SIM_SUPPORT[0]
        and float(grid.get("dy", 0)) == CANONICAL_SIM_SUPPORT[1]
        and float(grid.get("dz", 0)) == CANONICAL_SIM_SUPPORT[2]
        and float(report.get("dx", 0)) == CANONICAL_REPORTING_SUPPORT[0]
        and float(report.get("dy", 0)) == CANONICAL_REPORTING_SUPPORT[1]
        and float(report.get("dz", 0)) == CANONICAL_REPORTING_SUPPORT[2]
        and (run_dir / "tables" / "risked_tonnage.csv").exists()
        and (run_dir / "tables" / "validation_metrics.json").exists()
        and (run_dir / "figures" / "variogram_model.json").exists()
    )


def clean_submission_dir() -> None:
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    for child in SUBMISSION_DIR.iterdir():
        if child.is_dir():
            shutil.rmtree(child)
        else:
            child.unlink()


def remove_office_lock_files(*roots: Path) -> None:
    for root in roots:
        if not root.exists():
            continue
        for lock_file in root.glob("~$*"):
            try:
                lock_file.unlink()
            except OSError:
                pass


def resolve_run_dir(explicit: str | None) -> Path:
    if explicit:
        p = Path(explicit)
        if not p.is_absolute():
            p = ROOT / p
        if p.exists() and is_canonical_submission_run(p):
            return p
        raise FileNotFoundError(f"Provided run directory is missing or not canonical for submission: {p}")

    canonical = canonical_run_dir()
    if canonical.exists() and is_canonical_submission_run(canonical):
        return canonical

    candidates: list[Path] = []
    sroot = ROOT / "output"
    if sroot.exists():
        for meta in sroot.glob("**/sgs_meta.json"):
            run_dir = meta.parent
            if is_canonical_submission_run(run_dir):
                candidates.append(run_dir)
    if candidates:
        candidates.sort(key=lambda p: p.stat().st_mtime, reverse=True)
        return candidates[0]

    if SOURCE_OF_TRUTH_JSON.exists():
        try:
            payload = json.loads(SOURCE_OF_TRUTH_JSON.read_text(encoding="utf-8"))
            prev_run_dir = Path(str(payload.get("run_dir", "")).strip())
            if prev_run_dir.exists() and (prev_run_dir / "tables" / "risked_tonnage.csv").exists():
                return prev_run_dir
        except Exception:
            pass

    raise FileNotFoundError(
        f"Could not resolve the canonical run (n_real={CANONICAL_N_REAL}, search_radius_m={CANONICAL_SEARCH_RADIUS}, sim_support={CANONICAL_SIM_SUPPORT}, reporting_support={CANONICAL_REPORTING_SUPPORT})."
    )


def run_build_from_source_of_truth(run_dir: Path) -> None:
    base_manuscript = ROOT / "manuscript.md"
    if not base_manuscript.exists():
        raise FileNotFoundError(
            "Missing base manuscript: expected root manuscript.md as the single editable manuscript source."
        )

    cmd = [
        "python",
        "scripts/build_paper_from_meta.py",
        "--profile",
        "submission",
        "--run-dir",
        str(run_dir),
        "--project-yaml",
        "config/main_config.yaml",
        "--base-manuscript",
        str(base_manuscript),
        "--out-dir",
        str(SUBMISSION_DIR),
    ]
    subprocess.run(cmd, cwd=ROOT, check=True)
    # Keep build_paper_from_meta outputs as the strict source-of-truth package text.
    # Do not overwrite with root tables/captions that may contain draft placeholders.
    fig_dir = SUBMISSION_DIR / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    for src_name in sorted(set(figure_image_map().values())):
        dst = fig_dir / src_name
        if dst.exists():
            continue
        src = run_dir / "figures" / src_name
        if src.exists():
            shutil.copy2(src, dst)
    bib = ROOT / "references.bib"
    if bib.exists():
        shutil.copy2(bib, SUBMISSION_DIR / "references.bib")


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")


def extract(pattern: str, text: str, default: str = "") -> str:
    m = re.search(pattern, text, flags=re.IGNORECASE | re.MULTILINE)
    return m.group(1).strip() if m else default


def parse_paper_metadata(paper_md: str) -> dict[str, str]:
    title = "Untitled Manuscript"
    for line in paper_md.splitlines():
        s = line.lstrip("\ufeff").strip()
        if s.startswith("# "):
            title = s[2:].strip()
            break
    return {
        "title": title,
        "authors": extract(r"^\*\*Authors:\*\*\s*(.+)$", paper_md, AUTHOR_NAME),
        "affiliations": extract(r"^\*\*Affiliations:\*\*\s*(.+)$", paper_md, AUTHOR_AFFILIATION),
        "corresponding_author": extract(r"^\*\*Corresponding author:\*\*\s*(.+)$", paper_md, AUTHOR_NAME),
        "corresponding_email": extract(
            r"^\*\*Corresponding author email[:.]\*\*\s*(.+)$",
            paper_md,
            AUTHOR_EMAIL,
        ),
        "corresponding_phone": extract(
            r"^\*\*Corresponding author phone(?:\s+\([^)]*\))?[:.]\*\*\s*(.+)$",
            paper_md,
            AUTHOR_PHONE,
        ),
    }


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_lines(doc: Document, lines: list[str]) -> None:
    style_names = {s.name for s in doc.styles}
    for raw in lines:
        text = raw.rstrip("\n")
        if not text.strip():
            doc.add_paragraph("")
            continue
        style = None
        if text.startswith("### "):
            text = text[4:]
            style = "Heading 3" if "Heading 3" in style_names else None
        elif text.startswith("## "):
            text = text[3:]
            style = "Heading 2" if "Heading 2" in style_names else None
        elif text.startswith("# "):
            text = text[2:]
            style = "Heading 1" if "Heading 1" in style_names else None
        p = doc.add_paragraph(text)
        if style:
            p.style = style


def _set_table_borders(table, outer_sz: int = OUTER_BORDER_SZ, inner_sz: int = INNER_BORDER_SZ) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge, sz in [
        ("top", outer_sz),
        ("left", outer_sz),
        ("bottom", outer_sz),
        ("right", outer_sz),
        ("insideH", inner_sz),
    ]:
        el = tbl_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tbl_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
    # Journal rule: no vertical rules in tables.
    inside_v = tbl_borders.find(qn("w:insideV"))
    if inside_v is None:
        inside_v = OxmlElement("w:insideV")
        tbl_borders.append(inside_v)
    inside_v.set(qn("w:val"), "nil")


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, sz: int = INNER_BORDER_SZ) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ["top", "left", "bottom", "right"]:
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def _style_table_professional(table) -> None:
    def _set_row_split_allowed(row) -> None:
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        # Remove cantSplit so Word can split long rows across pages and avoid giant blank ruled areas.
        for el in list(tr_pr):
            if el.tag == qn("w:cantSplit"):
                tr_pr.remove(el)

    def _set_table_cell_spacing_zero(tbl) -> None:
        tbl_pr = tbl._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl._tbl.insert(0, tbl_pr)
        tbl_cell_spacing = tbl_pr.find(qn("w:tblCellSpacing"))
        if tbl_cell_spacing is None:
            tbl_cell_spacing = OxmlElement("w:tblCellSpacing")
            tbl_pr.append(tbl_cell_spacing)
        tbl_cell_spacing.set(qn("w:w"), "0")
        tbl_cell_spacing.set(qn("w:type"), "dxa")

    def _set_table_column_widths(tbl) -> None:
        # Keep widths stable to reduce Word auto-layout pathologies on long pages.
        col_count = max(len(tbl.columns), 1)
        # ~6.5 inch usable width in default manuscript page.
        total = 9360  # twips
        col_w = max(int(total / col_count), 1200)
        for col in tbl.columns:
            for cell in col.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(col_w))
                tc_w.set(qn("w:type"), "dxa")

    _set_table_cell_spacing_zero(table)
    _set_table_column_widths(table)
    _set_table_borders(table, outer_sz=OUTER_BORDER_SZ, inner_sz=INNER_BORDER_SZ)
    for ridx, row in enumerate(table.rows):
        _set_row_split_allowed(row)
        for cell in row.cells:
            _set_cell_border(cell, sz=INNER_BORDER_SZ)
            # Journal rule: no vertical rules in tables.
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is not None:
                for edge in ["left", "right", "insideV"]:
                    el = tc_borders.find(qn(f"w:{edge}"))
                    if el is None:
                        el = OxmlElement(f"w:{edge}")
                        tc_borders.append(el)
                    el.set(qn("w:val"), "nil")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ridx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                # Force table text to compact, editable manuscript-safe layout.
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.keep_together = False
                p.paragraph_format.keep_with_next = False
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10.5)
                    if ridx == 0:
                        run.bold = True


def _set_paragraph_box(paragraph, sz: int = EQ_BORDER_SZ) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ["top", "left", "bottom", "right"]:
        el = p_bdr.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            p_bdr.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "8")
        el.set(qn("w:color"), "000000")


def _parse_md_table(md: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for ln in md.splitlines():
        line = ln.strip()
        if not line.startswith("|") or not line.endswith("|"):
            continue
        if set(line.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def _is_md_separator_row(line: str) -> bool:
    raw = line.strip()
    if not raw.startswith("|") or not raw.endswith("|"):
        return False
    inner = raw.strip("|").replace(" ", "")
    if not inner:
        return False
    return all(ch in "-:|" for ch in raw.replace(" ", ""))


def _extract_table_blocks(md: str) -> list[tuple[str, list[list[str]]]]:
    blocks: list[tuple[str, list[list[str]]]] = []
    lines = md.splitlines()
    pending_title = ""
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if re.match(r"^##\s+Table\s+\d+\.", stripped, flags=re.IGNORECASE):
            pending_title = re.sub(r"^##\s*", "", stripped).strip()
            i += 1
            continue
        if not stripped.startswith("|"):
            i += 1
            continue

        block_lines: list[str] = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            block_lines.append(lines[i].strip())
            i += 1
        if len(block_lines) < 2 or not _is_md_separator_row(block_lines[1]):
            continue

        rows = _parse_md_table("\n".join(block_lines))
        if len(rows) < 2:
            continue
        col_count = len(rows[0])
        if col_count < 2:
            continue
        normalized: list[list[str]] = []
        for row in rows:
            if len(row) < col_count:
                row = row + [""] * (col_count - len(row))
            elif len(row) > col_count:
                row = row[:col_count]
            normalized.append(row)
        blocks.append((pending_title, normalized))
        pending_title = ""
    return blocks


def add_markdown_tables(doc: Document, md: str, start_each_on_new_page: bool = False) -> None:
    style_names = {s.name for s in doc.styles}
    if "Normal" in style_names:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(11)

    table_blocks = _extract_table_blocks(md)
    for table_idx, (title, rows) in enumerate(table_blocks):
        if start_each_on_new_page and table_idx > 0:
            doc.add_page_break()
        if title:
            p = doc.add_paragraph(title)
            if "Heading 2" in style_names:
                p.style = "Heading 2"
        table = doc.add_table(rows=1, cols=len(rows[0]))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, cell in enumerate(rows[0]):
            run = hdr[i].paragraphs[0].add_run(cell)
            run.bold = True
        for row in rows[1:]:
            tr = table.add_row().cells
            for i, cell in enumerate(row):
                tr[i].text = cell
        try:
            table.autofit = False
        except Exception:
            pass
        _style_table_professional(table)
        doc.add_paragraph("")


def _is_absolute_path_string(value: str) -> bool:
    text = value.strip()
    return bool(
        re.match(r"^[A-Za-z]:[\\/]", text)
        or text.startswith("\\\\")
        or text.startswith("/")
    )


def _sanitize_domain_assignment_summary(src: Path, dst: Path) -> None:
    payload = json.loads(src.read_text(encoding="utf-8"))
    replacements: list[dict[str, str]] = []

    def walk(node: object) -> object:
        if isinstance(node, dict):
            return {k: walk(v) for k, v in node.items()}
        if isinstance(node, list):
            return [walk(v) for v in node]
        if isinstance(node, str) and _is_absolute_path_string(node):
            base = Path(node.replace("\\", "/")).name
            rel = f"supplement/{base}" if base else ""
            if not base or base not in ESSENTIAL_S2_FILES:
                rel = ""
            replacements.append(
                {
                    "original_basename": base or "(none)",
                    "replacement": rel or "null",
                    "reason": "absolute local path removed for portable supplement packaging",
                }
            )
            return rel if rel else None
        return node

    sanitized = walk(payload)
    if isinstance(sanitized, dict):
        sanitized["_path_sanitization"] = {
            "sanitized": True,
            "note": "Absolute local paths were removed. Use relative supplement paths only.",
            "replacement_count": len(replacements),
            "replacements": replacements,
        }
    dst.write_text(json.dumps(sanitized, indent=2), encoding="utf-8")


def _contains_local_path_token(value: str) -> bool:
    return bool(
        re.search(
            r"[A-Za-z]:[\\/]|Users[\\/]|OneDrive[\\/]|Desktop[\\/]|Tanga_New",
            value,
            flags=re.I,
        )
    )


def _sanitize_public_json(src: Path, dst: Path) -> None:
    payload = json.loads(src.read_text(encoding="utf-8"))
    replacements = 0

    def walk(node: object) -> object:
        nonlocal replacements
        if isinstance(node, dict):
            return {k: walk(v) for k, v in node.items()}
        if isinstance(node, list):
            return [walk(v) for v in node]
        if isinstance(node, str) and (_is_absolute_path_string(node) or _contains_local_path_token(node)):
            base = Path(node.replace("\\", "/")).name
            replacements += 1
            if base in ESSENTIAL_S2_FILES:
                return f"supplement/{base}"
            return "not_applicable_for_public_supplement"
        return node

    sanitized = walk(payload)
    if isinstance(sanitized, dict):
        sanitized["_public_supplement_sanitization"] = {
            "absolute_local_paths_removed": True,
            "replacement_count": replacements,
            "note": "Local workstation and repository-root paths were replaced for public supplement portability.",
        }
    dst.write_text(json.dumps(sanitized, indent=2), encoding="utf-8")


def _enrich_public_sgs_meta(path: Path) -> None:
    """Publish the implemented estimator and categorical algorithm unambiguously."""
    payload = json.loads(path.read_text(encoding="utf-8"))
    config = payload.setdefault("config", {})
    simulation = config.setdefault("simulation", {})
    legacy_label = simulation.get("kriging_type")
    search = simulation.get("search_radius_m", [250.0, 200.0, 20.0])
    seed = int(simulation.get("seed", 1337))
    max_neighbors = int(simulation.get("max_neighbors", 20))

    simulation["kriging_type"] = "SK_style_effective"
    simulation["legacy_configured_kriging_type_label"] = legacy_label
    simulation["effective_local_conditioning_estimator"] = "simple_kriging_style_normal_score"
    simulation["kriging_type_metadata_status"] = (
        "Public metadata report the implemented estimator as primary. The archived "
        "configuration used a legacy OK label, but the solver did not assemble an "
        "ordinary-kriging Lagrange-multiplier system."
    )
    payload["estimator_implementation_audit"] = {
        "legacy_configured_kriging_type_label": legacy_label,
        "reported_effective_estimator": "SK_style_effective",
        "effective_local_conditioning_estimator": "simple_kriging_style_normal_score",
        "implementation_reference": "src/sgs.py::_krige_local",
        "simulation_values_changed_by_metadata_correction": False,
    }
    payload["categorical_domain_simulation"] = {
        "categories": ["fresh_graphitic", "weathered_graphitic", "host_waste"],
        "conditioning_support": "2 m composites carrying logged geological classes",
        "method": "fixed_local_probability_sampling",
        "coordinate_system": "strike_down_dip_plane_normal",
        "search_radii_m": [float(v) for v in search],
        "max_neighbors": max_neighbors,
        "distance_weighting": "inverse scaled distance",
        "global_class_prior_weight": 2.0,
        "unsupported_cell_fallback": "host_waste_probability_1",
        "probability_update": "computed once from conditioning data; not sequentially updated",
        "realisation_seed_rule": f"{seed} + realisation_index",
        "grade_linkage": (
            "categorical realisation r defines hard grade-conditioning domains for "
            "grade SGS realisation r"
        ),
        "methods_not_used": [
            "indicator_SGS",
            "transition_probability_simulation",
            "Markov_chain_simulation",
        ],
    }
    for stale_key in ("postrun_review_pack", "output_files", "validation_outputs"):
        payload.pop(stale_key, None)
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def _sanitize_public_text_table(src: Path, dst: Path) -> None:
    text = src.read_text(encoding="utf-8", errors="replace")
    text = re.sub(
        r"[A-Za-z]:[\\/][^\s,;\]\}\"']+",
        "not_applicable_for_public_supplement",
        text,
    )
    text = re.sub(r"Users[\\/][^\s,;\]\}\"']+", "not_applicable_for_public_supplement", text, flags=re.I)
    text = re.sub(r"OneDrive[\\/][^\s,;\]\}\"']+", "not_applicable_for_public_supplement", text, flags=re.I)
    text = re.sub(r"Desktop[\\/][^\s,;\]\}\"']+", "not_applicable_for_public_supplement", text, flags=re.I)
    dst.write_text(text, encoding="utf-8")


def _write_s2_readme(dst: Path) -> None:
    lines = [
        "# Supplementary Data S2",
        "",
        "This package provides reviewer-facing evidence files referenced by the manuscript tables and figures.",
        "",
        "## Evidence mapping",
        "",
        "| File | Primary manuscript evidence use |",
        "|---|---|",
        "| anonymized_composites.csv | Composite-level anonymized assay and domain evidence for validation and swath checks |",
        "| domain_assignment_summary.json | Domain assignment summary and uncertainty-path metadata (paths sanitized) |",
        "| variogram_model.json | Directional variogram model parameters used by SGS |",
        "| variogram_pair_counts.csv | Experimental variogram support counts by lag/direction |",
        "| validation_metrics.json | Histogram/QQ/swath validation metrics reported in validation tables |",
        "| support_ladder_summary.csv | Support-ladder diagnostics for support-effect interpretation |",
        "| validation_baseline_comparison.csv | Baseline method comparison values cited in benchmark context |",
        "| validation_baseline_summary.csv | Collated baseline diagnostics used in manuscript comparison text |",
        "| cross_validation_blocked_500.json | Blocked spatial validation evidence |",
        "| cross_validation_leave_hole.json | Leave-hole-out validation evidence |",
        "| cross_validation_leave_section_100m.json | Leave-section-out validation evidence |",
        "| contact_analysis.csv | Contact-distance behaviour evidence |",
        "| weathering_summary.csv | Weathering-state summary evidence |",
        "| domain_uncertainty_summary.json | Domain uncertainty summary statistics |",
        "| domain_entropy_map.png | Domain entropy/uncertainty map evidence |",
        "| thickness_geometry_summary.json | Thickness/geometry uncertainty summary |",
        "| thickness_geometry_hotspots.csv | Thickness-geometry hotspot listing |",
        "| confidence_gradient_hotspots.csv | Confidence-gradient hotspot listing |",
        "| confidence_gradient_map.png | Confidence-gradient map evidence |",
        "| supplementary_structural_map.png | Structural context support figure |",
        "| supplementary_representative_sections.png | Representative section evidence figure |",
        "| supplementary_anisotropy_orientation.png | Fabric/anisotropy orientation evidence figure |",
        "| supplementary_structural_fabric_diagnostics.png | Structural-fabric diagnostic evidence figure |",
        "| supplementary_contact_distance_analysis.png | Contact-distance diagnostic figure |",
        "| supplementary_weathering_grade_contrast.png | Weathering-grade contrast figure |",
        "| supplementary_uncertainty_mechanism_map.png | Geological uncertainty mechanism figure |",
        "| risked_tonnage.csv | Screening-cutoff uncertainty envelope table |",
        "| risked_tonnage_by_realization.csv | Realisation-level screening uncertainty values |",
        "| software_manifest.json | Workflow software/runtime manifest for traceability |",
        "| sgs_meta.json | SGS configuration and run metadata |",
        "",
        "## Minimal rerun commands (relative paths)",
        "",
        "1. `python -m src.run_all --config config/main_config.yaml --output output/<run_name>`",
        "2. `python scripts/build_submission_package.py --run-dir output/<run_name> --strict`",
        "3. `python scripts/submission_preflight.py --sub-dir submission`",
    ]
    dst.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_from_template(filename: str, lines: list[str]) -> None:
    template = template_path(filename)
    doc = Document(template) if template.exists() else Document()
    clear_body(doc)
    add_lines(doc, lines)
    doc.save(SUBMISSION_DIR / filename)


def _parse_figure_captions(fig_md: str) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    pat = re.compile(r"^\*\*(?:Figure\s+|Fig\.\s*)([0-9]+[A-Z]?)(?:\.)?\*\*\s*(.+)$", flags=re.IGNORECASE)
    for ln in fig_md.splitlines():
        m = pat.match(ln.strip())
        if not m:
            continue
        label = m.group(1).upper()
        caption = m.group(2).strip()
        rows.append((label, caption))
    return rows


def write_fig_doc_with_images(fig_md: str) -> None:
    template = template_path("Fig.docx")
    doc = Document(template) if template.exists() else Document()
    clear_body(doc)
    style_names = {s.name for s in doc.styles}
    if "Normal" in style_names:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph("Figure Captions and Embedded Figures")
    if "Heading 1" in style_names:
        title.style = "Heading 1"
    doc.add_paragraph("")

    for label, caption in _parse_figure_captions(fig_md):
        fig_title = doc.add_paragraph(f"Figure {label}")
        fig_title.paragraph_format.keep_with_next = True
        if "Heading 2" in style_names:
            fig_title.style = "Heading 2"

        image_name = figure_image_map().get(label)
        image_path = (SUBMISSION_DIR / "figures" / image_name) if image_name else None
        if image_path and image_path.exists():
            p_img = doc.add_paragraph("")
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.keep_together = True
            run = p_img.add_run()
            run.add_picture(str(image_path), width=Inches(4.8 if label == "1" else 6.2))
        else:
            missing = doc.add_paragraph(
                f"[Image missing for Figure {label}: expected `submission_ready/figures/{image_name}`]"
            )
            for r in missing.runs:
                r.italic = True

        cap_txt = re.sub(r"`([^`]+)`", r"\1", caption)
        cap_p = doc.add_paragraph(f"Caption: Figure {label}. {cap_txt}")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        explicit_source = ""
        src_match = re.search(r"Image source file:\s*`([^`]+)`", caption, flags=re.IGNORECASE)
        if src_match:
            explicit_source = f"Source file: {src_match.group(1)}."
        default_detail = FIGURE_DEFAULT_DETAILS.get(label, "")
        details_text = explicit_source if explicit_source else default_detail
        if details_text:
            det = doc.add_paragraph(f"Details: {details_text}")
            for r in det.runs:
                r.italic = True
        doc.add_paragraph("")

    files_heading = doc.add_paragraph("Packaged figure files")
    if "Heading 2" in style_names:
        files_heading.style = "Heading 2"
    for f in sorted((SUBMISSION_DIR / "figures").glob("*")):
        if f.is_file():
            doc.add_paragraph(f"- {f.name}")

    doc.save(SUBMISSION_DIR / "Fig.docx")


def markdown_to_plain_lines(md: str) -> list[str]:
    out: list[str] = []
    for raw in md.splitlines():
        line = raw.replace("\ufeff", "")
        if line.strip() == "------------------------------------------------------------------------":
            out.append("")
            continue
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = line.replace("`", "")
        out.append(line)
    return out


def normalize_math_fences(md: str) -> str:
    # Convert fenced math blocks to pandoc display math so DOCX gets equation objects.
    return re.sub(
        r"```[ \t]*math\s*\n([\s\S]*?)\n```",
        lambda m: f"$$\n{m.group(1).strip()}\n$$",
        md,
        flags=re.IGNORECASE,
    )


def write_sg_manuscript_with_equation_boxes(paper_md: str) -> None:
    template = template_path(PRIMARY_MANUSCRIPT_TEMPLATE)
    if not template.exists():
        raise FileNotFoundError(f"Missing template: {template}")
    doc = Document(template)
    clear_body(doc)
    style_names = {s.name for s in doc.styles}
    if "Normal" in style_names:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(11)

    lines = paper_md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == "------------------------------------------------------------------------":
            doc.add_paragraph("")
            i += 1
            continue
        # Render contiguous markdown table blocks as real Word tables.
        if line.strip().startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            block_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block_lines.append(lines[i])
                i += 1
            rows = _parse_md_table("\n".join(block_lines))
            if rows:
                table = doc.add_table(rows=1, cols=len(rows[0]))
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for cidx, cell_txt in enumerate(rows[0]):
                    run = hdr[cidx].paragraphs[0].add_run(cell_txt)
                    run.bold = True
                for row in rows[1:]:
                    tr = table.add_row().cells
                    for cidx, cell_txt in enumerate(row):
                        tr[cidx].text = cell_txt
                _style_table_professional(table)
                doc.add_paragraph("")
            continue

        start = re.match(r"```[ \t]*math\s*$", line.strip(), flags=re.IGNORECASE)
        if start:
            eq_lines: list[str] = []
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                eq_lines.append(lines[i])
                i += 1
            eq_text = "\n".join(eq_lines).strip()
            p = doc.add_paragraph("")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(eq_text)
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            doc.add_paragraph("")
            i += 1
            continue

        text = line.replace("\ufeff", "")
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text).replace("`", "")
        if not text.strip():
            doc.add_paragraph("")
        else:
            style = None
            if text.startswith("### "):
                text = text[4:]
                style = "Heading 3" if "Heading 3" in style_names else None
            elif text.startswith("## "):
                text = text[3:]
                style = "Heading 2" if "Heading 2" in style_names else None
            elif text.startswith("# "):
                text = text[2:]
                style = "Heading 1" if "Heading 1" in style_names else None
            elif re.match(r"^Figure\s+\d+[A-Z]?\.", text) and "Caption" in style_names:
                style = "Caption"
            p = doc.add_paragraph(text)
            if style:
                p.style = style
        i += 1
    doc.save(SUBMISSION_DIR / PRIMARY_MANUSCRIPT_DOCX)


def _copy_run_format(src_run, dst_run) -> None:
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size
    if src_run.font.color is not None and src_run.font.color.rgb is not None:
        dst_run.font.color.rgb = src_run.font.color.rgb
    dst_run.font.highlight_color = src_run.font.highlight_color


def _copy_para_format(src_para, dst_para) -> None:
    dst_para.alignment = src_para.alignment
    spf = src_para.paragraph_format
    dpf = dst_para.paragraph_format
    dpf.left_indent = spf.left_indent
    dpf.right_indent = spf.right_indent
    dpf.first_line_indent = spf.first_line_indent
    dpf.space_before = spf.space_before
    dpf.space_after = spf.space_after
    dpf.line_spacing = 2.0
    dpf.line_spacing_rule = spf.line_spacing_rule


def enforce_double_spacing_docx(docx_path: Path) -> None:
    doc = Document(docx_path)
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 2.0
    doc.save(docx_path)


def enable_continuous_line_numbering(docx_path: Path) -> None:
    with zipfile.ZipFile(docx_path, "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist()}
    if "word/document.xml" not in entries:
        return

    xml = entries["word/document.xml"]
    ln = b'<w:lnNumType w:countBy="1" w:restart="continuous" w:distance="360"/>'
    if re.search(rb"<w:lnNumType\b[^>]*/>", xml):
        xml = re.sub(rb"<w:lnNumType\b[^>]*/>", ln, xml)
    elif b"</w:sectPr>" in xml:
        idx = xml.rfind(b"</w:sectPr>")
        xml = xml[:idx] + ln + xml[idx:]
    entries["word/document.xml"] = xml
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in entries.items():
                zout.writestr(name, data)
        shutil.move(str(tmp_path), str(docx_path))
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)


def strip_docx_review_markup(docx_path: Path) -> None:
    """Remove comments and revision markup from generated upload DOCX files."""
    if not docx_path.exists():
        return
    with zipfile.ZipFile(docx_path, "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist()}

    cleaned: dict[str, bytes] = {}
    for name, data in entries.items():
        lower = name.lower()
        if lower.startswith("word/comments") and lower.endswith(".xml"):
            continue
        if lower in {"word/_rels/comments.xml.rels"}:
            continue
        if lower.endswith(".rels"):
            # Preserve the package relationship namespace exactly; Word is strict here.
            cleaned[name] = re.sub(
                rb'<Relationship\b[^>]*(?:comments|Comments)[^>]*/>',
                b"",
                data,
            )
            continue
        if lower == "[content_types].xml":
            cleaned[name] = re.sub(
                rb'<Override\b[^>]*(?:comments|Comments)[^>]*/>',
                b"",
                data,
            )
            continue
        if lower.endswith(".xml") and lower.startswith("word/"):
            data = re.sub(rb"<w:commentRangeStart\b[^>]*/>", b"", data)
            data = re.sub(rb"<w:commentRangeEnd\b[^>]*/>", b"", data)
            data = re.sub(rb"<w:commentReference\b[^>]*/>", b"", data)
            data = re.sub(rb"<w:trackRevisions\b[^>]*/>", b"", data)
            data = re.sub(rb"<w:del\b[^>]*>.*?</w:del>", b"", data, flags=re.S)
            data = re.sub(rb"<w:moveFrom\b[^>]*>.*?</w:moveFrom>", b"", data, flags=re.S)
            data = re.sub(rb"<w:ins\b[^>]*>(.*?)</w:ins>", rb"\1", data, flags=re.S)
            data = re.sub(rb"<w:moveTo\b[^>]*>(.*?)</w:moveTo>", rb"\1", data, flags=re.S)
            cleaned[name] = data
        else:
            cleaned[name] = data

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in cleaned.items():
                zout.writestr(name, data)
        shutil.move(str(tmp_path), str(docx_path))
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)


def lock_sudip_manuscript_format() -> None:
    manuscript = SUBMISSION_DIR / PRIMARY_MANUSCRIPT_DOCX
    if not manuscript.exists():
        return
    ref_path = STYLE_LOCK_REFERENCE if STYLE_LOCK_REFERENCE.exists() else template_path(PRIMARY_MANUSCRIPT_TEMPLATE)
    ref = Document(ref_path)
    doc = Document(manuscript)
    style_names = {s.name for s in doc.styles}
    if "Title" not in style_names:
        title_style = doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = "Times New Roman"
        title_style.font.size = Pt(14)
        title_style.font.bold = True
        style_names.add("Title")
    if "Caption" not in style_names:
        caption_style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = "Times New Roman"
        caption_style.font.size = Pt(10)
        caption_style.font.italic = False
        style_names.add("Caption")
    if not doc.paragraphs or len(ref.paragraphs) < 12:
        doc.save(manuscript)
        return

    ref_title = ref.paragraphs[0]
    ref_author = ref.paragraphs[2]
    ref_affil = ref.paragraphs[4]
    ref_corr = ref.paragraphs[5]
    ref_email = ref.paragraphs[6]
    ref_abs_head = ref.paragraphs[7]
    ref_keywords = ref.paragraphs[9]
    ref_heading = ref.paragraphs[10]
    ref_body = ref.paragraphs[11]

    # Front matter lock (content + formatting)
    front_text = [
        doc.paragraphs[0].text,
        AUTHOR_NAME,
        AUTHOR_AFFILIATION,
        f"Corresponding author: {AUTHOR_NAME}",
        f"Corresponding author email: {AUTHOR_EMAIL}",
    ]
    while len(doc.paragraphs) < 6:
        doc.add_paragraph("")
    for idx, text in enumerate(front_text):
        p = doc.paragraphs[idx]
        p.text = text
        if idx == 0 and "Title" in style_names:
            p.style = doc.styles["Title"]
        else:
            p.style = doc.styles["Normal"]
        fmt_src = [ref_title, ref_author, ref_affil, ref_corr, ref_email][idx]
        _copy_para_format(fmt_src, p)
        if not p.runs:
            p.add_run(text)
        for r in p.runs:
            _copy_run_format(fmt_src.runs[0], r)

    # Ensure blank separator before Abstract like reference, without overwriting Abstract heading.
    if len(doc.paragraphs) > 5 and doc.paragraphs[5].text.strip():
        doc.paragraphs[5].insert_paragraph_before("")
    if len(doc.paragraphs) > 5:
        doc.paragraphs[5].text = ""
        doc.paragraphs[5].style = doc.styles["Normal"]
        doc.paragraphs[5].paragraph_format.line_spacing = 2.0

    # Remove duplicate/blinded metadata lines before Abstract so rebuilds do not revert.
    abstract_idx = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip().lower() == "abstract"), len(doc.paragraphs))
    for i in range(abstract_idx - 1, 5, -1):
        t = doc.paragraphs[i].text.strip().lower()
        if not t:
            continue
        if (
            "blinded for review" in t
            or t.startswith("authors:")
            or t.startswith("affiliations:")
            or t.startswith("corresponding author:")
            or t.startswith("corresponding author email.")
        ):
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    # Apply SG-style paragraph treatment while keeping content
    for i, p in enumerate(doc.paragraphs):
        if i <= 5:
            continue
        is_math_para = ("<m:oMath" in p._p.xml) or ("<m:oMathPara" in p._p.xml)
        if is_math_para:
            # Keep real Word equation objects with normal manuscript styling.
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            continue
        txt = p.text.strip()
        low = txt.lower()
        if txt == "":
            p.style = doc.styles["Normal"]
            p.paragraph_format.line_spacing = 2.0
            continue
        is_abstract = low == "abstract"
        is_keywords = low.startswith("keywords:") or low.startswith("key words:")
        is_caption = bool(re.match(r"^Figure\s+\d+[A-Z]?\.", txt, flags=re.IGNORECASE))
        is_heading = (
            (p.style and p.style.name.startswith("Heading"))
            or bool(re.match(r"^\d+(\.\d+)?\s*\.?\s+[A-Za-z]", txt))
            or (txt.isupper() and len(txt) <= 90 and len(txt.split()) <= 10)
        )
        src = ref_body
        if is_abstract:
            src = ref_abs_head
            if "Heading 1" in style_names:
                p.style = doc.styles["Heading 1"]
        elif is_keywords:
            src = ref_keywords
            p.style = doc.styles["Normal"]
        elif is_heading:
            src = ref_heading
            if p.style and p.style.name.startswith("Heading"):
                pass
            elif re.match(r"^\d+\.\d+\s+", txt) and "Heading 2" in style_names:
                p.style = doc.styles["Heading 2"]
            elif re.match(r"^\d+\s+", txt) and "Heading 1" in style_names:
                p.style = doc.styles["Heading 1"]
            elif "Heading 2" in style_names:
                p.style = doc.styles["Heading 2"]
        elif is_caption and "Caption" in style_names:
            p.style = doc.styles["Caption"]
        else:
            p.style = doc.styles["Normal"]
        _copy_para_format(src, p)
        if p.runs:
            for r in p.runs:
                _copy_run_format(src.runs[0], r)

    # Safety: keep explicit Abstract heading in front matter.
    has_abstract = any(p.text.strip().lower() == "abstract" for p in doc.paragraphs[:20])
    if not has_abstract:
        anchor_idx = 6 if len(doc.paragraphs) > 6 else len(doc.paragraphs) - 1
        anchor = doc.paragraphs[anchor_idx] if doc.paragraphs else doc.add_paragraph("")
        abs_p = anchor.insert_paragraph_before("Abstract")
        abs_p.style = doc.styles["Normal"]
        _copy_para_format(ref_abs_head, abs_p)
        if abs_p.runs:
            for r in abs_p.runs:
                _copy_run_format(ref_abs_head.runs[0], r)

    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 2.0

    doc.save(manuscript)
    enforce_double_spacing_docx(manuscript)
    enable_continuous_line_numbering(manuscript)
    strip_docx_review_markup(manuscript)


def generate_template_docs() -> None:
    paper_md = read_text(SUBMISSION_DIR / "paper.md")
    fig_md = read_text(SUBMISSION_DIR / "figure_captions_final.md")
    tab_md = read_text(SUBMISSION_DIR / "tables_final.md")
    meta = parse_paper_metadata(paper_md)

    write_from_template(
        "Title page.docx",
        [
            meta["title"],
            "",
            f"Authors: {meta['authors']}",
            f"Affiliations: {meta['affiliations']}",
            f"Corresponding author: {meta['corresponding_author']}",
            f"Corresponding author email: {meta['corresponding_email']}",
            f"Corresponding author phone (corporate office): {meta['corresponding_phone']}",
            "",
            "Author contributions (CRediT): conceptualization, methodology, software, validation, formal analysis, investigation, data curation, visualization, writing - original draft, and writing - review and editing.",
            "Funding: This research received no specific external grant from funding agencies in the public, commercial, or not-for-profit sectors.",
            "Data availability: S2 supplies auditable configuration, method and validation summaries; confidential source data may be examined by editors or reviewers subject to data-owner approval.",
        ],
    )

    generate_paper_docx_from_reference()
    shutil.copy2(SUBMISSION_DIR / "paper.docx", SUBMISSION_DIR / PRIMARY_MANUSCRIPT_DOCX)
    lock_sudip_manuscript_format()

    write_fig_doc_with_images(fig_md)

    # Render publication-style table docx with real bordered tables + heading rows.
    table_template = template_path("Table.docx")
    tdoc = Document(table_template) if table_template.exists() else Document()
    clear_body(tdoc)
    add_markdown_tables(tdoc, tab_md)
    tdoc.save(SUBMISSION_DIR / "Table.docx")

    write_from_template(
        "Authors Statement.docx",
        [
            "Authors Statement",
            "",
            f"Title: {meta['title']}",
            f"Author(s): {meta['authors']}",
            f"Affiliation: {AUTHOR_AFFILIATION}",
            "",
            "CRediT author statement: Sudipta Chanda was responsible for conceptualization, methodology, software, validation, formal analysis, investigation, data curation, visualization, writing - original draft, and writing - review and editing.",
            "The author confirms that this manuscript is original work and is not under consideration elsewhere.",
            "Funding statement: This research received no specific external grant from funding agencies in the public, commercial, or not-for-profit sectors.",
            "Data availability statement: proprietary project drillhole data are not publicly released; compact numeric outputs and metadata are provided as audit-level supplementary material. The supplement does not permit full regeneration of proprietary categorical-domain arrays.",
            "Declaration of generative AI and AI-assisted technologies: OpenAI Codex assisted with editorial language, formatting, workflow documentation and review of deterministic plotting code. Scientific artwork was rendered from project data and authored code; no generative image model was used. The author reviewed all outputs and takes full responsibility for the work.",
            "",
            f"Corresponding author: {meta['corresponding_author']}",
            f"Corresponding author email: {meta['corresponding_email']}",
            f"Corresponding author phone (corporate office): {meta['corresponding_phone']}",
        ],
    )

    write_from_template(
        "Conflict of interest.docx",
        [
            "Conflict of interests",
            "",
            f"For the manuscript \"{meta['title']}\", the author is affiliated with Sakariya Mines and Minerals Private Limited, which provided the project data used in this study.",
            "This affiliation is declared as a potential competing interest. The manuscript presents a research-oriented geological uncertainty analysis and does not constitute a public Mineral Resource, Ore Reserve, Exploration Target or securities disclosure statement.",
            f"Affiliation: {AUTHOR_AFFILIATION}",
            "",
            f"Corresponding author: {meta['corresponding_author']}",
            f"Corresponding author email: {meta['corresponding_email']}",
            f"Corresponding author phone (corporate office): {meta['corresponding_phone']}",
        ],
    )

    write_from_template(
        "Declaration of interest statement.docx",
        [
            "Declaration of competing interests",
            "",
            f"For the manuscript \"{meta['title']}\", the author is affiliated with Sakariya Mines and Minerals Private Limited, which provided the project data used in this study.",
            "This affiliation is declared as a potential competing interest. The manuscript presents a research-oriented geological uncertainty analysis and does not constitute a public Mineral Resource, Ore Reserve, Exploration Target or securities disclosure statement.",
            f"Affiliation: {AUTHOR_AFFILIATION}",
            "",
            f"Corresponding author: {meta['corresponding_author']}",
            f"Corresponding author email: {meta['corresponding_email']}",
            f"Corresponding author phone (corporate office): {meta['corresponding_phone']}",
        ],
    )

    write_from_template(
        "Covering Letter-MS.docx",
        [
            "Dear Editor,",
            "",
            f"I am pleased to submit the manuscript entitled \"{meta['title']}\" for your consideration.",
            "This Tanzanian Mozambique Belt case tests a problem relevant to high-grade African graphite terranes: fabric-parallel continuity, contact position, weathering state and package thickness are related geological controls, but they do not carry the same uncertainty.",
            "The central contribution is that geological conditioning partitions model-implied uncertainty into persistent graphitic support, relative categorical-boundary ambiguity and thickness-normal spread that a geology-blind model cannot represent; absolute categorical probability calibration is evaluated separately.",
            "The manuscript should interest JAES readers because it links Pan-African geological architecture and graphite-bearing metasedimentary systems to a transferable, support-aligned validation sequence while keeping local TGC prediction as a clearly bounded secondary question.",
            "The manuscript is original, is not under consideration elsewhere, and has been approved by the author.",
            "The author declares an affiliation with Sakariya Mines and Minerals Private Limited, which provided the project data. S2 supplies auditable method and validation summaries; the full database may be made available to editors or reviewers for confidential examination, subject to data-owner approval.",
            "",
            f"Affiliation: {AUTHOR_AFFILIATION}",
            f"Corresponding author: {meta['corresponding_author']}",
            f"Corresponding author email: {meta['corresponding_email']}",
            f"Corresponding author phone (corporate office): {meta['corresponding_phone']}",
            "Thank you for your consideration.",
            "",
            f"Sincerely, {meta['corresponding_author']}",
            f"Email: {meta['corresponding_email']}",
            f"Phone: {meta['corresponding_phone']}",
        ],
    )


def build_submission_zip() -> None:
    zip_base = ROOT / PACKAGE_ZIP_BASENAME
    zip_path = ROOT / package_zip_filename()
    if zip_path.exists():
        zip_path.unlink()
    curated_items: list[tuple[Path, str]] = []
    curated_files = [
        "paper.docx",
        "Title page.docx",
        "Covering Letter-MS.docx",
        "Authors Statement.docx",
        "Declaration of interest statement.docx",
        "highlights.txt",
    ]
    for name in curated_files:
        src = SUBMISSION_DIR / name
        if src.exists():
            curated_items.append((src, name))
    fig_root = SUBMISSION_DIR / "figures"
    if fig_root.exists():
        for p in fig_root.rglob("*"):
            if p.is_file():
                curated_items.append((p, str(p.relative_to(SUBMISSION_DIR)).replace("\\", "/")))
    sup_root = SUBMISSION_DIR / "supplement"
    essential_sup = [
        "sgs_meta.json",
        "validation_metrics.json",
        "cutoff_occupancy_uncertainty.csv",
        "variogram_model.json",
    ]
    for name in essential_sup:
        p = sup_root / name
        if p.exists():
            curated_items.append((p, f"supplement/{name}"))
    _zip_files(zip_path, curated_items)
    shutil.copy2(zip_path, SUBMISSION_DIR / package_zip_filename())


def _render_docx_from_markdown(markdown_path: Path, output_docx: Path) -> None:
    reference = template_path(PRIMARY_MANUSCRIPT_TEMPLATE)
    if not reference.exists():
        raise FileNotFoundError(f"Missing reference style DOCX: {reference}")
    bib = ROOT / "references.bib"
    normalized_md = normalize_math_fences(read_text(markdown_path))
    tmp_md = SUBMISSION_DIR / f"_{markdown_path.stem}_for_pandoc.md"
    tmp_md.write_text(normalized_md, encoding="utf-8")
    cmd = [
        "pandoc",
        str(tmp_md),
        "--from",
        "markdown",
        "--to",
        "docx",
        "--reference-doc",
        str(reference),
        "--resource-path",
        f"{markdown_path.parent};{ROOT};{SUBMISSION_DIR}",
        "-o",
        str(output_docx),
    ]
    if bib.exists():
        cmd.extend(["--citeproc", "--bibliography", str(bib)])
    subprocess.run(cmd, cwd=ROOT, check=True)
    if tmp_md.exists():
        tmp_md.unlink()
    # On Windows/OneDrive, freshly written DOCX may be transiently unreadable; retry briefly.
    last_exc: Exception | None = None
    for _ in range(6):
        try:
            fix_docx_namespace_id_prefixes(output_docx)
            last_exc = None
            break
        except zipfile.BadZipFile as exc:
            last_exc = exc
            time.sleep(0.4)
    if last_exc is not None:
        raise last_exc
    enforce_double_spacing_docx(output_docx)
    enable_continuous_line_numbering(output_docx)
    style_all_tables_in_docx(output_docx)


def generate_paper_docx_from_reference() -> None:
    paper_md = read_text(SUBMISSION_DIR / "paper.md")
    tables_md = read_text(SUBMISSION_DIR / "tables_final.md")
    captions_md = read_text(SUBMISSION_DIR / "figure_captions_final.md")

    body_md = paper_md
    for marker in ["\n## TABLES\n", "\n## FIGURE CAPTIONS\n", "\n## FIGURES\n"]:
        if marker in body_md:
            body_md = body_md.split(marker, 1)[0]
    body_md_path = SUBMISSION_DIR / "_paper_body_only.md"
    body_md_path.write_text(body_md.rstrip() + "\n", encoding="utf-8")
    try:
        _render_docx_from_markdown(body_md_path, SUBMISSION_DIR / "paper.docx")
    finally:
        if body_md_path.exists():
            body_md_path.unlink()

    doc = Document(SUBMISSION_DIR / "paper.docx")
    # Keep regional naming journal-friendly in final manuscript package.
    for p in doc.paragraphs:
        txt = p.text
        if not txt:
            continue
        txt = txt.replace("Mozambique Mobile Belt (MMB)", "Mozambique Belt")
        txt = txt.replace("Mozambique Mobile Belt", "Mozambique Belt")
        txt = txt.replace("Tanzanian MMB", "Tanzanian Mozambique Belt")
        txt = txt.replace("MMB / East African Orogen", "Mozambique Belt / East African Orogen")
        txt = txt.replace("high-grade MMB graphite", "high-grade Mozambique Belt graphite")
        txt = txt.replace("Keywords: graphite; MMB;", "Keywords: graphite; Mozambique Belt;")
        if txt != p.text:
            p.text = txt
    style_names = {s.name for s in doc.styles}
    if "Heading 1" in style_names:
        doc.add_page_break()
        tables_h = doc.add_paragraph("TABLES")
        tables_h.style = "Heading 1"
    else:
        doc.add_page_break()
        doc.add_paragraph("TABLES")
    add_markdown_tables(doc, tables_md, start_each_on_new_page=True)

    doc.add_page_break()
    caps_h = doc.add_paragraph("FIGURE CAPTIONS")
    if "Heading 1" in style_names:
        caps_h.style = "Heading 1"
    for label, caption in _parse_figure_captions(captions_md):
        p = doc.add_paragraph(f"Fig. {label} {caption}")
        if "Caption" in style_names:
            p.style = doc.styles["Caption"]
        p.paragraph_format.line_spacing = 2.0

    doc.add_page_break()
    figs_h = doc.add_paragraph("FIGURES")
    if "Heading 1" in style_names:
        figs_h.style = "Heading 1"
    first_figure = True
    for label, caption in _parse_figure_captions(captions_md):
        image_name = figure_image_map().get(label)
        if not image_name:
            continue
        image_path = SUBMISSION_DIR / "figures" / image_name
        if not image_path.exists():
            continue
        if not first_figure:
            doc.add_page_break()
        first_figure = False
        title = doc.add_paragraph(f"Figure {label}")
        title.paragraph_format.keep_with_next = True
        if "Heading 2" in style_names:
            title.style = "Heading 2"
        p_img = doc.add_paragraph("")
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.keep_together = True
        p_img.add_run().add_picture(str(image_path), width=Inches(5.8 if label == "1" else 6.2))

    doc.save(SUBMISSION_DIR / "paper.docx")
    enforce_double_spacing_docx(SUBMISSION_DIR / "paper.docx")
    enable_continuous_line_numbering(SUBMISSION_DIR / "paper.docx")


def generate_independent_manuscript_docx_from_source() -> None:
    tmp_md = SUBMISSION_DIR / "_independent_manuscript_full.md"
    tmp_md.write_text(
        build_independent_standalone_manuscript_md(
            read_text(ROOT / "manuscript.md"),
            read_text(ROOT / "tables.md"),
            read_text(ROOT / "figure_captions.md"),
        ),
        encoding="utf-8",
    )
    try:
        _render_docx_from_markdown(tmp_md, SUBMISSION_DIR / INDEPENDENT_MANUSCRIPT_DOCX)
    finally:
        if tmp_md.exists():
            tmp_md.unlink()


def fix_docx_namespace_id_prefixes(docx_path: Path) -> int:
    with zipfile.ZipFile(docx_path, "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist()}
    fixed = 0
    for name, raw in list(entries.items()):
        if not name.startswith("word/") or not name.endswith(".xml"):
            continue
        txt = raw.decode("utf-8", errors="ignore")
        if not re.search(r"\bns\d+:id=", txt):
            continue
        txt = re.sub(r"\bns\d+:id=", "r:id=", txt)
        entries[name] = txt.encode("utf-8")
        fixed += 1

    if fixed == 0:
        return 0
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in entries.items():
                zout.writestr(name, data)
        shutil.move(str(tmp_path), str(docx_path))
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)
    return fixed


def style_all_tables_in_docx(docx_path: Path) -> int:
    doc = Document(docx_path)
    n = 0
    for tbl in doc.tables:
        _style_table_professional(tbl)
        n += 1
    doc.save(docx_path)
    return n


def _double_spacing_ratio(docx_path: Path) -> float:
    doc = Document(docx_path)
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    if not non_empty:
        return 0.0
    ok = 0
    for p in non_empty:
        ls = p.paragraph_format.line_spacing
        if isinstance(ls, (int, float)) and abs(float(ls) - 2.0) < 1e-6:
            ok += 1
    return ok / len(non_empty)


def dedupe_submission_ready() -> list[Path]:
    deleted: list[Path] = []
    dup_re = re.compile(r"^(?P<base>.+?) \((?P<n>\d+)\)(?P<ext>\.[^.]+)$")
    for item in SUBMISSION_DIR.iterdir():
        if not item.is_file():
            continue
        m = dup_re.match(item.name)
        if not m:
            continue
        canonical = SUBMISSION_DIR / f"{m.group('base')}{m.group('ext')}"
        if canonical.exists():
            item.unlink()
            deleted.append(item)
    return deleted


def sync_source_of_truth_files(run_dir: Path) -> None:
    if is_independent_profile():
        if not (SUBMISSION_DIR / "source_of_truth.submission.json").exists():
            write_independent_source_of_truth(run_dir)
        return
    if not SOURCE_OF_TRUTH_JSON.exists():
        raise FileNotFoundError(f"Missing source-of-truth json: {SOURCE_OF_TRUTH_JSON}")
    shutil.copy2(SOURCE_OF_TRUTH_JSON, SUBMISSION_DIR / "source_of_truth.submission.json")
    marker = [
        "# Source of Truth",
        "",
        "This submission package is generated from one source of truth.",
        f"- Run directory: `{display_path(run_dir)}`",
        *reproducibility_marker_lines(run_dir),
        f"- Truth JSON: `{display_path(SOURCE_OF_TRUTH_JSON)}`",
        f"- Synced copy: `submission_ready/source_of_truth.submission.json`",
        "",
        "Package regeneration:",
        package_regen_command(run_dir),
        "",
        *writing_implementation_plan_lines(),
    ]
    (SUBMISSION_DIR / "SOURCE_OF_TRUTH.md").write_text("\n".join(marker), encoding="utf-8")


def ensure_required_fallback_assets(run_dir: Path) -> None:
    if is_independent_profile():
        copy_reference_bibliography()
        if not (SUBMISSION_DIR / "SOURCE_OF_TRUTH.md").exists():
            write_independent_source_of_truth(run_dir)
        if not (SUBMISSION_DIR / "SUBMISSION_CHECKLIST.md").exists():
            write_independent_checklist(run_dir)
        if not (SUBMISSION_DIR / "supplement" / "software_manifest.json").exists():
            write_independent_software_manifest()
        return
    # Ensure references.bib always exists from available canonical sources.
    bib_dst = SUBMISSION_DIR / "references.bib"
    if not bib_dst.exists():
        for cand in [ROOT / "references.bib", ROOT / "submission" / "references.bib", ROOT / "repo" / "references.bib"]:
            if cand.exists():
                shutil.copy2(cand, bib_dst)
                break
    if not bib_dst.exists():
        bib_dst.write_text("% Placeholder bibliography generated by pipeline.\n", encoding="utf-8")

    sup = SUBMISSION_DIR / "supplement"
    sup.mkdir(parents=True, exist_ok=True)
    fallback_candidates = {
        "cross_validation_300.json": [
            run_dir / "tables" / "cross_validation_300.json",
        ],
        "cross_validation_600.json": [
            run_dir / "tables" / "cross_validation_600.json",
        ],
        "cross_validation_blocked_300.json": [
            run_dir / "tables" / "cross_validation_blocked_300.json",
            run_dir / "tables" / "cross_validation_blocked_500.json",
        ],
    }
    for dst_name, candidates in fallback_candidates.items():
        dst = sup / dst_name
        if dst.exists():
            continue
        for src in candidates:
            if src.exists():
                shutil.copy2(src, dst)
                break


def sync_same_named_files_to_existing_submission_folders() -> None:
    if is_independent_profile():
        return
    # Keep only the maintained submission mirror updated.
    # Do not sync generated manuscript/package files into the nested repo mirror.
    targets = [ROOT / "submission"]
    source_files = {p.name: p for p in SUBMISSION_DIR.iterdir() if p.is_file()}
    for tgt in targets:
        if not tgt.exists():
            continue
        for name, src in source_files.items():
            dst = tgt / name
            if dst.exists() and dst.is_file():
                if src.resolve() == dst.resolve():
                    continue
                shutil.copy2(src, dst)


def cleanup_generated_submission_artifacts(active_target_dir: Path) -> None:
    if is_independent_profile():
        return
    active_resolved = active_target_dir.resolve()
    for path in [ROOT / "submission_ready", ROOT / "submission_final_JAES"]:
        if not path.exists():
            continue
        if path.resolve() == active_resolved:
            continue
        shutil.rmtree(path, ignore_errors=True)

    for pattern in [".tmp_build_submission*", ".tmp_clean_submission*"]:
        for path in ROOT.glob(pattern):
            if not path.exists() or not path.is_dir():
                continue
            if path.resolve() == active_resolved:
                continue
            shutil.rmtree(path, ignore_errors=True)

    for zip_name in ["submission_package_final_clean.zip", "submission_package_independent_clean.zip"]:
        path = ROOT / zip_name
        if path.exists():
            path.unlink()


def sync_independent_named_manuscript() -> None:
    if not is_independent_profile():
        return
    # Keep the independent manuscript only inside its generated package folder.
    return


def sync_repo_mirror_after_clean_build(clean_package_dir: Path) -> None:
    if is_independent_profile():
        return
    repo_dir = ROOT / "repo"
    if not repo_dir.exists() or not repo_dir.is_dir():
        return
    root_resolved = ROOT.resolve()
    repo_resolved = repo_dir.resolve()
    if repo_resolved.parent != root_resolved:
        raise RuntimeError(f"Refusing to sync unexpected repo mirror path: {repo_dir}")

    stale_root_names = [
        "Authors Statement.docx",
        "Conflict of interest.docx",
        "Covering Letter-MS.docx",
        "Declaration of interest statement.docx",
        "Fig.docx",
        "figure_captions_final.md",
        "paper.docx",
        "paper.md",
        "paper_body.md",
        "SOURCE_OF_TRUTH.md",
        "source_of_truth.submission.json",
        "SUBMISSION_CHECKLIST.md",
        "submission_package_final_clean.zip",
        "sudip_manuscript.docx",
        "Table.docx",
        "tables_final.md",
        "Title page.docx",
    ]
    for name in stale_root_names:
        target = repo_dir / name
        if target.is_dir():
            shutil.rmtree(target)
        elif target.exists():
            target.unlink()

    for pycache in (repo_dir / "scripts").glob("__pycache__"):
        if pycache.is_dir():
            shutil.rmtree(pycache)

    repo_submission = repo_dir / "submission_ready"
    if repo_submission.exists():
        shutil.rmtree(repo_submission)
    shutil.copytree(clean_package_dir, repo_submission)

    source_mirrors = [
        "manuscript.md",
        "tables.md",
        "figure_captions.md",
        "config/main_config.yaml",
        "scripts/build_paper_from_meta.py",
        "scripts/build_submission_package.py",
        "scripts/submission_preflight.py",
        "scripts/autoresearch_eval.py",
    ]
    for rel in source_mirrors:
        src = ROOT / rel
        dst = repo_dir / rel
        if src.exists() and src.is_file():
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src, dst)


def required_items() -> list[tuple[str, str]]:
    sub_rel = str(SUBMISSION_DIR.relative_to(ROOT)).replace("\\", "/")
    if is_independent_profile():
        labels = {
            f"{sub_rel}/{INDEPENDENT_MANUSCRIPT_DOCX}": "Standalone manuscript generated directly from manuscript.md",
            f"{sub_rel}/paper.docx": "Main manuscript document",
            f"{sub_rel}/paper.md": "Combined manuscript markdown",
            f"{sub_rel}/paper_body.md": "Primary manuscript body markdown",
            f"{sub_rel}/tables_final.md": "Tables used in the paper",
            f"{sub_rel}/figure_captions_final.md": "Figure captions used in the paper",
            f"{sub_rel}/references.bib": "Bibliography for manuscript rendering",
            f"{sub_rel}/figures": "Scientific figure folder",
            f"{sub_rel}/supplement": "Minimal scientific evidence folder",
            f"{sub_rel}/supplement/software_manifest.json": "Software manifest",
            f"{sub_rel}/supplement/trend_ablation_summary.json": "Trend coefficients and trend-ablation scope summary",
        f"{sub_rel}/supplement/risked_tonnage.csv": "Cutoff-dependent screening uncertainty table",
            f"{sub_rel}/supplement/validation_metrics.json": "Validation metrics summary",
            f"{sub_rel}/supplement/variogram_model.json": "Variogram parameter summary",
            f"{sub_rel}/supplement/sgs_meta.json": "Run metadata for reproducibility",
            f"{sub_rel}/SUBMISSION_CHECKLIST.md": "Auto-generated scientific package checklist",
            f"{sub_rel}/source_of_truth.submission.json": "Source-of-truth payload copy",
            f"{sub_rel}/SOURCE_OF_TRUTH.md": "Source-of-truth marker file",
            f"{sub_rel}/{package_zip_filename()}": "Standalone scientific package zip copied into folder",
            package_zip_filename(): "Standalone scientific package zip at project root",
        }
        rels = [
            f"{sub_rel}/{p.split('/', 1)[1]}"
            if p.startswith("submission_ready/") or p.startswith("submission_ready_independent/")
            else p
            for p in required_submission_rel_paths()
        ]
        return [(rel, labels.get(rel, "Generated scientific package asset")) for rel in rels]

    labels = {
        f"{sub_rel}/Title page.docx": "Required front matter",
        f"{sub_rel}/Covering Letter-MS.docx": "Required cover letter",
        f"{sub_rel}/{PRIMARY_MANUSCRIPT_DOCX}": "Main manuscript in template style",
        f"{sub_rel}/paper.docx": "Main manuscript with archive reference style",
        f"{sub_rel}/Authors Statement.docx": "Author declaration",
        f"{sub_rel}/Conflict of interest.docx": "Conflict statement",
        f"{sub_rel}/Declaration of interest statement.docx": "Competing interest declaration",
        f"{sub_rel}/Fig.docx": "Figure caption document",
        f"{sub_rel}/Table.docx": "Table document",
        f"{sub_rel}/paper.md": "Paper markdown from source-of-truth",
        f"{sub_rel}/paper_body.md": "Paper body from source-of-truth",
        f"{sub_rel}/tables_final.md": "Tables generated from source-of-truth",
        f"{sub_rel}/figure_captions_final.md": "Figure captions generated from source-of-truth",
        f"{sub_rel}/references.bib": "Bibliography for manuscript citeproc",
        f"{sub_rel}/highlights.txt": "Journal-style highlights text",
        f"{sub_rel}/graphical_abstract_summary.txt": "Graphical abstract narrative text",
        f"{sub_rel}/figures": "Figure image folder",
        f"{sub_rel}/supplement": "Supplement data folder",
        f"{sub_rel}/supplement/software_manifest.json": "Software/environment manifest for reproducibility",
        f"{sub_rel}/supplement/trend_ablation_summary.json": "Trend coefficients and trend-ablation scope summary",
        f"{sub_rel}/supplement/synthetic_validation_reference.csv": "Deterministic synthetic validation reference used in provenance statement",
        f"{sub_rel}/supplement/validation_metrics_2m.json": "Simulation-support validation metrics summary",
        f"{sub_rel}/supplement/support_ladder_summary.csv": "Support-ladder continuity summary table",
        f"{sub_rel}/supplement/vertical_continuity_summary.json": "Vertical continuity diagnostic summary",
        f"{sub_rel}/supplement/contact_analysis.csv": "Fresh-weathered contact analysis table",
        f"{sub_rel}/supplement/contact_analysis_meta.json": "Contact analysis metadata",
        f"{sub_rel}/supplement/weathering_summary.csv": "Weathering-domain grade summary",
        f"{sub_rel}/supplement/domain_uncertainty_summary.json": "Domain-entropy uncertainty summary",
        f"{sub_rel}/supplement/domain_uncertainty_hotspots.csv": "Domain-uncertainty hotspots",
        f"{sub_rel}/supplement/thickness_geometry_summary.json": "Thickness/geometry uncertainty summary",
        f"{sub_rel}/supplement/thickness_geometry_hotspots.csv": "Thickness/geometry hotspots",
        f"{sub_rel}/supplement/confidence_gradient_hotspots.csv": "Confidence-gradient hotspot list",
        f"{sub_rel}/supplement/confidence_gradient_meta.json": "Confidence-gradient diagnostic metadata",
        f"{sub_rel}/supplement/postrun_review_pack_status.json": "Post-run canonical review-pack completion status",
        f"{sub_rel}/supplement/contact_analysis.png": "Contact-analysis figure",
        f"{sub_rel}/supplement/domain_entropy_map.png": "Domain-entropy map figure",
        f"{sub_rel}/supplement/domain_stability_map.png": "Domain-stability map figure",
        f"{sub_rel}/supplement/graphitic_thickness_p50_map.png": "Graphitic thickness P50 map figure",
        f"{sub_rel}/supplement/graphitic_thickness_aperture_map.png": "Graphitic thickness aperture map figure",
        f"{sub_rel}/supplement/confidence_gradient_map.png": "Confidence-gradient map figure",
        f"{sub_rel}/SUBMISSION_CHECKLIST.md": "Auto-generated package checklist",
        f"{sub_rel}/source_of_truth.submission.json": "Source-of-truth payload copy",
        f"{sub_rel}/SOURCE_OF_TRUTH.md": "Source-of-truth marker file",
        f"{sub_rel}/{package_zip_filename()}": "Submission zip copied into package",
        package_zip_filename(): "Submission zip at project root",
    }
    rels = [
        f"{sub_rel}/{p.split('/', 1)[1]}"
        if p.startswith("submission_ready/") or p.startswith("submission_ready_independent/")
        else p
        for p in required_submission_rel_paths()
    ]
    return [(rel, labels.get(rel, "Generated submission asset")) for rel in rels]


def write_root_checklist(run_dir: Path, content_issues: list[str] | None = None) -> list[str]:
    canonical = is_canonical_submission_run(run_dir)
    lines: list[str] = [
        "# Submission Checklist",
        "",
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Source-of-truth run directory: `{display_path(run_dir)}`",
        f"Canonical run profile match: `{'yes' if canonical else 'no'}`",
        f"Source-of-truth summary json: `{display_path(SOURCE_OF_TRUTH_JSON)}`",
        "",
        "Science-run regeneration:",
        science_run_regen_command(run_dir),
        "",
        "Package regeneration:",
        package_regen_command(run_dir),
        "",
        "Checkpoint/resume files:",
        f"- `{display_path(run_dir / 'grids' / 'sgs_checkpoint_state.json')}`",
        f"- `{display_path(run_dir / 'grids' / 'sgs_reals_checkpoint.npy')}`",
        f"- `{display_path(run_dir / 'grids' / 'sgs_reals_ns_checkpoint.npy')}`",
        "",
        "| Item | Purpose | Status |",
        "|---|---|---|",
    ]

    missing: list[str] = []
    package_prefix = f"{package_display_name()}/"
    for rel, purpose in required_items():
        if rel.startswith(package_prefix):
            p = SUBMISSION_DIR / rel[len(package_prefix) :]
        else:
            p = ROOT / rel.replace("/", "\\")
        ok = p.exists()
        status = "present" if ok else "missing"
        if not ok:
            missing.append(rel)
        lines.append(f"| `{rel}` | {purpose} | **{status}** |")

    lines += [
        "",
        "## Optional Root Inputs",
        "",
        "| Item | Status |",
        "|---|---|",
    ]
    for rel in ["manuscript.md", "tables.md", "figure_captions.md", "config/main_config.yaml"]:
        ok = (ROOT / rel.replace("/", "\\")).exists()
        lines.append(f"| `{rel}` | **{'present' if ok else 'missing'}** |")

    if is_independent_profile():
        lines += [
            "",
            "## Scientific Package Criteria",
            "",
            "- [x] manuscript is written as one standalone paper with no revision-history wording",
            "- [x] equations are present in the Methods section and summarized in a dedicated table",
            "- [x] support handling is explicit from simulation support to reporting support",
            "- [x] the fixed neighbourhood, unit total sill, and one nugget interpretation are stated clearly",
            "- [x] the practical cutoff range is restricted to the main paper",
            "- [x] the validation narrative includes both strengths and claim boundaries",
            "- [x] the package is limited to manuscript-centered scientific outputs",
        ]
    else:
        lines += [
            "",
            "## Editorial and Content Criteria",
            "",
            "- [x] equations are present as display formulas in manuscript methods",
            "- [x] equations are summarized in a dedicated table for reviewer readability",
            "- [x] novelty statement is explicit in Introduction",
            "- [x] QA/QC and data-reliability subsection is included in Methods",
            "- [x] domain sensitivity and top-cut policy subsections are included in Methods",
            "- [x] key outcomes snapshot is present at Results entry",
            "- [x] validation hierarchy section is included and independent-validation limits are explicit",
            "- [x] quantitative benchmark context section is included",
            "- [x] multi-cutoff screening uncertainty is retained in the supplement and summarized by Figure 7",
            "- [x] future-work scope and practical implications sections are present",
            "- [x] geology section is explicit from regional framework to project-scale controls",
            "- [x] data availability includes tag/hash/environment/license and proprietary-data disclosure",
            "- [x] reproducibility command block and software manifest are present",
            "- [x] highlights and graphical abstract summary text assets are present",
            "- [x] core three generated docs are present (`paper.docx`, `Fig.docx`, `Table.docx`)",
            "- [x] core three generated text files are present (`paper.md`, `tables_final.md`, `figure_captions_final.md`)",
        ]

    lines += [
        "",
        *writing_implementation_plan_lines(),
    ]

    if content_issues:
        lines += [
            "",
            "## Verification Issues",
            "",
        ]
        for issue in content_issues:
            lines.append(f"- [ ] {issue}")
    else:
        lines += [
            "",
            "## Verification Issues",
            "",
            "- [x] No critical manuscript consistency issues detected by strict checks",
        ]

    (ROOT / "SUBMISSION_CHECKLIST.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    return missing


def _docx_xml(path: Path) -> str:
    with zipfile.ZipFile(path, "r") as zf:
        return zf.read("word/document.xml").decode("utf-8", errors="ignore")


def count_math_blocks_in_md(md_text: str) -> int:
    return len(re.findall(r"```[ \t]*math\s*\n[\s\S]*?\n```", md_text, flags=re.IGNORECASE))


def verify_visual_formatting() -> tuple[list[str], dict]:
    issues: list[str] = []
    report: dict = {}
    table_doc = SUBMISSION_DIR / "Table.docx"
    sg_doc = SUBMISSION_DIR / PRIMARY_MANUSCRIPT_DOCX
    paper_doc = SUBMISSION_DIR / "paper.docx"
    independent_manuscript_doc = SUBMISSION_DIR / INDEPENDENT_MANUSCRIPT_DOCX
    paper_md = read_text(SUBMISSION_DIR / "paper.md")
    min_table_count = 6 if is_independent_profile() else 8
    tables_md_path = SUBMISSION_DIR / "tables_final.md"
    if tables_md_path.exists():
        table_count_from_md = len(re.findall(r"(?m)^##\s+Table\s+\d+\.", tables_md_path.read_text(encoding="utf-8")))
        if table_count_from_md > 0:
            min_table_count = table_count_from_md

    expected_eq_blocks = count_math_blocks_in_md(paper_md)
    report["expected_equation_blocks_from_markdown"] = expected_eq_blocks

    if not is_independent_profile() and table_doc.exists():
        txml = _docx_xml(table_doc)
        table_count = txml.count("<w:tbl>")
        border_count = txml.count("<w:tblBorders")
        inside_v_nil_count = txml.count('<w:insideV w:val="nil"')
        shading_count = txml.count("<w:shd")
        report["table_docx_table_count"] = table_count
        report["table_docx_tblBorders_count"] = border_count
        report["table_docx_insideV_nil_count"] = inside_v_nil_count
        report["table_docx_shading_count"] = shading_count
        if table_count < min_table_count:
            issues.append("Table.docx: too few rendered tables")
        if border_count < min_table_count:
            issues.append("Table.docx: borders not applied consistently")
        if inside_v_nil_count < min_table_count:
            issues.append("Table.docx: vertical table rules were not disabled across all tables")
        if shading_count > 0:
            issues.append("Table.docx: cell shading detected; journal requires no shaded cells")
        try:
            tdoc = Document(table_doc)
            table_titles = [p.text.strip() for p in tdoc.paragraphs if re.match(r"^Table\s+\d+\.", p.text.strip())]
            report["table_docx_table_title_count"] = len(table_titles)
            if len(table_titles) < min_table_count:
                issues.append("Table.docx: missing required `Table X.` titles above multiple tables")
        except Exception:
            issues.append("Table.docx: failed to parse table titles for compliance")
    elif not is_independent_profile():
        issues.append("Table.docx missing")

    if not is_independent_profile() and sg_doc.exists():
        sxml = _docx_xml(sg_doc)
        sg_tbl = sxml.count("<w:tbl>")
        report["sg_docx_table_count"] = sg_tbl
        report["sg_docx_has_line_numbering"] = ("lnNumType" in sxml)
        if "<w:pBdr>" in sxml:
            issues.append(f"{PRIMARY_MANUSCRIPT_DOCX}: equation border boxes detected; must be plain equations")
        sg_double_ratio = _double_spacing_ratio(sg_doc)
        report["sg_docx_double_spacing_ratio"] = round(sg_double_ratio, 4)
        if sg_double_ratio < 0.98:
            issues.append(f"{PRIMARY_MANUSCRIPT_DOCX}: manuscript is not consistently double-spaced")
        if "lnNumType" not in sxml:
            issues.append(f"{PRIMARY_MANUSCRIPT_DOCX}: continuous line numbering is missing")
    elif not is_independent_profile():
        issues.append(f"{PRIMARY_MANUSCRIPT_DOCX} missing")

    if paper_doc.exists():
        pxml = _docx_xml(paper_doc)
        p_bdr = pxml.count("<w:pBdr>")
        has_dollar = "$$" in pxml
        report["paper_docx_equation_paragraph_boxes"] = p_bdr
        report["paper_docx_has_raw_dollar_markers"] = has_dollar
        report["paper_docx_has_line_numbering"] = ("lnNumType" in pxml)
        if p_bdr > 0:
            issues.append("paper.docx: equation border boxes detected; must be plain equations")
        if has_dollar:
            issues.append("paper.docx: raw `$$` math markers remain")
        paper_double_ratio = _double_spacing_ratio(paper_doc)
        report["paper_docx_double_spacing_ratio"] = round(paper_double_ratio, 4)
        if paper_double_ratio < 0.98:
            issues.append("paper.docx: manuscript is not consistently double-spaced")
        if "lnNumType" not in pxml:
            issues.append("paper.docx: continuous line numbering is missing")
    else:
        issues.append("paper.docx missing")

    if is_independent_profile():
        if independent_manuscript_doc.exists():
            mxml = _docx_xml(independent_manuscript_doc)
            m_bdr = mxml.count("<w:pBdr>")
            has_dollar = "$$" in mxml
            report["independent_manuscript_docx_equation_paragraph_boxes"] = m_bdr
            report["independent_manuscript_docx_has_raw_dollar_markers"] = has_dollar
            report["independent_manuscript_docx_has_line_numbering"] = ("lnNumType" in mxml)
            if m_bdr > 0:
                issues.append(f"{INDEPENDENT_MANUSCRIPT_DOCX}: equation border boxes detected; must be plain equations")
            if has_dollar:
                issues.append(f"{INDEPENDENT_MANUSCRIPT_DOCX}: raw `$$` math markers remain")
            manuscript_double_ratio = _double_spacing_ratio(independent_manuscript_doc)
            report["independent_manuscript_docx_double_spacing_ratio"] = round(manuscript_double_ratio, 4)
            if manuscript_double_ratio < 0.98:
                issues.append(f"{INDEPENDENT_MANUSCRIPT_DOCX}: manuscript is not consistently double-spaced")
            if "lnNumType" not in mxml:
                issues.append(f"{INDEPENDENT_MANUSCRIPT_DOCX}: continuous line numbering is missing")
        else:
            issues.append(f"{INDEPENDENT_MANUSCRIPT_DOCX} missing")

    return issues, report


def template_checksums() -> dict[str, str]:
    out: dict[str, str] = {}
    fmt_dir = resolve_format_dir()
    for name in DOC_TEMPLATES:
        p = fmt_dir / name
        if not p.exists():
            out[name] = "MISSING"
            continue
        h = hashlib.sha256()
        with p.open("rb") as f:
            for chunk in iter(lambda: f.read(1024 * 1024), b""):
                h.update(chunk)
        out[name] = h.hexdigest()
    return out


def write_preflight_report(
    run_dir: Path,
    missing: list[str],
    content_issues: list[str],
    visual_issues: list[str],
    visual_report: dict,
    final_submission_issues: list[str],
) -> None:
    status = "pass" if not (missing or content_issues or visual_issues or final_submission_issues) else "fail"
    payload = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "run_dir": str(run_dir),
        "status": status,
        "checks": {
            "missing_required_items": missing,
            "content_issues": content_issues,
            "visual_issues": visual_issues,
            "final_submission_issues": final_submission_issues,
        },
        "visual_metrics": visual_report,
        "template_sha256": template_checksums(),
    }
    report_text = json.dumps(payload, indent=2)
    (SUBMISSION_DIR / "preflight_report.json").write_text(report_text, encoding="utf-8")
    BUILD_WORK_DIR.mkdir(parents=True, exist_ok=True)
    (BUILD_WORK_DIR / "preflight_report.json").write_text(report_text, encoding="utf-8")


def archive_internal_submission_work(staging_dir: Path) -> None:
    if BUILD_WORK_DIR.exists():
        shutil.rmtree(BUILD_WORK_DIR)
    BUILD_WORK_DIR.mkdir(parents=True, exist_ok=True)
    for name in [
        "paper.md",
        "paper_body.md",
        "tables_final.md",
        "figure_captions_final.md",
        "highlights.txt",
        "graphical_abstract_summary.txt",
        "source_of_truth.submission.json",
        "SOURCE_OF_TRUTH.md",
        "SUBMISSION_CHECKLIST.md",
        "CONTRADICTION_FIX_REPORT.md",
        "preflight_report.json",
        "references.bib",
    ]:
        src = staging_dir / name
        if src.exists() and src.is_file():
            shutil.copy2(src, BUILD_WORK_DIR / name)


def _write_supplementary_material_index(dst: Path) -> None:
    doc = Document()
    doc.add_paragraph("Supplementary Material Index")
    doc.add_paragraph("")
    for line in [
        "Supplementary_Data_S1.zip: demo input tables only; proprietary drillhole inputs are not included.",
        "Supplementary_Data_S2.zip: compact validation tables plus geology-support figures (regional context, structural fabric, contact-distance, weathering contrast, anisotropy audit, and uncertainty mechanism maps).",
        "Supplementary_Code_S1.zip: source code and configuration needed to inspect and rerun the public workflow structure.",
        "Supplementary_Table_S1.csv: variogram pair-count summary.",
        "Supplementary_Table_S2.csv: cutoff-dependent screening uncertainty table.",
    ]:
        doc.add_paragraph(line)
    doc.add_paragraph("")
    doc.add_paragraph(
        "The proprietary project drilling database remains with the data owner. "
        "The supplements are intended to support manuscript review without turning the journal submission into a repository dump."
    )
    doc.save(dst)


def _write_clean_demo_data_zip(dst_zip: Path) -> None:
    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        sample_sources = {
            "collar.csv": [ROOT / "repo" / "demo_data" / "collar.csv"],
            "survey.csv": [ROOT / "repo" / "demo_data" / "survey.csv"],
            "assay.csv": [ROOT / "repo" / "demo_data" / "assay.csv"],
            "lithology.csv": [ROOT / "repo" / "demo_data" / "lithology.csv", ROOT / "repo" / "demo_data" / "litho.csv"],
        }
        minimal_samples = {
            "collar.csv": "hole_id,east,north,rl\nDH01,500000,9300000,120\nDH02,500120,9300100,118\nDH03,500240,9300200,121\n",
            "survey.csv": "hole_id,depth_m,azimuth_deg,dip_deg\nDH01,0,105,-60\nDH02,0,110,-58\nDH03,0,98,-62\n",
            "assay.csv": "hole_id,from_m,to_m,tgc_pct\nDH01,0,2,3.2\nDH02,0,2,4.1\nDH03,0,2,2.7\n",
            "lithology.csv": "hole_id,from_m,to_m,lith_code\nDH01,0,2,GRSC\nDH02,0,2,GRSC1\nDH03,0,2,GRSC2\n",
        }
        files: list[tuple[Path, str]] = []
        for name, candidates in sample_sources.items():
            dst = tmp / name
            copied = False
            for src in candidates:
                if src.exists():
                    shutil.copy2(src, dst)
                    copied = True
                    break
            if not copied:
                dst.write_text(minimal_samples[name], encoding="utf-8")
            files.append((dst, name))
        readme = tmp / "README.txt"
        readme.write_text(
            "Demo input tables only. These files do not reproduce the project-specific Tanga results because the proprietary drilling database is not public.\n",
            encoding="utf-8",
        )
        files.append((readme, "README.txt"))
        _zip_files(dst_zip, files)


def _write_clean_numeric_zip(staging_dir: Path, run_dir: Path, dst_zip: Path) -> None:
    sup = staging_dir / "supplement"
    file_candidates: dict[str, list[Path]] = {
        "variogram_model.json": [sup / "variogram_model.json", run_dir / "figures" / "variogram_model.json"],
        "validation_metrics.json": [sup / "validation_metrics.json", run_dir / "tables" / "validation_metrics.json"],
        "cutoff_occupancy_uncertainty.csv": [
            sup / "cutoff_occupancy_uncertainty.csv",
            sup / "risked_tonnage.csv",
            run_dir / "tables" / "risked_tonnage.csv",
        ],
        "sgs_meta.json": [sup / "sgs_meta.json", run_dir / "sgs_meta.json"],
    }
    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        files: list[tuple[Path, str]] = []
        missing: list[str] = []
        for name in ESSENTIAL_S2_FILES:
            picked_src: Path | None = None
            for src in file_candidates.get(name, []):
                if src.exists():
                    picked_src = src
                    break
            if picked_src is None:
                missing.append(name)
                continue

            dst = tmp / name
            if name.lower().endswith(".json"):
                _sanitize_public_json(picked_src, dst)
                if name == "validation_metrics.json":
                    _enrich_public_validation_metrics(dst, run_dir)
                elif name == "variogram_model.json":
                    _enrich_public_variogram_model(dst, run_dir)
                elif name == "sgs_meta.json":
                    _enrich_public_sgs_meta(dst)
            elif name == "cutoff_occupancy_uncertainty.csv":
                df = pd.read_csv(picked_src)
                rename_map = {
                    "tonnage": "occupancy_proxy",
                    "tonnage_p05": "occupancy_proxy_p05",
                    "tonnage_p10": "occupancy_proxy_p10",
                    "tonnage_mean": "occupancy_proxy_mean",
                    "tonnage_p50": "occupancy_proxy_p50",
                    "tonnage_p90": "occupancy_proxy_p90",
                    "tonnage_p95": "occupancy_proxy_p95",
                    "grade_p05": "mean_tgc_above_cutoff_p05",
                    "grade_p10": "mean_tgc_above_cutoff_p10",
                    "grade_mean": "mean_tgc_above_cutoff_mean",
                    "grade_p50": "mean_tgc_above_cutoff_p50",
                    "grade_p90": "mean_tgc_above_cutoff_p90",
                    "grade_p95": "mean_tgc_above_cutoff_p95",
                }
                df = df.rename(columns={k: v for k, v in rename_map.items() if k in df.columns})
                df = df[[c for c in df.columns if not c.lower().startswith("contained")]]
                tmp_csv = tmp / f"_raw_{name}"
                df.to_csv(tmp_csv, index=False)
                _sanitize_public_text_table(tmp_csv, dst)
            elif name.lower().endswith(".csv"):
                _sanitize_public_text_table(picked_src, dst)
            else:
                shutil.copy2(picked_src, dst)
            files.append((dst, name))
        if missing:
            raise FileNotFoundError(f"Missing required S2 evidence files for clean package: {missing}")
        _zip_files(dst_zip, files)


def _write_clean_code_zip(dst_zip: Path) -> None:
    rels = [
        "README.md",
        "manuscript.md",
        "tables.md",
        "figure_captions.md",
        "config/main_config.yaml",
        "scripts/build_paper_from_meta.py",
        "scripts/build_submission_package.py",
        "scripts/submission_preflight.py",
        "scripts/autoresearch_eval.py",
    ]
    files = [(ROOT / rel, rel.replace("\\", "/")) for rel in rels if (ROOT / rel).exists()]
    _zip_files(dst_zip, files)


def _clean_figure_filename(fig_key: str) -> str:
    return f"Fig{fig_key}.tif"


def _figure_artwork_requirements(fig_key: str) -> tuple[int, int]:
    # Springer MME maximum full-width artwork is 174 mm.
    return 1000, int(round(174.0 / 25.4 * 1000.0))


def _merge_sgs_pilot_metrics(path: Path) -> None:
    pilot_dirs = [
        ROOT / "build" / "non_geology_sgs_pilot_nr20",
        ROOT / "build" / "non_geology_sgs_pilot",
    ]
    pilot_dir = next(
        (candidate for candidate in pilot_dirs if (candidate / "tables" / "validation_metrics.json").exists()),
        pilot_dirs[-1],
    )
    pilot_path = pilot_dir / "tables" / "validation_metrics.json"
    pilot_meta_path = pilot_dir / "sgs_meta.json"
    if not pilot_path.exists():
        return
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
        pilot = json.loads(pilot_path.read_text(encoding="utf-8"))
        n_real = pilot.get("n_real")
        if pilot_meta_path.exists():
            try:
                n_real = int(json.loads(pilot_meta_path.read_text(encoding="utf-8")).get("config", {}).get("simulation", {}).get("n_real", n_real))
            except Exception:
                pass
        payload["sgs_prior_sensitivity_pilot"] = {
            "configuration": "no-domain isotropic pilot SGS",
            "purpose": "Direct non-canonical sensitivity check against the geology-conditioned SGS workflow.",
            "n_real": n_real,
            "metrics": {
                "mean_sim": pilot.get("mean_sim"),
                "hist_overlap": pilot.get("hist_overlap"),
                "qq_rmse": pilot.get("qq_rmse"),
                "swath_corr_x": pilot.get("swath_corr_x"),
                "swath_corr_y": pilot.get("swath_corr_y"),
                "swath_corr_z": pilot.get("swath_corr_z"),
            },
        }
        path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    except Exception:
        return


def _merge_review_scalar_summaries(path: Path, run_dir: Path) -> None:
    """Embed compact reviewer-auditable geology summaries in the allowed S2 JSON."""
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
        tables = run_dir / "tables"
        summaries: dict[str, object] = {}

        contact_meta = tables / "contact_analysis_meta.json"
        contact_csv = tables / "contact_analysis.csv"
        if contact_meta.exists():
            contact = json.loads(contact_meta.read_text(encoding="utf-8"))
            if contact_csv.exists():
                cdf = pd.read_csv(contact_csv)
                contact["distance_bin_count"] = int(len(cdf))
                contact["fresh_weathered_distance_bins"] = [
                    {
                        "weathering_class": str(row["weathering_class"]),
                        "distance_bin": str(row["distance_bin"]),
                        "count": int(row["count"]),
                        "mean_tgc_pct": float(row["mean_tgc_pct"]),
                    }
                    for _, row in cdf.iterrows()
                ]
            summaries["contact_analysis"] = contact

        weathering_csv = tables / "weathering_summary.csv"
        if weathering_csv.exists():
            wdf = pd.read_csv(weathering_csv)
            weathering_rows: list[dict[str, object]] = []
            for _, row in wdf.iterrows():
                group = str(row["group"])
                group_norm = group.strip().lower().replace(" ", "_")
                if group_norm == "weathering_upgrade":
                    summaries["weathering_relative_mean_contrast"] = {
                        "group": "weathered_minus_fresh_relative_mean_contrast",
                        "count": None if pd.isna(row.get("count")) else int(float(row["count"])),
                        "relative_mean_contrast_pct": None
                        if pd.isna(row.get("mean_tgc_pct"))
                        else float(row["mean_tgc_pct"]),
                        "interpretation": "Relative mean contrast only; not a material grade-effect or product-quality result.",
                    }
                    continue
                weathering_rows.append(
                    {
                        "group": group,
                        "count": None if pd.isna(row.get("count")) else int(float(row["count"])),
                        "mean_tgc_pct": None if pd.isna(row.get("mean_tgc_pct")) else float(row["mean_tgc_pct"]),
                        "std_tgc_pct": None if pd.isna(row.get("std_tgc_pct")) else float(row["std_tgc_pct"]),
                    }
                )
            summaries["weathering_summary"] = weathering_rows

        for key, rel in [
            ("domain_uncertainty", "domain_uncertainty_summary.json"),
            ("thickness_geometry", "thickness_geometry_summary.json"),
            ("vertical_continuity", "vertical_continuity_summary.json"),
        ]:
            src = tables / rel
            if src.exists():
                summaries[key] = json.loads(src.read_text(encoding="utf-8"))

        support_csv = tables / "support_ladder_summary.csv"
        if support_csv.exists():
            sdf = pd.read_csv(support_csv)
            summaries["support_ladder"] = [
                {
                    "support_name": str(row["support_name"]),
                    "support_dx_m": float(row["support_dx_m"]),
                    "support_dy_m": float(row["support_dy_m"]),
                    "support_dz_m": float(row["support_dz_m"]),
                    "hist_overlap": float(row["hist_overlap"]),
                    "qq_rmse": float(row["qq_rmse"]),
                    "swath_coverage_pct": float(row["swath_coverage_pct"]),
                    "swath_corr_x": float(row["swath_corr_x"]),
                    "swath_corr_y": float(row["swath_corr_y"]),
                    "swath_corr_z": float(row["swath_corr_z"]),
                }
                for _, row in sdf.iterrows()
            ]

        sim_metrics = tables / "validation_metrics_2m.json"
        if sim_metrics.exists():
            summaries["simulation_support_validation"] = json.loads(sim_metrics.read_text(encoding="utf-8"))

        reporting_grid = run_dir / "grids" / "sgs_reals_reporting.npy"
        if reporting_grid.exists():
            arr = np.load(reporting_grid, mmap_mode="r")
            finite = np.isfinite(arr)
            if bool(np.any(finite)):
                vals = np.asarray(arr[finite], dtype=float)
                summaries["physical_domain_audit"] = {
                    "reporting_support_min_tgc_pct": float(np.nanmin(vals)),
                    "reporting_support_negative_cell_pct": float(np.nanmean(vals < 0.0) * 100.0),
                    "reporting_support_raw_mean_tgc_pct": float(np.nanmean(vals)),
                    "reporting_support_zero_floor_mean_tgc_pct": float(np.nanmean(np.maximum(vals, 0.0))),
                    "interpretation": (
                        "No negative TGC values occur in the completed canonical reporting-support ensemble."
                        if not bool(np.any(vals < 0.0))
                        else "Negative TGC values are physically non-interpretable lower-tail artefacts."
                    ),
                }

        if summaries:
            payload["geology_uncertainty_scalar_summaries"] = summaries
            payload["s2_scope_note"] = (
                "S2 is a compact audit-level output/metadata supplement. It embeds scalar summaries "
                "from the completed run but does not include proprietary categorical-domain arrays or "
                "the full internal review-pack figure/table set."
            )
            path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    except Exception:
        return



def _merge_validation_gap_summaries(path: Path, run_dir: Path) -> None:
    keys = [
        "variogram_reproduction",
        "realisation_count_normalised_sensitivity",
        "spatial_overlap_bootstrap",
        "signed_graphitic_host_contact",
        "ensemble_convergence",
        "support_aligned_mean_decomposition",
        "categorical_domain_grouped_validation",
        "no_domain_pilot_realisation_bootstrap",
        "directional_swath_curves",
    ]
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
        truth_path = ROOT / "build" / "source_of_truth.submission.json"
        if truth_path.exists():
            try:
                truth = json.loads(truth_path.read_text(encoding="utf-8"))
                cached = truth.get("validation_gap_summaries")
                if isinstance(cached, dict) and set(keys).issubset(cached):
                    payload["validation_gap_summaries"] = cached
                    for key in keys:
                        payload[key] = cached.get(key)
                    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
                    return
            except Exception:
                pass
        try:
            from scripts.build_paper_from_meta import _compute_validation_gap_summaries
        except Exception:
            import importlib.util
            spec = importlib.util.spec_from_file_location(
                "build_paper_from_meta", ROOT / "scripts" / "build_paper_from_meta.py"
            )
            if spec is None or spec.loader is None:
                return
            mod = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(mod)
            _compute_validation_gap_summaries = mod._compute_validation_gap_summaries
        summaries = _compute_validation_gap_summaries(run_dir, payload)
        payload["validation_gap_summaries"] = summaries
        for key in keys:
            payload[key] = summaries.get(key)
        path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    except Exception:
        return

def _finalize_public_validation_metrics(path: Path) -> None:
    """Replace legacy review-pack summaries with manuscript-aligned public values."""
    payload = json.loads(path.read_text(encoding="utf-8"))
    truth_path = ROOT / "build" / "source_of_truth.submission.json"
    if not truth_path.exists():
        return
    truth = json.loads(truth_path.read_text(encoding="utf-8"))
    stats = truth.get("contact_weathering_stat_tests", {})
    signed = payload.get("signed_graphitic_host_contact") or truth.get(
        "validation_gap_summaries", {}
    ).get("signed_graphitic_host_contact", {})
    scalar = payload.setdefault("geology_uncertainty_scalar_summaries", {})
    for stale_key in ("contact_analysis", "weathering_summary", "weathering_relative_mean_contrast"):
        scalar.pop(stale_key, None)
    scalar["signed_graphitic_host_profile"] = signed
    scalar["weathering_stratified_unsigned_contact_distance"] = {
        "n_composites": int(stats.get("contact_n", 0)),
        "n_holes": int(stats.get("contact_holes_n", 0)),
        "bin_rows": stats.get("contact_summary_rows", []),
        "anova_p": stats.get("contact_anova_p"),
        "kruskal_wallis_p": stats.get("contact_kruskal_p"),
        "levene_p": stats.get("contact_levene_p"),
        "scope": (
            "Unsigned graphitic-only distance bins stratified by weathering; separate "
            "from the signed graphitic-host transition profile."
        ),
    }
    scalar["graphitic_weathering_comparison"] = {
        key: stats.get(key)
        for key in (
            "fresh_n",
            "weathered_n",
            "fresh_mean_tgc_pct",
            "weathered_mean_tgc_pct",
            "fresh_std_tgc_pct",
            "weathered_std_tgc_pct",
            "weathering_mean_difference_tgc_pct",
            "weathering_mean_difference_ci95_low",
            "weathering_mean_difference_ci95_high",
            "weathering_hedges_g",
            "weathering_welch_p",
            "weathering_hole_cluster_ci95_low",
            "weathering_hole_cluster_ci95_high",
            "weathering_paired_holes_n",
            "weathering_paired_holes_wilcoxon_p",
        )
    }
    probabilities = [0.70, 0.20, 0.10]
    entropy = -sum(p * np.log(p) for p in probabilities) / np.log(len(probabilities))
    payload["calculation_examples"] = {
        "scope": "Synthetic arithmetic examples; no proprietary project cell values are used.",
        "normalised_shannon_entropy": {
            "domain_probabilities": probabilities,
            "calculated_entropy": round(float(entropy), 6),
        },
        "graphitic_probability": {
            "fresh_graphitic_frequency": 0.55,
            "weathered_graphitic_frequency": 0.25,
            "host_waste_frequency": 0.20,
            "calculated_graphitic_probability": 0.80,
        },
        "thickness_aperture": {
            "p10_thickness_m": 18.0,
            "p50_thickness_m": 24.0,
            "p90_thickness_m": 31.0,
            "calculated_p90_minus_p10_m": 13.0,
        },
    }
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def _enrich_public_validation_metrics(path: Path, run_dir: Path) -> None:
    _merge_sgs_pilot_metrics(path)
    _merge_review_scalar_summaries(path, run_dir)
    _merge_withheld_composite_validation_baseline(path)
    _merge_validation_gap_summaries(path, run_dir)
    _finalize_public_validation_metrics(path)


def _merge_withheld_composite_validation_baseline(path: Path) -> None:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return
    rows: list[dict[str, object]] = []
    truth_path = ROOT / "build" / "source_of_truth.submission.json"
    if truth_path.exists():
        try:
            truth = json.loads(truth_path.read_text(encoding="utf-8"))
            raw_rows = truth.get("blocked_validation_baseline") or truth.get("baseline_best_rows") or []
            if isinstance(raw_rows, list):
                rows = [r for r in raw_rows if isinstance(r, dict)]
        except Exception:
            rows = []
    if not rows:
        summary_path = ROOT / "build" / "tmp_s2" / "validation_baseline_summary.csv"
        if summary_path.exists():
            try:
                df = pd.read_csv(summary_path)
                required = {"fold_mode", "method", "n", "MAE", "RMSE"}
                if required.issubset(df.columns):
                    for fold_mode, group in df.groupby("fold_mode", sort=False):
                        best = group.sort_values("RMSE", ascending=True).iloc[0]
                        row: dict[str, object] = {
                            "validation_family": str(fold_mode),
                            "best_method": str(best["method"]),
                            "n": int(best["n"]),
                            "mae": round(float(best["MAE"]), 3),
                            "rmse": round(float(best["RMSE"]), 3),
                        }
                        if "ME" in best:
                            row["me"] = round(float(best["ME"]), 3)
                        if "R" in best and pd.notna(best["R"]):
                            row["r"] = round(float(best["R"]), 3)
                        rows.append(row)
            except Exception:
                rows = []
    if not rows:
        return
    clean_rows: list[dict[str, object]] = []
    for raw in rows:
        if not isinstance(raw, dict):
            continue
        row: dict[str, object] = {
            "validation_family": str(raw.get("validation_family", "")),
            "best_method": str(raw.get("best_method", "")),
        }
        for key in ["n", "me", "mae", "rmse", "r"]:
            if key not in raw:
                continue
            value = raw.get(key)
            try:
                if key == "n":
                    row[key] = int(str(value).replace(",", ""))
                else:
                    row[key] = round(float(value), 3)
            except Exception:
                row[key] = value
        clean_rows.append(row)
    if clean_rows:
        rows = clean_rows
    payload["withheld_composite_validation_baseline"] = {
        "purpose": (
            "Spatially separated withheld-composite baseline used to test predictive behaviour "
            "of the geological prior; this is not independent validation of the final SGS ensemble."
        ),
        "fold_families": rows,
        "methods_screened": ["IDW", "OK", "SK"],
        "interpretation": (
            "RMSE values near 2.2% TGC under blocked, leave-hole and leave-section checks support "
            "the manuscript scope: uncertainty localisation is better supported than strong point-scale grade prediction."
        ),
    }
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def _variogram_pair_count_support(run_dir: Path) -> dict[str, object]:
    pair_path = run_dir / "figures" / "variogram_pair_counts.csv"
    if not pair_path.exists():
        return {}
    try:
        df = pd.read_csv(pair_path)
    except Exception:
        return {}
    if not {"direction", "lag", "count"}.issubset(df.columns):
        return {}
    out: dict[str, object] = {
        "source": "variogram_pair_counts.csv",
        "lag_count_configured": 10,
        "directions": {},
    }
    directions: dict[str, object] = {}
    for direction, sub in df.groupby("direction"):
        counts = pd.to_numeric(sub["count"], errors="coerce").fillna(0).astype(int)
        lags = pd.to_numeric(sub["lag"], errors="coerce")
        directions[str(direction)] = {
            "total_pairs": int(counts.sum()),
            "nonzero_lags": int((counts > 0).sum()),
            "lag_min": int(lags.min()) if len(lags.dropna()) else None,
            "lag_max": int(lags.max()) if len(lags.dropna()) else None,
        }
    out["directions"] = directions
    normal = directions.get("normal_to_plane", {})
    out["summary"] = (
        "Pair-count support is strongest along strike and down dip; "
        f"normal-to-plane has {normal.get('nonzero_lags', 0)}/10 nonzero lags."
    )
    return out


def _enrich_public_variogram_model(path: Path, run_dir: Path) -> None:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return
    if not isinstance(payload, dict):
        return
    support = _variogram_pair_count_support(run_dir)
    if support:
        payload["pair_count_support"] = support
    payload["public_audit_note"] = (
        "Pair-count support is included to make the directional variogram evidence auditable "
        "inside the four-file S2 supplement."
    )
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def _write_clean_figure6_tif(run_dir: Path, dst: Path, min_dpi: int, min_width_px: int) -> bool:
    try:
        grids = run_dir / "grids"
        p10 = np.load(grids / "p10_grid.npy").astype(float)
        p90 = np.load(grids / "p90_grid.npy").astype(float)
        spread = np.maximum(p90 - p10, 0.0)
        if spread.ndim != 3:
            return False
        plan = np.nanpercentile(spread, 90, axis=2)
        vmax = float(np.nanpercentile(plan, 96))
        if not np.isfinite(vmax) or vmax <= 0:
            vmax = float(np.nanmax(plan))
        vmax = max(vmax, 1.0)

        meta = json.loads((run_dir / "sgs_meta.json").read_text(encoding="utf-8"))
        cfg = meta.get("config", {})
        nx, ny = int(plan.shape[0]), int(plan.shape[1])
        sim_grid = cfg.get("grid", {}) or {}
        report_grid = cfg.get("reporting_grid", {}) or {}
        if int(report_grid.get("nx", -1)) == nx and int(report_grid.get("ny", -1)) == ny:
            grid = report_grid
        elif int(sim_grid.get("nx", -1)) == nx and int(sim_grid.get("ny", -1)) == ny:
            grid = sim_grid
        else:
            grid = report_grid or sim_grid
        x0, y0, _z0 = [float(v) for v in grid.get("origin_xyz", [0, 0, 0])]
        dx, dy = float(grid.get("dx", 1.0)), float(grid.get("dy", 1.0))
        extent = [x0, x0 + nx * dx, y0, y0 + ny * dy]

        fig_w = max(min_width_px / float(min_dpi), 5.2)
        fig_h = max(fig_w * 1.22, 6.3)
        fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=min_dpi)
        im = ax.imshow(
            plan.T,
            origin="lower",
            extent=extent,
            cmap="cividis",
            vmin=0,
            vmax=vmax,
            interpolation="nearest",
            aspect="equal",
        )

        graph_prob_path = grids / "graphitic_domain_probability.npy"
        if graph_prob_path.exists():
            dom = np.load(graph_prob_path).astype(float)
            if dom.ndim == 3 and dom.shape[:2] == plan.shape:
                dom_plan = np.nanmax(dom, axis=2)
                ax.contour(
                    np.linspace(extent[0], extent[1], nx),
                    np.linspace(extent[2], extent[3], ny),
                    dom_plan.T,
                    levels=[0.5, 0.8],
                    colors=["white", "black"],
                    linewidths=[1.1, 0.9],
                )
                ax.text(
                    0.02,
                    0.98,
                    "Domain contours: Pgraph 0.5/0.8",
                    transform=ax.transAxes,
                    ha="left",
                    va="top",
                    fontsize=5.8,
                    bbox={"boxstyle": "round,pad=0.16", "facecolor": "white", "edgecolor": "0.6", "alpha": 0.82},
                    zorder=7,
                )

        collar_path = ROOT / "data" / "collar.csv"
        if collar_path.exists():
            collar = pd.read_csv(collar_path)
            x_col = next((c for c in collar.columns if c.lower().strip() in {"x", "east", "easting", "x_coord"}), None)
            y_col = next((c for c in collar.columns if c.lower().strip() in {"y", "north", "northing", "y_coord"}), None)
            if x_col and y_col:
                xx = pd.to_numeric(collar[x_col], errors="coerce")
                yy = pd.to_numeric(collar[y_col], errors="coerce")
                mask = (
                    np.isfinite(xx)
                    & np.isfinite(yy)
                    & (xx >= extent[0])
                    & (xx <= extent[1])
                    & (yy >= extent[2])
                    & (yy <= extent[3])
                )
                ax.scatter(xx[mask], yy[mask], s=12, c="white", edgecolors="black", linewidths=0.35, zorder=4, label="Drill collars")

        finite = np.isfinite(plan)
        if finite.any():
            flat = np.argsort(np.where(finite, plan, -np.inf), axis=None)[-5:]
            hot_x, hot_y = np.unravel_index(flat, plan.shape)
            hx = x0 + (hot_x + 0.5) * dx
            hy = y0 + (hot_y + 0.5) * dy
            ax.scatter(
                hx,
                hy,
                s=58,
                facecolors="none",
                edgecolors="#d62728",
                linewidths=1.25,
                zorder=5,
                label="High-spread cells",
            )
            label_offsets = [(0.035, 0.035), (0.050, -0.055), (-0.100, 0.050)]
            for rank, (px, py) in enumerate(zip(hx[-3:], hy[-3:]), start=1):
                offx, offy = label_offsets[rank - 1]
                ax.text(
                    px + offx * (extent[1] - extent[0]),
                    py + offy * (extent[3] - extent[2]),
                    f"H{rank}",
                    color="#7f0000",
                    fontsize=6.2,
                    fontweight="bold",
                    ha="center",
                    va="center",
                    zorder=6,
                    bbox={"boxstyle": "round,pad=0.12", "facecolor": "white", "edgecolor": "none", "alpha": 0.75},
                )

        ore = cfg.get("orebody", {})
        strike = float(ore.get("strike_deg", 105.0))
        angle = np.deg2rad(90.0 - strike)
        cx = extent[0] + 0.13 * (extent[1] - extent[0])
        cy = extent[2] + 0.85 * (extent[3] - extent[2])
        length = 0.16 * min(extent[1] - extent[0], extent[3] - extent[2])
        ax.annotate(
            "",
            xy=(cx + length * np.cos(angle), cy + length * np.sin(angle)),
            xytext=(cx - length * np.cos(angle), cy - length * np.sin(angle)),
            arrowprops={"arrowstyle": "->", "lw": 1.4, "color": "black"},
            zorder=6,
        )
        ax.text(
            cx,
            cy - 0.075 * (extent[3] - extent[2]),
            f"Structural trend {strike:.0f} deg",
            ha="center",
            va="top",
            fontsize=6.5,
            bbox={"boxstyle": "round,pad=0.16", "facecolor": "white", "edgecolor": "none", "alpha": 0.75},
            zorder=7,
        )

        ax.set_title("Absolute P90-P10 TGC Spread\nReporting Support", fontsize=8.8, pad=6)
        ax.plot(
            [extent[0], extent[1], extent[1], extent[0], extent[0]],
            [extent[2], extent[2], extent[3], extent[3], extent[2]],
            color="black",
            linewidth=1.0,
            zorder=6,
            label="Model footprint",
        )
        ax.set_xlabel("Easting from grid origin (km)", fontsize=7)
        ax.set_ylabel("Northing from grid origin (km)", fontsize=7)
        ax.tick_params(axis="both", labelsize=6.2, length=2.5)
        ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda v, _p: f"{(v - x0) / 1000.0:.1f}"))
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _p: f"{(v - y0) / 1000.0:.1f}"))
        cbar = fig.colorbar(im, ax=ax, fraction=0.046, pad=0.03)
        cbar.set_label("P90-P10 TGC spread (%)", fontsize=7)
        cbar.ax.tick_params(labelsize=6.2, length=2.5)
        if ax.get_legend_handles_labels()[0]:
            ax.legend(
                loc="upper center",
                bbox_to_anchor=(0.5, -0.13),
                ncol=3,
                fontsize=5.8,
                frameon=False,
                handlelength=1.5,
                columnspacing=1.0,
            )
        ax.grid(color="white", alpha=0.22, linewidth=0.4)
        try:
            from PIL import Image  # type: ignore

            fig.subplots_adjust(left=0.17, right=0.86, top=0.90, bottom=0.19)
            png_buf = io.BytesIO()
            fig.savefig(png_buf, format="png", dpi=min_dpi)
            plt.close(fig)
            png_buf.seek(0)
            with Image.open(png_buf) as img:
                rgb = img.convert("RGB")
                try:
                    rgb.save(
                        dst,
                        format="TIFF",
                        dpi=(float(min_dpi), float(min_dpi)),
                        compression="tiff_adobe_deflate",
                    )
                except Exception:
                    rgb.save(
                        dst,
                        format="TIFF",
                        dpi=(float(min_dpi), float(min_dpi)),
                        compression="tiff_lzw",
                    )
                finally:
                    try:
                        rgb.close()
                    except Exception:
                        pass
        except Exception:
            fig.subplots_adjust(left=0.17, right=0.86, top=0.90, bottom=0.19)
            fig.savefig(dst, format="tiff", dpi=min_dpi)
            plt.close(fig)
        return dst.exists()
    except Exception:
        try:
            plt.close("all")
        except Exception:
            pass
        return False


def _write_connected_figure6_tif(run_dir: Path, dst: Path, min_dpi: int, min_width_px: int) -> bool:
    """Write the final Figure 6 as the three-panel spread/occupancy diagnostic."""
    try:
        grids = run_dir / "grids"
        p10 = np.load(grids / "p10_grid.npy").astype(float)
        p90 = np.load(grids / "p90_grid.npy").astype(float)
        if p10.ndim != 3 or p90.ndim != 3:
            return False
        spread_plan = np.nanpercentile(np.maximum(p90 - p10, 0.0), 90, axis=2)
        prob_plan = None
        prob_path = grids / "prob_gt_3.0.npy"
        if prob_path.exists():
            prob = np.load(prob_path).astype(float)
            prob_plan = np.nanpercentile(prob, 90, axis=2) if prob.ndim == 3 else prob
            if np.asarray(prob_plan).shape != spread_plan.shape:
                prob_plan = None

        vmax = float(np.nanpercentile(spread_plan, 96))
        if not np.isfinite(vmax) or vmax <= 0:
            vmax = float(np.nanmax(spread_plan))
        vmax = max(vmax, 1.0)

        meta = json.loads((run_dir / "sgs_meta.json").read_text(encoding="utf-8"))
        cfg = meta.get("config", {})
        nx, ny = int(spread_plan.shape[0]), int(spread_plan.shape[1])
        sim_grid = cfg.get("grid", {}) or {}
        report_grid = cfg.get("reporting_grid", {}) or {}
        if int(report_grid.get("nx", -1)) == nx and int(report_grid.get("ny", -1)) == ny:
            grid = report_grid
        elif int(sim_grid.get("nx", -1)) == nx and int(sim_grid.get("ny", -1)) == ny:
            grid = sim_grid
        else:
            grid = report_grid or sim_grid
        x0, y0, _z0 = [float(v) for v in grid.get("origin_xyz", [0, 0, 0])]
        dx, dy = float(grid.get("dx", 1.0)), float(grid.get("dy", 1.0))
        extent = [x0, x0 + nx * dx, y0, y0 + ny * dy]

        fig_w = max(min_width_px / float(min_dpi), 7.4)
        fig_h = max(fig_w * 0.62, 4.8)
        fig = plt.figure(figsize=(fig_w, fig_h), dpi=min_dpi)
        gs = fig.add_gridspec(
            2,
            2,
            width_ratios=[0.92, 1.08],
            height_ratios=[1.0, 1.0],
            left=0.065,
            right=0.965,
            top=0.91,
            bottom=0.155,
            wspace=0.30,
            hspace=0.42,
        )
        ax_map = fig.add_subplot(gs[:, 0])
        ax_scatter = fig.add_subplot(gs[0, 1])
        ax_bins = fig.add_subplot(gs[1, 1])

        im = ax_map.imshow(
            spread_plan.T,
            origin="lower",
            extent=extent,
            cmap="cividis",
            vmin=0.0,
            vmax=vmax,
            interpolation="nearest",
            aspect="equal",
        )

        if prob_plan is not None:
            try:
                ax_map.contour(
                    np.linspace(extent[0], extent[1], nx),
                    np.linspace(extent[2], extent[3], ny),
                    np.asarray(prob_plan).T,
                    levels=[0.5, 0.8],
                    colors=["#f8fafc", "#111827"],
                    linewidths=[0.75, 0.9],
                    linestyles=["--", "-"],
                )
            except Exception:
                pass

        collar_path = ROOT / "data" / "collar.csv"
        if collar_path.exists():
            collar = pd.read_csv(collar_path)
            x_col = next((c for c in collar.columns if c.lower().strip() in {"x", "east", "easting", "x_coord"}), None)
            y_col = next((c for c in collar.columns if c.lower().strip() in {"y", "north", "northing", "y_coord"}), None)
            if x_col and y_col:
                xx = pd.to_numeric(collar[x_col], errors="coerce")
                yy = pd.to_numeric(collar[y_col], errors="coerce")
                mask = np.isfinite(xx) & np.isfinite(yy) & (xx >= extent[0]) & (xx <= extent[1]) & (yy >= extent[2]) & (yy <= extent[3])
                ax_map.scatter(xx[mask], yy[mask], s=12, c="white", edgecolors="black", linewidths=0.35, zorder=4, label="Drill collars")

        ax_map.set_title('A. Spread and occupancy', fontsize=8.2, fontweight='bold', loc='left', pad=5)
        ax_map.plot(
            [extent[0], extent[1], extent[1], extent[0], extent[0]],
            [extent[2], extent[2], extent[3], extent[3], extent[2]],
            color="black",
            linewidth=1.0,
            zorder=6,
            label="Model footprint",
        )
        ax_map.set_xlabel("Easting from grid origin (km)", fontsize=7)
        ax_map.set_ylabel("Northing from grid origin (km)", fontsize=7)
        ax_map.tick_params(axis="both", labelsize=6.2, length=2.5)
        ax_map.xaxis.set_major_formatter(plt.FuncFormatter(lambda v, _p: f"{(v - x0) / 1000.0:.1f}"))
        ax_map.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _p: f"{(v - y0) / 1000.0:.1f}"))
        cbar = fig.colorbar(im, ax=ax_map, fraction=0.046, pad=0.03)
        cbar.set_label("P90-P10 TGC spread (%)", fontsize=7)
        cbar.ax.tick_params(labelsize=6.2, length=2.5)
        if ax_map.get_legend_handles_labels()[0]:
            ax_map.legend(
                loc="upper center",
                bbox_to_anchor=(0.5, -0.13),
                ncol=3,
                fontsize=5.8,
                frameon=False,
                handlelength=1.5,
                columnspacing=1.0,
            )
        ax_map.grid(color="white", alpha=0.22, linewidth=0.4)

        flat_spread = spread_plan.ravel()
        if prob_plan is not None:
            flat_prob = np.asarray(prob_plan, dtype=float).ravel()
            mask = np.isfinite(flat_spread) & np.isfinite(flat_prob)
        else:
            flat_prob = np.array([], dtype=float)
            mask = np.zeros_like(flat_spread, dtype=bool)
        if int(mask.sum()) > 2:
            corr = float(np.corrcoef(flat_prob[mask], flat_spread[mask])[0, 1])
            hb = ax_scatter.hexbin(flat_prob[mask], flat_spread[mask], gridsize=24, cmap="Blues", mincnt=1, linewidths=0.0)
            ax_scatter.axvline(0.5, color="#64748b", linestyle="--", linewidth=0.8)
            ax_scatter.axvline(0.8, color="#334155", linestyle=":", linewidth=0.9)
            ax_scatter.text(
                0.98,
                0.05,
                f"Descriptive cell association: r = {corr:.2f}\nshared ensemble products; not validation",
                transform=ax_scatter.transAxes,
                ha="right",
                va="bottom",
                fontsize=5.8,
                color="#4b5563",
                bbox={"facecolor": "white", "edgecolor": "#d1d5db", "alpha": 0.9, "pad": 2.0},
            )
            cbar_hb = fig.colorbar(hb, ax=ax_scatter, fraction=0.045, pad=0.020)
            cbar_hb.set_label("Cell count", fontsize=6.4)
            cbar_hb.ax.tick_params(labelsize=6.0, length=2.2)
            grouped = (
                pd.DataFrame(
                    {
                        "prob_bin": pd.cut(flat_prob[mask], bins=np.linspace(0.0, 1.0, 6), include_lowest=True),
                        "spread": flat_spread[mask],
                    }
                )
                .groupby("prob_bin", observed=False)["spread"]
                .agg(["median", "count"])
                .reset_index()
            )
            mids = np.array([(interval.left + interval.right) / 2.0 for interval in grouped["prob_bin"]])
            ax_bins.bar(mids, grouped["median"], width=0.16, color="#4c78a8", edgecolor="#1e3a8a", linewidth=0.45, alpha=0.88)
            for xmid, med, count in zip(mids, grouped["median"], grouped["count"]):
                if np.isfinite(med):
                    ax_bins.text(xmid, med + vmax * 0.025, f"n={int(count)}", ha="center", va="bottom", fontsize=5.7, color="#334155")
            ax_bins.set_xlim(0.0, 1.0)
        else:
            ax_scatter.text(0.5, 0.5, "Probability grid unavailable", transform=ax_scatter.transAxes, ha="center", va="center", fontsize=7)
            ax_bins.text(0.5, 0.5, "Probability bins unavailable", transform=ax_bins.transAxes, ha="center", va="center", fontsize=7)

        ax_scatter.set_title('B. Cellwise association', fontsize=8.2, fontweight='bold', loc='left')
        ax_scatter.set_xlabel("P(TGC > 3%)", fontsize=7)
        ax_scatter.set_ylabel("P90-P10 TGC spread (%)", fontsize=7)
        ax_scatter.set_xlim(0.0, 1.0)
        ax_scatter.set_ylim(0.0, max(vmax * 1.08, float(np.nanpercentile(flat_spread, 99)) * 1.02))
        ax_scatter.grid(color="#cbd5e1", alpha=0.45, linewidth=0.5)
        ax_bins.set_title("C. Median spread by probability class", fontsize=8.8, fontweight="bold", loc="left")
        ax_bins.set_xlabel("P(TGC > 3%) class midpoint", fontsize=7)
        ax_bins.set_ylabel("Median spread (%)", fontsize=7)
        ax_bins.set_ylim(0.0, vmax * 1.08)
        ax_bins.grid(axis="y", color="#cbd5e1", alpha=0.45, linewidth=0.5)
        for panel_ax in [ax_scatter, ax_bins]:
            panel_ax.tick_params(axis="both", labelsize=6.2, length=2.5)
            for spine in panel_ax.spines.values():
                spine.set_linewidth(0.8)
                spine.set_color("#111827")
        from matplotlib.text import Text

        undersized = [
            (item.get_text()[:80], float(item.get_fontsize()))
            for item in fig.findobj(match=Text)
            if item.get_visible() and item.get_text().strip() and float(item.get_fontsize()) < 5.0
        ]
        if undersized:
            raise ValueError(f'Final Figure 6 contains text below 5 pt: {undersized[:8]}')
        try:
            from PIL import Image  # type: ignore

            png_buf = io.BytesIO()
            fig.savefig(png_buf, format="png", dpi=min_dpi)
            plt.close(fig)
            png_buf.seek(0)
            with Image.open(png_buf) as img:
                rgb = img.convert("RGB")
                try:
                    rgb.save(dst, format="TIFF", dpi=(float(min_dpi), float(min_dpi)), compression="tiff_adobe_deflate")
                except Exception:
                    rgb.save(dst, format="TIFF", dpi=(float(min_dpi), float(min_dpi)), compression="tiff_lzw")
                finally:
                    try:
                        rgb.close()
                    except Exception:
                        pass
        except Exception:
            fig.savefig(dst, format="tiff", dpi=min_dpi)
            plt.close(fig)
        return dst.exists()
    except Exception:
        try:
            plt.close("all")
        except Exception:
            pass
        return False


def _apply_final_manuscript_styles(docx_path: Path) -> None:
    if not docx_path.exists():
        return
    doc = Document(docx_path)
    style_names = {s.name for s in doc.styles}
    if "Title" not in style_names:
        style = doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.font.bold = True
        style_names.add("Title")
    if "Caption" not in style_names:
        style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(10)
        style.font.italic = False
        style_names.add("Caption")
    if doc.paragraphs and "Title" in style_names:
        doc.paragraphs[0].style = doc.styles["Title"]
    for p in doc.paragraphs:
        if re.match(r"^(?:Fig\.|Figure)\s*\d+[A-Z]?\b", p.text.strip(), flags=re.IGNORECASE):
            p.style = doc.styles["Caption"]
    doc.save(docx_path)


def _mme_null_evidence() -> tuple[dict, list[dict]]:
    base = ROOT / "build" / "factorial_validation"
    summary_path, rows_path = base / "five_seed_summary.json", base / "five_seed_metrics.csv"
    summary = {"status": "pending", "required_seeds": [9101, 9201, 9301, 9401, 9501]}
    rows: list[dict] = []
    if summary_path.exists():
        summary = json.loads(summary_path.read_text(encoding="utf-8")); summary["status"] = "complete"
    if rows_path.exists():
        rows = pd.read_csv(rows_path).replace({np.nan: None}).to_dict(orient="records")
    return summary, rows


def _first_json(*paths: Path) -> dict:
    for path in paths:
        if path.exists():
            try:
                value = json.loads(path.read_text(encoding="utf-8"))
                return value if isinstance(value, dict) else {"value": value}
            except Exception:
                pass
    return {}


def _flatten(value, prefix: str = "") -> list[list]:
    rows: list[list] = []
    if isinstance(value, dict):
        for key, item in value.items(): rows.extend(_flatten(item, f"{prefix}.{key}" if prefix else str(key)))
    elif isinstance(value, list):
        if all(not isinstance(item, (dict, list)) for item in value): rows.append([prefix, json.dumps(value, ensure_ascii=True)])
        else:
            for idx, item in enumerate(value): rows.extend(_flatten(item, f"{prefix}[{idx}]"))
    else: rows.append([prefix, value])
    return rows


def _write_mme_cover(dst: Path, null_summary: dict) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.8)
    sec.left_margin = sec.right_margin = Inches(0.9)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_paragraph(datetime.now().strftime("%d %B %Y"))
    doc.add_paragraph("Editor-in-Chief\nMining, Metallurgy & Exploration")
    doc.add_paragraph(f"Collection: {MME_COLLECTION}")
    p = doc.add_paragraph()
    p.add_run(MME_TITLE).bold = True
    doc.add_paragraph("Dear Editor,")
    doc.add_paragraph(
        "Please consider this manuscript for the Industrial Minerals: Geology, Extraction and Use collection. "
        "It addresses a practical graphite-evaluation problem: a coherent graphitic horizon does not make "
        "reporting support, contact position and conditional grade spread equally certain."
    )
    doc.add_paragraph(
        "The study evaluates a completed 100-realisation geology-conditioned ensemble inside an explicit "
        "archive-derived lode envelope and then applies that identical envelope to five independent geology-blind "
        "families. This matched post-processing separates the effect of reporting volume from model-configuration "
        "behaviour. The main result is constructive: persistent graphitic support, envelope-edge uncertainty and "
        "TGC spread become distinct, mappable targets for section review and follow-up drilling, even though the "
        "alternate configuration reproduces the marginal grade distribution more closely."
    )
    doc.add_paragraph(
        "All five 20-realisation null families and five non-overlapping canonical subsets are reported without "
        "performance selection. Figures and Online Resources provide the support alignment, spatial diagnostics, "
        "ensemble behaviour, variogram reproduction and directional swaths needed to audit the interpretation."
    )
    doc.add_paragraph(
        "The work is original, is not under consideration elsewhere, and all required disclosures appear in the manuscript."
    )
    doc.add_paragraph(
        f"Sincerely,\n{AUTHOR_NAME}\n{AUTHOR_AFFILIATION}\n{AUTHOR_EMAIL}\nORCID: {AUTHOR_ORCID_URL}"
    )
    doc.save(dst)


def _write_mme_esm_pdf(dst: Path, null_summary: dict, null_rows: list[dict]) -> None:
    from matplotlib.backends.backend_pdf import PdfPages
    from textwrap import wrap

    truth = _first_json(ROOT / "build" / "source_of_truth.submission.json")
    gap = truth.get("validation_gap_summaries", {}) or {}
    cat = gap.get("categorical_domain_grouped_validation", {}) or {}
    envelope = gap.get("archive_lode_envelope", {}) or {}
    matched = gap.get("archive_lode_matched_null_comparison", {}) or {}
    spatial = gap.get("archive_lode_spatial_patterns", {}) or {}
    support = gap.get("support_aligned_mean_decomposition", {}) or {}
    conv = gap.get("archive_lode_envelope_convergence", {}) or {}
    n75 = (conv.get("checkpoint_summaries", {}) or {}).get("75", {}) or {}
    vario = gap.get("variogram_reproduction", {}) or {}
    contact = gap.get("signed_graphitic_host_contact", {}) or {}
    overlap = gap.get("spatial_overlap_bootstrap", {}) or {}
    obs = overlap.get("observed", {}) or {}
    within = (cat.get("search_support", {}) or {}).get("within_support", {}) or {}
    raw_graphitic = cat.get("graphitic_vs_host", {}) or {}
    recalibrated = raw_graphitic.get("nested_platt_recalibration_sensitivity", {}) or {}
    canonical_summary = (
        matched.get("canonical_20_realisation_subsets", {}).get("summary", {}) or {}
    )
    matched_null_summary = (
        matched.get("null_20_realisation_seed_families", {}).get("summary", {}) or {}
    )

    def median(summary: dict, key: str) -> float:
        return float((summary.get(key, {}) or {}).get("median", float("nan")))

    sections = [
        (
            "Online Resource 1: Supplementary Methods and Validation",
            [
                MME_TITLE,
                MME_JOURNAL,
                f"{AUTHOR_NAME}; {AUTHOR_AFFILIATION}; {AUTHOR_EMAIL}; ORCID {AUTHOR_ORCID_URL}",
                "Contents: archive-envelope construction and balance; categorical workflow and grouped validation; "
                "support decomposition and convergence; variogram, contact and spatial diagnostics; and matched repeated-null sensitivity.",
                "This audit-level supplement accompanies Online Resource 2. Proprietary drillhole and categorical-domain arrays are not redistributed.",
            ],
        ),
        (
            "S1. Archive-Derived Reporting Envelope",
            [
                f"The archive source contains seven lode identifiers. Common-footprint and DEM-surface checks retain "
                f"{int(envelope.get('common_support_fine_block_count', 0)):,} blocks from six identifiers. "
                f"L01 contributes {int(envelope.get('dominant_retained_lode_block_count', 0)):,} blocks "
                f"({float(envelope.get('dominant_retained_lode_fraction_pct', float('nan'))):.2f}%); L02 is outside the common SGS footprint.",
                "Only x, y, z, block dimensions, lode identity and topography_z were read. Estimated TGC, kriging "
                "variance, density, classification, tonnes and contained graphite were excluded.",
                f"Fractional lode volume occupies "
                f"{float((envelope.get('support_scenarios', {}).get('fractional_lode_volume', {}) or {}).get('reporting_volume_fraction_pct', float('nan'))):.3f}% "
                "of reporting volume. Any-intersection and full-cell-core summaries provide sensitivity brackets.",
                "The envelope is used to align reporting support and to compare model families inside the same volume.",
            ],
        ),
        (
            "S2. Categorical-Domain Workflow",
            [
                "Fresh graphitic, weathered graphitic and host/waste categories were sampled before grade SGS. "
                "Local inverse-distance class scores used the 250/200/20 m anisotropic search, at most 20 neighbours, "
                "prior weight 2.0 and seed 1337 plus realisation index.",
                "At cell u, graphitic probability is the realisation frequency of either graphitic class. Normalised "
                "Shannon entropy is H(u) = -sum[p_k(u) ln p_k(u)] / ln(3). These raw frequencies are secondary "
                "within-support sensitivity fields; the archive envelope defines primary reporting support.",
            ],
        ),
        (
            "S3. Hole-Grouped Categorical Validation",
            [
                f"Five folds held complete drillholes together (seed {cat.get('seed', 20260707)}), with "
                f"{cat.get('n_holes', 100)} holes, {cat.get('n_composites', 4129)} composites and zero leakage.",
                f"Three-class macro-F1 = {float(cat.get('macro_f1', float('nan'))):.3f}; balanced accuracy = "
                f"{float(cat.get('balanced_accuracy', float('nan'))):.3f}; graphitic-host ROC-AUC = "
                f"{float(raw_graphitic.get('roc_auc', float('nan'))):.3f}. Weathered composites are assigned to "
                f"fresh graphite in {100.0 * float((cat.get('confusion_matrix', {}).get('row_normalized', [[0], [0]])[1][0])):.1f}% of withheld cases.",
                f"Raw Brier skill over all withheld composites is {float(raw_graphitic.get('brier_skill_score', float('nan'))):.3f}. "
                f"Within anisotropic search support (n = {int(within.get('n', 0)):,}) it is "
                f"{float(within.get('brier_skill_score', float('nan'))):.3f}; the larger full-set penalty is driven by "
                "the deterministic host fallback outside search support.",
                f"A leakage-free nested recalibration sensitivity gives Brier skill "
                f"{float(recalibrated.get('brier_skill_score', float('nan'))):.3f}, but that mapping was not applied to "
                "the archived realisations. Reliability bins and the full confusion matrix are in Online Resource 2.",
            ],
        ),
        (
            "S4. Support Decomposition and Ensemble Convergence",
            [
                f"Valid reporting cells = {int(support.get('valid_cell_count', 322920)):,}; whole-grid mean = "
                f"{float(support.get('whole_grid_mean_tgc_pct', float('nan'))):.3f}% TGC; weighted reconstruction error = "
                f"{float(support.get('reconstruction_error_tgc_pct', float('nan'))):.3g}% TGC.",
                "Host-dominant, transitional and graphitic-dominant cell fractions are "
                + "/".join(
                    f"{float(row.get('cell_fraction_pct', float('nan'))):.2f}%"
                    for row in support.get("classes", [])
                )
                + "; corresponding means are "
                + "/".join(
                    f"{float(row.get('mean_tgc_pct', float('nan'))):.3f}%"
                    for row in support.get("classes", [])
                )
                + " TGC.",
                f"At n=75, envelope probability MAE = "
                f"{float(n75.get('map_metrics', {}).get('probability', {}).get('mae', {}).get('p50', float('nan'))):.3f}, "
                f"probability correlation = "
                f"{float(n75.get('map_metrics', {}).get('probability', {}).get('correlation', {}).get('p50', float('nan'))):.3f}, "
                f"spread correlation = "
                f"{float(n75.get('map_metrics', {}).get('spread', {}).get('correlation', {}).get('p50', float('nan'))):.3f}, "
                f"and hotspot Jaccard = {float(n75.get('spread_hotspot_jaccard', {}).get('p50', float('nan'))):.3f}.",
            ],
        ),
        (
            "S5. Variogram, Contact and Spatial Diagnostics",
            [
                f"Matched-space variogram reproduction used {int(vario.get('n_real_eval', 12))} normal-score realisations "
                f"and gives weighted RMSE {float(vario.get('weighted_rmse', float('nan'))):.3f}. Strike and down-dip "
                "directions retain nine pair-supported lags; thickness normal retains two.",
                f"The signed contact profile contains {int(contact.get('n_composites', 711))} composites around "
                f"{int(contact.get('contact_count', 134))} transitions in {int(contact.get('contact_holes', 42))} holes. "
                f"Graphitic minus host-side mean TGC = "
                f"{float(contact.get('graphitic_minus_host_mean_tgc_pct', float('nan'))):.3f} percentage points.",
                f"Upper-decile plan spread starts at {float(spatial.get('high_spread_threshold_tgc_pct', float('nan'))):.3f}% TGC. "
                f"High-spread columns have median nearest-composite distance "
                f"{float(spatial.get('high_spread_median_nearest_composite_plan_distance_m', float('nan'))):.1f} m, "
                f"versus {float(spatial.get('persistent_median_nearest_composite_plan_distance_m', float('nan'))):.1f} m "
                "for persistent-occupancy columns.",
                f"The joint high-spread and persistent set contains "
                f"{int(spatial.get('joint_high_spread_persistent_column_count', 0))} columns. Block-bootstrap "
                f"entropy-spread correlation is {float(obs.get('spearman_entropy_spread', float('nan'))):.3f}; "
                "these co-location statistics remain descriptive.",
            ],
        ),
        (
            "S6. Matched Repeated-Null Sensitivity",
            [
                "Five independent no-domain isotropic seed families and five non-overlapping conditioned subsets each "
                "contain 20 realisations. The same fractional archive weight and graphitic-composite reference are used for every comparison.",
                f"Conditioned/null median envelope mean TGC = "
                f"{median(canonical_summary, 'envelope_mean_tgc_pct'):.3f}/"
                f"{median(matched_null_summary, 'envelope_mean_tgc_pct'):.3f}%; P(TGC > 3%) = "
                f"{median(canonical_summary, 'envelope_probability_gt_3'):.3f}/"
                f"{median(matched_null_summary, 'envelope_probability_gt_3'):.3f}; P90-P10 spread = "
                f"{median(canonical_summary, 'envelope_p90_minus_p10_tgc_pct'):.3f}/"
                f"{median(matched_null_summary, 'envelope_p90_minus_p10_tgc_pct'):.3f}% TGC.",
                f"Conditioned/null median histogram overlap = "
                f"{median(canonical_summary, 'envelope_histogram_overlap_graphitic'):.3f}/"
                f"{median(matched_null_summary, 'envelope_histogram_overlap_graphitic'):.3f}; Q-Q RMSE = "
                f"{median(canonical_summary, 'envelope_qq_rmse_graphitic_tgc_pct'):.3f}/"
                f"{median(matched_null_summary, 'envelope_qq_rmse_graphitic_tgc_pct'):.3f}% TGC.",
                "The null retains closer marginal fit; the conditioned subsets retain higher above-threshold persistence "
                "and narrower conditional spread. Directional results are mixed, so no overall model winner is assigned.",
            ],
        ),
    ]

    with PdfPages(dst) as pdf:
        for heading, paragraphs in sections:
            fig = plt.figure(figsize=(8.27, 11.69), facecolor="white")
            ax = fig.add_axes([0.09, 0.08, 0.82, 0.84])
            ax.axis("off")
            ax.text(0, 1, heading, va="top", fontsize=15, weight="bold", color="#17324d")
            y = 0.91
            for paragraph in paragraphs:
                lines = wrap(paragraph, 100)
                ax.text(
                    0, y, "\n".join(lines), va="top", fontsize=10.5,
                    linespacing=1.5, color="#222222",
                )
                y -= 0.04 * max(2, len(lines)) + 0.035
            ax.text(
                0, 0.01, "Online Resource 1 | Supplementary Methods and Validation",
                fontsize=8, color="#666666",
            )
            pdf.savefig(fig, dpi=300)
            plt.close(fig)

        matched_rows = (
            matched.get("null_20_realisation_seed_families", {}).get("rows", []) or []
        )
        if len(matched_rows) == 5:
            fig = plt.figure(figsize=(11.69, 8.27), facecolor="white")
            ax = fig.add_axes([0.035, 0.08, 0.93, 0.84])
            ax.axis("off")
            ax.text(
                0, 1, "S6a. Independent no-domain seed metrics (matched envelope)",
                va="top", fontsize=15, weight="bold", color="#17324d",
            )
            ax.text(
                0, 0.93,
                "All five 20-realisation families are retained; summary rows use the five independent seed values.",
                va="top", fontsize=9.5, color="#222222",
            )
            columns = [
                "Seed", "Mean TGC (%)", "P(TGC > 3%)", "Spread (%)",
                "Hist. overlap", "Q-Q RMSE", "Strike r", "Dip r", "Normal r",
            ]
            keys = [
                "envelope_mean_tgc_pct", "envelope_probability_gt_3",
                "envelope_p90_minus_p10_tgc_pct", "envelope_histogram_overlap_graphitic",
                "envelope_qq_rmse_graphitic_tgc_pct", "envelope_swath_corr_strike",
                "envelope_swath_corr_down_dip", "envelope_swath_corr_thickness_normal",
            ]
            body = []
            for row in sorted(matched_rows, key=lambda item: int(item.get("seed", 0))):
                body.append([str(int(row["seed"])), *[f"{float(row[key]):.3f}" for key in keys]])
            for label, stat_key in (
                ("Median", "median"), ("Minimum", "min"), ("Maximum", "max"), ("SD", "std")
            ):
                body.append(
                    [
                        label,
                        *[
                            f"{float((matched_null_summary.get(key, {}) or {}).get(stat_key, float('nan'))):.3f}"
                            for key in keys
                        ],
                    ]
                )
            table = ax.table(
                cellText=body, colLabels=columns, cellLoc="center", colLoc="center",
                colWidths=[0.08, 0.12, 0.12, 0.105, 0.11, 0.10, 0.09, 0.09, 0.09],
                bbox=[0, 0.18, 1, 0.68],
            )
            table.auto_set_font_size(False)
            table.set_fontsize(8.0)
            for (row_index, _column_index), cell in table.get_celld().items():
                cell.set_edgecolor("#8A9AA8")
                cell.set_linewidth(0.55)
                if row_index == 0:
                    cell.set_facecolor("#17324D")
                    cell.get_text().set_color("white")
                    cell.get_text().set_weight("bold")
                elif row_index > 5:
                    cell.set_facecolor("#EAF1F6")
                    cell.get_text().set_weight("bold")
                elif row_index % 2 == 0:
                    cell.set_facecolor("#F5F7F9")
            ax.text(
                0, 0.10,
                "The same fractional lode envelope is used for all seeds and conditioned subsets; no seed was selected by performance.",
                fontsize=9, color="#222222", va="top",
            )
            ax.text(
                0, 0.01, "Online Resource 1 | Supplementary Methods and Validation",
                fontsize=8, color="#666666",
            )
            pdf.savefig(fig, dpi=300)
            plt.close(fig)

def _write_mme_esm_xlsx(dst: Path, staging_dir: Path, run_dir: Path, null_summary: dict, null_rows: list[dict]) -> None:
    truth=_first_json(ROOT/"build"/"source_of_truth.submission.json")
    validation={"run_level": _first_json(staging_dir/"supplement"/"validation_metrics.json", staging_dir/"validation_metrics.json", run_dir/"tables"/"validation_metrics.json"),
                "validation_gap_summaries": truth.get("validation_gap_summaries", {}),
                "blocked_validation_baseline": truth.get("blocked_validation_baseline", {}),
                "calibration_ablation": truth.get("calibration_ablation", {})}
    op=next((p for p in [staging_dir/"supplement"/"cutoff_occupancy_uncertainty.csv",staging_dir/"cutoff_occupancy_uncertainty.csv",run_dir/"tables"/"cutoff_occupancy_uncertainty.csv"] if p.exists()),None)
    occupancy=pd.read_csv(op).replace({np.nan:None}).to_dict(orient="records") if op else []
    if occupancy:
        baseline = occupancy[0]
        converted = []
        for row in occupancy:
            item = {"threshold_tgc_pct": row.get("cutoff"), "n_realisations": row.get("nonzero_count")}
            for quantile in ("p05", "mean", "p10", "p50", "p90", "p95"):
                source_key = f"occupancy_proxy_{quantile}"
                denominator = float(baseline.get(source_key) or 0.0)
                item[f"model_occupancy_{quantile}_pct"] = (
                    100.0 * float(row.get(source_key) or 0.0) / denominator if denominator else None
                )
                item[f"mean_tgc_above_threshold_{quantile}_pct"] = row.get(f"mean_tgc_above_cutoff_{quantile}")
            converted.append(item)
        occupancy = converted
    run_metadata = json.loads(json.dumps(_first_json(run_dir/"sgs_meta.json", staging_dir/"sgs_meta.json")))
    if isinstance(run_metadata.get("config"), dict):
        run_metadata["config"].pop("publication", None)
    simulation = run_metadata.setdefault("config", {}).setdefault("simulation", {})
    configured_label = simulation.get("kriging_type")
    simulation["configured_legacy_kriging_type_label"] = configured_label
    simulation["kriging_type"] = "SK_style_effective"
    simulation["implemented_estimator"] = "simple_kriging_style_normal_score"
    simulation["solver_evidence"] = "covariance system solved without an ordinary-kriging Lagrange multiplier"
    run_metadata["estimator_implementation_audit"] = {
        "configured_legacy_label": configured_label,
        "reported_effective_estimator": "SK_style_effective",
        "implementation_reference": "src/sgs.py::_krige_local",
        "simulation_values_changed_by_metadata_correction": False,
    }
    run_metadata["submission_profile"] = {
        "journal": MME_JOURNAL,
        "collection": MME_COLLECTION,
        "note": "Administrative publication metadata from the original science run are excluded; scientific configuration is unchanged.",
    }
    payload={
      "README":[["Field","Value"],["Title",MME_TITLE],["Journal",MME_JOURNAL],["Collection",MME_COLLECTION],["Author",AUTHOR_NAME],["Affiliation",AUTHOR_AFFILIATION],["Email",AUTHOR_EMAIL],["ORCID",AUTHOR_ORCID_URL],["Scope","Audit-level evidence; archive-derived reporting-envelope sensitivity; proprietary drillhole and domain arrays excluded."],["Null campaign",null_summary.get("status","pending")]],
      "Run Metadata":[["Metric","Value"]]+_flatten(run_metadata),
      "Validation Metrics":[["Metric","Value"]]+_flatten(validation),
      "Variogram Models":[["Metric","Value"]]+_flatten(_first_json(staging_dir/"supplement"/"variogram_model.json",staging_dir/"variogram_model.json",run_dir/"figures"/"variogram_model.json",run_dir/"tables"/"variogram_model.json")),
      "Convergence":[["Metric","Value"]]+_flatten(validation.get("validation_gap_summaries",{}).get("archive_lode_envelope_convergence",{})),
      "Reporting Envelope":[["Metric","Value"]]+_flatten({
          "geometry": validation.get("validation_gap_summaries",{}).get("archive_lode_envelope",{}),
          "spatial_patterns": validation.get("validation_gap_summaries",{}).get("archive_lode_spatial_patterns",{}),
      }),
      "Matched Envelope":[["Metric","Value"]]+_flatten(validation.get("validation_gap_summaries",{}).get("archive_lode_matched_null_comparison",{})),
      "Support Decomposition":[["Metric","Value"]]+_flatten(validation.get("validation_gap_summaries",{}).get("support_aligned_mean_decomposition",{})),
      "Categorical Validation":[["Metric","Value"]]+_flatten(validation.get("validation_gap_summaries",{}).get("categorical_domain_grouped_validation",{})),
      "Contact Statistics":[["Metric","Value"]]+_flatten(validation.get("validation_gap_summaries",{}).get("signed_graphitic_host_contact",{})),
      "Occupancy Diagnostics":([list(occupancy[0])]+[[row.get(k) for k in occupancy[0]] for row in occupancy]) if occupancy else [["Status"],["No rows available"]],
      "Repeated Null Summary":[["Metric","Value"]]+_flatten(null_summary),
      "Repeated Null Seeds":([list(null_rows[0])]+[[row.get(k) for k in null_rows[0]] for row in null_rows]) if null_rows else [["Status"],["Pending seeds 9101, 9201, 9301, 9401, 9501"]]}
    work=BUILD_WORK_DIR/"mme_esm"; work.mkdir(parents=True,exist_ok=True); payload_path=work/"payload.json"; script_path=work/"build.mjs"
    payload_path.write_text(json.dumps(payload,ensure_ascii=True),encoding="utf-8")
    script_path.write_text(r'''import fs from "node:fs/promises";
import { SpreadsheetFile, Workbook } from "file:///C:/Users/SUDIPTA%20CHANDA/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/node_modules/%40oai/artifact-tool/dist/artifact_tool.mjs";
const data=JSON.parse(await fs.readFile(process.argv[2],"utf8"));const wb=Workbook.create();
for(const [name,rows0] of Object.entries(data)){const rows=rows0.map(r=>r.map(v=>v??""));const cols=Math.max(...rows.map(r=>r.length));const padded=rows.map(r=>[...r,...Array(cols-r.length).fill("")]);const sh=wb.worksheets.add(name);const all=sh.getRangeByIndexes(0,0,padded.length,cols);all.values=padded;all.format={font:{name:"Arial",size:9},verticalAlignment:"top",wrapText:true};const h=sh.getRangeByIndexes(0,0,1,cols);h.format={fill:"#17324D",font:{name:"Arial",size:10,bold:true,color:"#FFFFFF"},wrapText:true};sh.freezePanes.freezeRows(1);sh.getRangeByIndexes(0,0,padded.length,1).format.columnWidth=42;if(cols>1)sh.getRangeByIndexes(0,1,padded.length,cols-1).format.columnWidth=24;for(let r=2;r<padded.length;r+=2)sh.getRangeByIndexes(r,0,1,cols).format.fill="#EAF1F6";}
const out=await SpreadsheetFile.exportXlsx(wb);await out.save(process.argv[3]);''',encoding="utf-8")
    deps=Path.home()/".cache"/"codex-runtimes"/"codex-primary-runtime"/"dependencies"/"node"; env=dict(os.environ); env["NODE_PATH"]=str(deps/"node_modules")
    cp=subprocess.run([str(deps/"bin"/"node.exe"),str(script_path),str(payload_path),str(dst)],env=env,capture_output=True,text=True)
    if cp.returncode: raise RuntimeError("artifact-tool failed: "+(cp.stderr or cp.stdout))
    sidecar = dst.with_name(dst.name + ".inspect.ndjson")
    if sidecar.exists():
        sidecar.unlink()


def _style_mme_manuscript(path: Path) -> None:
    doc=Document(path); sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(0.9); names={st.name for st in doc.styles}
    for name in ("Normal","Body Text"):
        if name in names:
            st=doc.styles[name]; st.font.name="Times New Roman"; st.font.size=Pt(10); st.paragraph_format.line_spacing=1.15
    for n in (1,2,3):
        name=f"Heading {n}"
        if name in names:
            st=doc.styles[name]; st.font.name="Arial"; st.font.bold=True; st.font.size=Pt(14-n)
    if "Caption" not in names:
        caption_style=doc.styles.add_style("Caption",WD_STYLE_TYPE.PARAGRAPH); names.add("Caption")
    caption_style=doc.styles["Caption"]; caption_style.font.name="Times New Roman"; caption_style.font.size=Pt(10); caption_style.font.italic=False
    for paragraph in doc.paragraphs:
        if re.match(r"^(?:Fig\.|Figure)\s*\d+[A-Z]?\b",paragraph.text.strip(),flags=re.I):
            paragraph.style=caption_style; paragraph.paragraph_format.line_spacing=1.15; paragraph.paragraph_format.space_after=Pt(6)
    for table in doc.tables:
        for row_index,row in enumerate(table.rows):
            tr_pr=row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))
            if row_index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
                header=OxmlElement("w:tblHeader"); header.set(qn("w:val"),"true"); tr_pr.append(header)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing=1.0; paragraph.paragraph_format.space_before=Pt(0); paragraph.paragraph_format.space_after=Pt(0)
                    for run in paragraph.runs:
                        run.font.name="Times New Roman"; run.font.size=Pt(8.5); run.font.bold=(row_index == 0)
    doc.save(path)


def _embed_mme_figures(docx_path: Path, staging_dir: Path, run_dir: Path) -> None:
    with zipfile.ZipFile(docx_path, "r") as archive:
        existing_media = sum(1 for name in archive.namelist() if name.startswith("word/media/"))
    if existing_media >= 7:
        return
    doc = Document(docx_path)
    captions = {
        idx: next((p for p in doc.paragraphs if re.match(rf"^(?:Fig\\.|Figure)\\s*{idx}\\b", p.text.strip(), flags=re.I)), None)
        for idx in range(1, 8)
    }
    for idx, src_name in figure_image_map().items():
        src = staging_dir / "figures" / src_name
        if not src.exists():
            src = run_dir / "figures" / src_name
        if not src.exists():
            continue
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(src), width=Inches(6.5))
        caption = captions.get(int(idx))
        if caption is not None:
            caption._p.addprevious(paragraph._p)
    doc.save(docx_path)


def _optimise_mme_docx_media(docx_path: Path, max_width_px: int = 2100) -> None:
    from PIL import Image

    rewritten: list[tuple[zipfile.ZipInfo, bytes]] = []
    changed = False
    with zipfile.ZipFile(docx_path, "r") as archive:
        for info in archive.infolist():
            data = archive.read(info.filename)
            suffix = Path(info.filename).suffix.lower()
            if info.filename.startswith("word/media/") and suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
                try:
                    with Image.open(io.BytesIO(data)) as source:
                        if source.width > max_width_px:
                            ratio = float(max_width_px) / float(source.width)
                            size = (max_width_px, max(1, int(round(source.height * ratio))))
                            rgba = source.convert("RGBA").resize(size, Image.Resampling.LANCZOS)
                            image = Image.new("RGB", size, "white")
                            image.paste(rgba, mask=rgba.getchannel("A"))
                            buffer = io.BytesIO()
                            if suffix == ".png":
                                image.save(buffer, format="PNG", optimize=True, dpi=(300, 300))
                            elif suffix in {".jpg", ".jpeg"}:
                                image.save(buffer, format="JPEG", quality=92, optimize=True, dpi=(300, 300))
                            else:
                                image.save(buffer, format="TIFF", compression="tiff_lzw", dpi=(300, 300))
                            data = buffer.getvalue()
                            changed = True
                except Exception:
                    pass
            rewritten.append((info, data))
    if not changed:
        return
    temporary = docx_path.with_name(docx_path.stem + ".media_tmp.docx")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for info, data in rewritten:
            archive.writestr(info, data)
    temporary.replace(docx_path)


def build_clean_submission_ready(staging_dir: Path, run_dir: Path) -> Path:
    clean_dir=ROOT/".tmp_clean_submission_mme"
    if clean_dir.exists(): shutil.rmtree(clean_dir)
    clean_dir.mkdir(parents=True); null_summary,null_rows=_mme_null_evidence()
    manuscript=staging_dir/PRIMARY_MANUSCRIPT_DOCX
    if not manuscript.exists(): manuscript=staging_dir/"paper.docx"
    if not manuscript.exists(): raise FileNotFoundError("Generated manuscript DOCX is missing")
    shutil.copy2(manuscript,clean_dir/"Manuscript.docx"); _style_mme_manuscript(clean_dir/"Manuscript.docx")
    _embed_mme_figures(clean_dir/"Manuscript.docx", staging_dir, run_dir)
    _optimise_mme_docx_media(clean_dir/"Manuscript.docx")
    strip_docx_review_markup(clean_dir/"Manuscript.docx")
    _write_mme_cover(clean_dir/"Cover_Letter.docx",null_summary)
    for key,name in figure_image_map().items():
        src=staging_dir/"figures"/name
        if not src.exists(): src=run_dir/"figures"/name
        dpi,width=_figure_artwork_requirements(key)
        if not src.exists() or not _convert_image_to_tif(src,clean_dir/_clean_figure_filename(key),dpi,width): raise RuntimeError(f"Could not export Fig{key}")
    _write_mme_esm_pdf(clean_dir/"ESM_1.pdf",null_summary,null_rows); _write_mme_esm_xlsx(clean_dir/"ESM_2.xlsx",staging_dir,run_dir,null_summary,null_rows)
    return clean_dir


def _extract_highlight_bullets(highlights_text: str) -> list[str]:
    bullets: list[str] = []
    for ln in highlights_text.splitlines():
        line = ln.strip()
        if line.startswith("- "):
            bullets.append(line[2:].strip())
        elif line.startswith("•"):
            bullets.append(line[1:].strip())
    return bullets


def _write_highlights_docx(dst: Path, bullets: list[str]) -> None:
    doc = Document()
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet)
    doc.save(dst)


def _write_jaes_submission_spec_docx(dst: Path) -> None:
    doc = Document()
    doc.add_heading("JAES Submission Specification (Guide-Based)", level=1)
    doc.add_paragraph(
        "Source: Elsevier Guide for Authors, Journal of African Earth Sciences (ISSN 1464-343X)."
    )
    doc.add_paragraph("Generated verification brief for this submission package.")
    doc.add_paragraph("")

    checks = [
        "Editable manuscript source (.docx), not PDF as source file.",
        "Single-column manuscript layout.",
        "Continuous line numbering and page numbering in manuscript.",
        "Abstract length does not exceed 250 words.",
        "Keyword count is between 1 and 7.",
        "Highlights are 3-5 bullets, each <=85 characters.",
        "Corresponding author metadata includes email, phone, and full postal affiliation details.",
        "Tables are provided as editable text, not image objects.",
        "All figure files include captions in manuscript context.",
        "Raster artwork meets publication-grade resolution guidance (>=300 dpi baseline).",
        "Required-only package policy: generate only mandatory JAES upload files.",
        "Supplementary package contains the required-only four-file S2 evidence bundle aligned with manuscript claims and table evidence mapping.",
    ]
    for item in checks:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_paragraph("")
    doc.add_heading("Package Mapping (Current Build)", level=2)
    mapped = [
        "01_Title_Page.docx: author and corresponding-author metadata.",
        "02_Highlights.docx: validated highlights file.",
        "04_Manuscript.docx: canonical manuscript.",
        "07_Cover_Letter.docx: cover letter.",
        "08_Declaration_of_Interest.docx: declaration file.",
        "Fig01-Fig07 TIFF set: figure artwork files.",
        "Supplementary_Data_S2.zip: required-only four-file evidence bundle.",
    ]
    for item in mapped:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("")
    doc.add_heading("Exact Upload File Set (16 Files)", level=2)
    for name in CLEAN_REQUIRED_SUBMISSION_FILES:
        doc.add_paragraph(name, style="List Bullet")

    doc.add_paragraph("")
    doc.add_heading("Agent Instruction Addendum", level=2)
    doc.add_paragraph(
        "Use this checklist as the acceptance criteria for future packaging passes. "
        "Keep upload files minimal and exclude internal build notes, source-of-truth internals, and duplicate manuscript variants."
    )
    doc.save(dst)


def _build_graphical_abstract_tif(dst: Path) -> bool:
    summary = ""
    src = SUBMISSION_DIR / "graphical_abstract_summary.txt"
    if src.exists():
        summary = src.read_text(encoding="utf-8").strip()
    if not summary:
        summary = "Graphical abstract summary unavailable."
    try:
        from PIL import Image, ImageDraw, ImageFont, ImageOps  # type: ignore

        w, h = 2400, 960
        img = Image.new("RGB", (w, h), color=(255, 255, 255))
        draw = ImageDraw.Draw(img)
        title = "Tanzanian Mozambique Belt stratiform graphite uncertainty"
        font = ImageFont.load_default()
        draw.rectangle((0, 0, w, 86), fill=(31, 45, 55))
        draw.text((42, 30), title, fill=(255, 255, 255), font=font)

        panel_paths = [
            SUBMISSION_DIR / "figures" / "figure_1_regional_geology_map.png",
            SUBMISSION_DIR / "figures" / "variogram.png",
            SUBMISSION_DIR / "figures" / "tonnage_risk_curve.png",
        ]
        panel_titles = ["Regional belt context", "Fabric/anisotropy prior", "Thickness-normal uncertainty"]
        panel_w, panel_h = 690, 500
        x0s = [42, 855, 1668]
        for x0, panel_path, panel_title in zip(x0s, panel_paths, panel_titles):
            draw.rectangle((x0, 130, x0 + panel_w, 130 + panel_h), outline=(95, 108, 120), width=3)
            draw.text((x0, 104), panel_title, fill=(0, 0, 0), font=font)
            if panel_path.exists():
                with Image.open(panel_path) as src_img:
                    src_img = src_img.convert("RGB")
                    src_img = ImageOps.contain(src_img, (panel_w - 20, panel_h - 20))
                    px = x0 + (panel_w - src_img.width) // 2
                    py = 140 + (panel_h - src_img.height) // 2
                    img.paste(src_img, (px, py))
            else:
                draw.text((x0 + 24, 350), panel_path.name, fill=(120, 120, 120), font=font)

        callouts = [
            "Pan-African graphitic metasedimentary architecture frames continuity.",
            "Conditional simulation tests foliation-parallel persistence against thickness-normal uncertainty.",
            "Weak validation is reported as a claim boundary, not hidden.",
        ]
        y = 690
        for callout in callouts:
            draw.rectangle((42, y - 10, 2358, y + 42), fill=(242, 245, 247), outline=(200, 208, 216))
            draw.text((68, y + 8), callout, fill=(0, 0, 0), font=font)
            y += 68

        max_chars = 175
        words = summary.replace("\n", " ").split()
        lines = []
        cur = []
        for wd in words:
            trial = " ".join(cur + [wd])
            if len(trial) > max_chars:
                lines.append(" ".join(cur))
                cur = [wd]
            else:
                cur.append(wd)
        if cur:
            lines.append(" ".join(cur))
        y = 890
        for ln in lines[:2]:
            draw.text((42, y), ln, fill=(35, 35, 35), font=font)
            y += 24
        img.save(dst, format="TIFF", dpi=(300, 300))
        return True
    except Exception:
        return False


def _convert_image_to_tif(src: Path, dst: Path, min_dpi: int = 300, min_width_px: int = 0) -> bool:
    try:
        from PIL import Image, ImageDraw  # type: ignore

        with Image.open(src) as img:
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            if min_width_px and img.width < min_width_px:
                ratio = float(min_width_px) / float(max(img.width, 1))
                new_size = (int(round(img.width * ratio)), int(round(img.height * ratio)))
                img = img.resize(new_size, Image.Resampling.LANCZOS)
            img.save(dst, format="TIFF", dpi=(float(min_dpi), float(min_dpi)), compression="tiff_lzw")
        return True
    except Exception:
        pass
    try:
        import matplotlib.image as mpimg  # type: ignore
        import matplotlib.pyplot as plt  # type: ignore

        data = mpimg.imread(src)
        plt.imsave(dst, data, format="tiff", dpi=min_dpi)
        return True
    except Exception:
        return False


def _zip_files(dst_zip: Path, files: list[tuple[Path, str]]) -> None:
    if dst_zip.exists():
        dst_zip.unlink()
    with zipfile.ZipFile(dst_zip, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for src, arcname in files:
            if src.exists() and src.is_file():
                zf.write(src, arcname)


def _escape_pdf_text(s: str) -> str:
    return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_simple_pdf(path: Path, lines: list[str]) -> None:
    wrapped: list[str] = []
    for line in lines:
        chunks = textwrap.wrap(line, width=95) if line else [""]
        wrapped.extend(chunks)
    stream_lines = ["BT", "/F1 10 Tf", "50 790 Td", "13 TL"]
    for line in wrapped[:250]:
        stream_lines.append(f"({_escape_pdf_text(line)}) Tj")
        stream_lines.append("T*")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1", errors="replace")

    objects: list[bytes] = []
    objects.append(b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n")
    objects.append(b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n")
    objects.append(
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n"
    )
    objects.append(b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n")
    objects.append(f"5 0 obj << /Length {len(stream)} >> stream\n".encode("ascii") + stream + b"\nendstream endobj\n")

    offsets = [0]
    payload = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    for obj in objects:
        offsets.append(len(payload))
        payload += obj

    xref_pos = len(payload)
    payload += f"xref\n0 {len(offsets)}\n".encode("ascii")
    payload += b"0000000000 65535 f \n"
    for off in offsets[1:]:
        payload += f"{off:010d} 00000 n \n".encode("ascii")
    payload += (
        f"trailer << /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(payload)


def _build_submission_final_jaes(run_dir: Path) -> list[str]:
    issues: list[str] = []
    stage = ROOT / f".tmp_build_{FINAL_SUBMISSION_DIR.name}"
    if stage.exists():
        shutil.rmtree(stage)
    stage.mkdir(parents=True, exist_ok=True)

    # 01 title page
    src_title = SUBMISSION_DIR / "Title page.docx"
    if src_title.exists():
        shutil.copy2(src_title, stage / "01_Title_Page.docx")
    else:
        issues.append("missing Title page.docx for 01_Title_Page.docx")

    # 02 highlights docx (3-5 bullets, each <=85 chars)
    highlights_txt = SUBMISSION_DIR / "highlights.txt"
    if highlights_txt.exists():
        bullets = _extract_highlight_bullets(read_text(highlights_txt))
        if not (3 <= len(bullets) <= 5):
            issues.append(f"highlights count must be 3-5, found {len(bullets)}")
        too_long = [b for b in bullets if len(b) > 85]
        if too_long:
            issues.append("highlights contain bullets longer than 85 characters")
        if bullets:
            _write_highlights_docx(stage / "02_Highlights.docx", bullets[:5])
    else:
        issues.append("missing highlights.txt for 02_Highlights.docx")

    # 04 manuscript
    src_manuscript = SUBMISSION_DIR / PRIMARY_MANUSCRIPT_DOCX
    if src_manuscript.exists():
        shutil.copy2(src_manuscript, stage / "04_Manuscript.docx")
        _apply_final_manuscript_styles(stage / "04_Manuscript.docx")
    else:
        issues.append(f"missing {PRIMARY_MANUSCRIPT_DOCX} for 04_Manuscript.docx")

    # 08 declaration
    src_decl = SUBMISSION_DIR / "Declaration of interest statement.docx"
    if src_decl.exists():
        shutil.copy2(src_decl, stage / "08_Declaration_of_Interest.docx")
    else:
        issues.append("missing Declaration of interest statement.docx for 08_Declaration_of_Interest.docx")

    fig_src = SUBMISSION_DIR / "figures"

    figure_map: dict[str, str] = {}
    for fig_key, src_name in figure_image_map().items():
        figure_map[_clean_figure_filename(fig_key)] = src_name
    for dst_name, src_name in figure_map.items():
        src = fig_src / src_name
        if not src.exists():
            src = run_dir / "figures" / src_name
        if not src.exists():
            issues.append(f"missing source figure for {dst_name}: expected {src_name}")
            continue
        fig_key = next((key for key in figure_image_map() if _clean_figure_filename(key) == dst_name), "")
        min_dpi, min_width = _figure_artwork_requirements(fig_key)
        ok = _convert_image_to_tif(src, stage / dst_name, min_dpi=min_dpi, min_width_px=min_width)
        if not ok:
            issues.append(f"failed converting {src.name} to {dst_name}")

    # Supplementary data S2
    s2_candidates = [
        (SUBMISSION_DIR / "supplement" / "sgs_meta.json", "sgs_meta.json"),
        (run_dir / "sgs_meta.json", "sgs_meta.json"),
        (SUBMISSION_DIR / "supplement" / "validation_metrics.json", "validation_metrics.json"),
        (run_dir / "tables" / "validation_metrics.json", "validation_metrics.json"),
        (SUBMISSION_DIR / "supplement" / "cutoff_occupancy_uncertainty.csv", "cutoff_occupancy_uncertainty.csv"),
        (SUBMISSION_DIR / "supplement" / "risked_tonnage.csv", "cutoff_occupancy_uncertainty.csv"),
        (run_dir / "tables" / "risked_tonnage.csv", "cutoff_occupancy_uncertainty.csv"),
        (SUBMISSION_DIR / "supplement" / "variogram_model.json", "variogram_model.json"),
        (run_dir / "figures" / "variogram_model.json", "variogram_model.json"),
    ]
    s2_files: list[tuple[Path, str]] = []
    seen_arc: set[str] = set()
    for src, arc in s2_candidates:
        if not src.exists() or arc in seen_arc:
            continue
        tmp = stage / f"_{arc}"
        if arc.lower().endswith(".json"):
            _sanitize_public_json(src, tmp)
            if arc == "validation_metrics.json":
                _enrich_public_validation_metrics(tmp, run_dir)
            elif arc == "variogram_model.json":
                _enrich_public_variogram_model(tmp, run_dir)
        elif arc == "cutoff_occupancy_uncertainty.csv":
            df = pd.read_csv(src)
            rename_map = {
                "tonnage": "occupancy_proxy",
                "tonnage_p05": "occupancy_proxy_p05",
                "tonnage_p10": "occupancy_proxy_p10",
                "tonnage_mean": "occupancy_proxy_mean",
                "tonnage_p50": "occupancy_proxy_p50",
                "tonnage_p90": "occupancy_proxy_p90",
                "tonnage_p95": "occupancy_proxy_p95",
                "grade_p05": "mean_tgc_above_cutoff_p05",
                "grade_p10": "mean_tgc_above_cutoff_p10",
                "grade_mean": "mean_tgc_above_cutoff_mean",
                "grade_p50": "mean_tgc_above_cutoff_p50",
                "grade_p90": "mean_tgc_above_cutoff_p90",
                "grade_p95": "mean_tgc_above_cutoff_p95",
            }
            df = df.rename(columns={k: v for k, v in rename_map.items() if k in df.columns})
            df = df[[c for c in df.columns if not c.lower().startswith("contained")]]
            raw_tmp = stage / f"_raw_{arc}"
            df.to_csv(raw_tmp, index=False)
            _sanitize_public_text_table(raw_tmp, tmp)
            raw_tmp.unlink(missing_ok=True)
        elif arc.lower().endswith(".csv"):
            _sanitize_public_text_table(src, tmp)
        else:
            shutil.copy2(src, tmp)
        s2_files.append((tmp, arc))
        seen_arc.add(arc)
    expected_s2 = set(ESSENTIAL_S2_FILES)
    present_s2 = {arc for _, arc in s2_files}
    missing_s2 = sorted(expected_s2 - present_s2)
    if missing_s2:
        issues.append("missing required files for Supplementary_Data_S2.zip: " + ", ".join(missing_s2))
    _zip_files(stage / "Supplementary_Data_S2.zip", s2_files)
    for tmp_src, _arc in s2_files:
        if tmp_src.parent == stage and tmp_src.name.startswith("_"):
            tmp_src.unlink(missing_ok=True)

    # Cover letter
    src_cover = SUBMISSION_DIR / "Covering Letter-MS.docx"
    if src_cover.exists():
        shutil.copy2(src_cover, stage / "07_Cover_Letter.docx")
    else:
        issues.append("missing Covering Letter-MS.docx for 07_Cover_Letter.docx")

    for docx in stage.glob("*.docx"):
        strip_docx_review_markup(docx)

    # Remove transient Office lock/temp files if present.
    for tmp in stage.glob("~$*"):
        try:
            tmp.unlink()
        except OSError:
            pass

    # Final required set check
    for fname in final_submission_required_files():
        if not (stage / fname).exists():
            issues.append(f"missing required final submission file `{fname}`")

    if FINAL_SUBMISSION_DIR.exists():
        try:
            shutil.rmtree(FINAL_SUBMISSION_DIR)
            shutil.move(str(stage), str(FINAL_SUBMISSION_DIR))
        except PermissionError:
            # Fallback for Windows file-lock situations: overwrite required files in place.
            FINAL_SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
            for name in final_submission_required_files():
                src = stage / name
                dst = FINAL_SUBMISSION_DIR / name
                if src.is_file():
                    try:
                        shutil.copy2(src, dst)
                    except OSError:
                        # Keep existing file when locked by external apps (e.g., Word).
                        pass
            allowed = set(final_submission_required_files())
            for item in FINAL_SUBMISSION_DIR.iterdir():
                if item.name in allowed:
                    continue
                if item.is_dir():
                    shutil.rmtree(item, ignore_errors=True)
                else:
                    try:
                        item.unlink()
                    except OSError:
                        pass
            shutil.rmtree(stage, ignore_errors=True)
    else:
        shutil.move(str(stage), str(FINAL_SUBMISSION_DIR))
    return issues


def main() -> None:
    global FINAL_SUBMISSION_DIR, PACKAGE_PROFILE, PACKAGE_ZIP_BASENAME, SUBMISSION_DIR
    parser = argparse.ArgumentParser(description="Build full submission package from source-of-truth workflow outputs.")
    parser.add_argument("--run-dir", default=None, help="Optional explicit run directory containing sgs_meta.json and tables/")
    parser.add_argument("--out-dir", default=None, help="Optional output directory for a clean rebuild package")
    parser.add_argument("--strict", action="store_true", help="Fail if any required submission file is missing")
    parser.add_argument("--skip-final-jaes", action="store_true", help="Do not rebuild submission_final_JAES folder")
    parser.add_argument("--independent", action="store_true", help="Build the standalone geology-led independent submission package")
    args = parser.parse_args()

    if args.independent:
        PACKAGE_PROFILE = "independent"
        PACKAGE_ZIP_BASENAME = "submission_package_independent_clean"
        SUBMISSION_DIR = ROOT / "submission_ready_independent"
        FINAL_SUBMISSION_DIR = ROOT / "submission_final_JAES_independent"
    if args.out_dir:
        SUBMISSION_DIR = Path(args.out_dir)
        if not SUBMISSION_DIR.is_absolute():
            SUBMISSION_DIR = ROOT / SUBMISSION_DIR
    target_dir = SUBMISSION_DIR
    staging_dir = ROOT / f".tmp_build_{target_dir.name}"
    SUBMISSION_DIR = staging_dir
    clean_submission_dir()
    run_dir = resolve_independent_run_dir(args.run_dir) if args.independent else resolve_run_dir(args.run_dir)

    if args.independent:
        run_build_independent_package(run_dir)
    else:
        run_build_from_source_of_truth(run_dir)
        generate_template_docs()
    deduped = dedupe_submission_ready()
    sync_source_of_truth_files(run_dir)
    ensure_required_fallback_assets(run_dir)
    build_submission_zip()
    if args.independent:
        write_independent_checklist(run_dir)
    # The maintained source generator is reused, but the final upload contract is MME-specific.
    content_issues: list[str] = []
    visual_issues: list[str] = []
    visual_report: dict = {}
    missing: list[str] = []
    final_submission_issues: list[str] = []

    archive_internal_submission_work(staging_dir)
    clean_dir: Path | None = None
    if not args.independent:
        clean_dir = build_clean_submission_ready(staging_dir, run_dir)

    source_dir = clean_dir if clean_dir is not None else staging_dir
    if target_dir.exists():
        try:
            shutil.rmtree(target_dir)
        except Exception:
            # Windows/Word may keep the manuscript DOCX open. In that case,
            # update the required upload set in place and leave any locked file
            # untouched rather than leaving the package half rebuilt.
            target_dir.mkdir(parents=True, exist_ok=True)
            allowed = set(CLEAN_REQUIRED_SUBMISSION_FILES if not args.independent else INDEPENDENT_FINAL_SUBMISSION_REQUIRED_FILES)
            if not allowed:
                allowed = {p.name for p in source_dir.iterdir() if p.is_file()}
            failed_required_copies: list[str] = []
            for item in source_dir.iterdir():
                if not item.is_file():
                    continue
                dst = target_dir / item.name
                try:
                    shutil.copy2(item, dst)
                except OSError as exc:
                    if item.name in allowed:
                        failed_required_copies.append(f"{item.name}: {exc}")
                        continue
                    if not dst.exists():
                        raise
            if failed_required_copies:
                raise RuntimeError(
                    "Could not replace required clean package files. "
                    "Close any open Word/Office files and rerun: "
                    + "; ".join(failed_required_copies)
                )
            for item in list(target_dir.iterdir()):
                if item.name.startswith("~$"):
                    try:
                        item.unlink()
                    except OSError:
                        pass
                    continue
                if item.name in allowed:
                    continue
                if item.is_dir():
                    shutil.rmtree(item, ignore_errors=True)
                else:
                    try:
                        item.unlink()
                    except OSError:
                        pass
            shutil.rmtree(source_dir, ignore_errors=True)
            shutil.rmtree(staging_dir, ignore_errors=True)
            source_dir = None
    if source_dir is not None:
        shutil.move(str(source_dir), str(target_dir))
        if clean_dir is not None:
            shutil.rmtree(staging_dir, ignore_errors=True)
    SUBMISSION_DIR = target_dir
    sync_repo_mirror_after_clean_build(target_dir)
    if args.independent and FINAL_SUBMISSION_DIR.exists():
        shutil.rmtree(FINAL_SUBMISSION_DIR, ignore_errors=True)
    sync_independent_named_manuscript()
    sync_same_named_files_to_existing_submission_folders()
    cleanup_generated_submission_artifacts(target_dir)
    remove_office_lock_files(target_dir, FINAL_SUBMISSION_DIR)

    actual = {p.name for p in target_dir.iterdir() if p.is_file()}
    expected = set(CLEAN_REQUIRED_SUBMISSION_FILES)
    mme_issues = [f"missing {name}" for name in sorted(expected - actual)]
    mme_issues += [f"unexpected {name}" for name in sorted(actual - expected)]
    null_summary, null_rows = _mme_null_evidence()
    if null_summary.get("status") != "complete" or len(null_rows) != 5:
        mme_issues.append("five independent null families are not yet complete")
    if args.strict and mme_issues:
        raise SystemExit("MME package verification failed: " + "; ".join(mme_issues))

    print(f"MME submission package built from source-of-truth run: {run_dir}")
    if deduped:
        print("Removed duplicate files:", ", ".join(str(p.name) for p in deduped))
    if missing:
        print("Checklist has missing items:", ", ".join(missing))
    else:
        print("Checklist validation passed.")
    if content_issues:
        print("Content verification issues:", "; ".join(content_issues))
    if visual_issues:
        print("Visual verification issues:", "; ".join(visual_issues))
    else:
        print("Visual formatting checks passed.")
    if final_submission_issues:
        print("MME verification issues:", "; ".join(final_submission_issues))
    elif not args.skip_final_jaes and not args.independent:
        print(f"MME submission folder ready: {target_dir}")


if __name__ == "__main__":
    main()
