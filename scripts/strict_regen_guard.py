from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
import xml.etree.ElementTree as ET

import yaml
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SUBMISSION_DIR = ROOT / "submission_final_JAES"
DEFAULT_MANIFEST = ROOT / "review" / "STRICT_BASELINE_MANIFEST.json"
DEFAULT_POLICY = ROOT / "review" / "STRICT_ALLOWED_DELTAS.yaml"

TRACKED_DOCX = [
    "04_Manuscript.docx",
    "Tables_with_Captions.docx",
    "Figures_with_Captions.docx",
]
TRACKED_ZIP = "Supplementary_Data_S2.zip"


def abs_path(p: Path) -> Path:
    return p if p.is_absolute() else (ROOT / p)


def sha256_bytes(data: bytes) -> str:
    h = hashlib.sha256()
    h.update(data)
    return h.hexdigest()


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def normalize_line(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def docx_lines(path: Path) -> list[str]:
    doc = Document(path)
    out = [normalize_line(p.text) for p in doc.paragraphs if normalize_line(p.text)]
    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_line(c.text) for c in row.cells]
            row_text = " | ".join([c for c in cells if c])
            if row_text:
                out.append(row_text)
    return out


def _line_hash(lines: list[str]) -> str:
    return sha256_bytes("\n".join(lines).encode("utf-8"))


def docx_invariants(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path, "r") as zf:
        doc_xml = zf.read("word/document.xml")
    root = ET.fromstring(doc_xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    ln_num = bool(root.findall(".//w:lnNumType", ns))
    sects = root.findall(".//w:sectPr", ns)
    single_col = True
    for sect in sects:
        cols = sect.find("w:cols", ns)
        if cols is None:
            continue
        val = cols.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num")
        if val and int(val) > 1:
            single_col = False
            break
    doc = Document(path)
    return {
        "line_numbering_present": ln_num,
        "single_column_layout": single_col,
        "section_count": len(sects),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
    }


def heading_positions(lines: list[str]) -> dict[str, int]:
    keys = {
        "abstract": r"^abstract\b",
        "keywords": r"^keywords\b",
        "data_availability": r"^data availability\b",
        "references": r"^references\b",
        "tables": r"^tables\b",
        "figure_captions": r"^figure captions\b",
    }
    out: dict[str, int] = {}
    for k, pat in keys.items():
        out[k] = -1
        for i, ln in enumerate(lines):
            if re.search(pat, ln, flags=re.IGNORECASE):
                out[k] = i
                break
    return out


def table_ids_from_text(text: str) -> set[int]:
    return set(int(x) for x in re.findall(r"\bTable\s+(\d+)\b", text, flags=re.IGNORECASE))


def figure_ids_from_text(text: str) -> set[int]:
    return set(int(x) for x in re.findall(r"\bFigure\s+(\d+)\b", text, flags=re.IGNORECASE))


def table_ids_from_caption_doc(text: str) -> set[int]:
    return set(int(x) for x in re.findall(r"\bTable\s+(\d+)\s*[:.]", text, flags=re.IGNORECASE))


def figure_ids_from_caption_doc(text: str) -> set[int]:
    return set(int(x) for x in re.findall(r"\bFigure\s+(\d+)\s*[:.]", text, flags=re.IGNORECASE))


def parse_author_year_keys(text: str) -> set[str]:
    keys: set[str] = set()
    patterns = [
        r"([A-Z][A-Za-z\-]+)\s+et al\.,?\s*\(?(\d{4})\)?",
        r"([A-Z][A-Za-z\-]+)\s+and\s+[A-Z][A-Za-z\-]+,?\s*\(?(\d{4})\)?",
        r"([A-Z][A-Za-z\-]+)\s*\((\d{4})\)",
    ]
    for pat in patterns:
        for m in re.finditer(pat, text):
            keys.add(f"{m.group(1).lower()}_{m.group(2)}")
    return keys


def extract_reference_block(manuscript_text: str) -> str:
    parts = re.split(r"\bReferences\b", manuscript_text, flags=re.IGNORECASE, maxsplit=1)
    if len(parts) < 2:
        return ""
    tail = parts[1]
    tail = re.split(r"\bTables\b", tail, flags=re.IGNORECASE, maxsplit=1)[0]
    return tail


def snapshot(submission_dir: Path, manifest_path: Path) -> None:
    submission_dir = abs_path(submission_dir)
    manifest_path = abs_path(manifest_path)
    manifest: dict[str, Any] = {
        "schema_version": 1,
        "created_utc": datetime.now(timezone.utc).isoformat(),
        "submission_dir": str(submission_dir.relative_to(ROOT)).replace("\\", "/"),
        "docx": {},
        "zip": {},
    }

    for rel in TRACKED_DOCX:
        p = submission_dir / rel
        lines = docx_lines(p)
        text = "\n".join(lines)
        item = {
            "file_sha256": sha256_file(p),
            "text_sha256": _line_hash(lines),
            "lines_count": len(lines),
            "invariants": docx_invariants(p),
        }
        if rel == "04_Manuscript.docx":
            item["heading_positions"] = heading_positions(lines)
            item["table_refs"] = sorted(table_ids_from_text(text))
            item["figure_refs"] = sorted(figure_ids_from_text(text))
        if rel == "Tables_with_Captions.docx":
            item["table_ids"] = sorted(table_ids_from_caption_doc(text))
        if rel == "Figures_with_Captions.docx":
            item["figure_ids"] = sorted(figure_ids_from_caption_doc(text))
        manifest["docx"][rel] = item

    all_cite_keys = parse_author_year_keys(
        "\n".join(docx_lines(submission_dir / "04_Manuscript.docx"))
        + "\n"
        + "\n".join(docx_lines(submission_dir / "Tables_with_Captions.docx"))
        + "\n"
        + "\n".join(docx_lines(submission_dir / "Figures_with_Captions.docx"))
    )
    manifest["citation_keys_all"] = sorted(all_cite_keys)

    zpath = submission_dir / TRACKED_ZIP
    with zipfile.ZipFile(zpath, "r") as zf:
        members = {}
        for name in sorted(zf.namelist()):
            members[name] = {"sha256": sha256_bytes(zf.read(name))}
    manifest["zip"][TRACKED_ZIP] = {"members": members}

    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(f"Wrote manifest: {manifest_path}")


def load_policy(path: Path) -> dict[str, Any]:
    path = abs_path(path)
    if not path.exists():
        return {"allowed_text_replacements": [], "mutable_zip_members": []}
    return yaml.safe_load(path.read_text(encoding="utf-8")) or {}


def apply_replacements(lines: list[str], replacements: list[dict[str, str]]) -> list[str]:
    if not replacements:
        return lines
    joined = "\n".join(lines)
    for rep in replacements:
        joined = joined.replace(rep["from"], rep["to"])
    return joined.split("\n")


def write_report(report_path: Path, issues: list[str], notes: list[str]) -> None:
    report_path.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        "# Strict Drift Audit Report",
        "",
        f"- Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"- Result: {'PASS' if not issues else 'FAIL'}",
        f"- Issues: {len(issues)}",
        "",
        "## Findings",
    ]
    if issues:
        lines.extend([f"- [ ] {x}" for x in issues])
    else:
        lines.append("- None")
    if notes:
        lines.extend(["", "## Notes"])
        lines.extend([f"- {x}" for x in notes])
    report_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote report: {report_path}")


def audit(submission_dir: Path, manifest_path: Path, policy_path: Path, report_path: Path) -> int:
    submission_dir = abs_path(submission_dir)
    manifest_path = abs_path(manifest_path)
    policy_path = abs_path(policy_path)
    report_path = abs_path(report_path)
    baseline = json.loads(manifest_path.read_text(encoding="utf-8"))
    policy = load_policy(policy_path)
    allowed_reps = policy.get("allowed_text_replacements", [])
    mutable_zip_members = set(policy.get("mutable_zip_members", []))
    issues: list[str] = []
    notes: list[str] = []

    doc_lines: dict[str, list[str]] = {}
    doc_text: dict[str, str] = {}
    for rel in TRACKED_DOCX:
        p = submission_dir / rel
        lines = docx_lines(p)
        rel_reps = [x for x in allowed_reps if x.get("file") == rel]
        lines_for_cmp = apply_replacements(lines, rel_reps)
        actual_text_sha = _line_hash(lines_for_cmp)
        expected = baseline["docx"][rel]["text_sha256"]
        if actual_text_sha != expected:
            issues.append(f"{rel}: text drift outside allowed replacements.")
        inv = docx_invariants(p)
        if inv != baseline["docx"][rel]["invariants"]:
            issues.append(f"{rel}: formatting/layout invariant drift. expected={baseline['docx'][rel]['invariants']} actual={inv}")
        doc_lines[rel] = lines
        doc_text[rel] = "\n".join(lines)

    man_text = doc_text["04_Manuscript.docx"]
    tab_text = doc_text["Tables_with_Captions.docx"]
    fig_text = doc_text["Figures_with_Captions.docx"]

    if "scenario factor ()" in man_text:
        issues.append("04_Manuscript.docx: unresolved placeholder `scenario factor ()` present.")

    hp = heading_positions(doc_lines["04_Manuscript.docx"])
    expected_hp = baseline["docx"]["04_Manuscript.docx"].get("heading_positions")
    if expected_hp and hp != expected_hp:
        issues.append(f"04_Manuscript.docx: heading-position drift. expected={expected_hp} actual={hp}")

    table_ids = table_ids_from_caption_doc(tab_text)
    fig_ids = figure_ids_from_caption_doc(fig_text)
    man_table_refs = table_ids_from_text(man_text)
    man_fig_refs = figure_ids_from_text(man_text)
    uncited_tables = sorted(x for x in table_ids if x not in man_table_refs)
    uncited_figs = sorted(x for x in fig_ids if x not in man_fig_refs)
    if uncited_tables:
        issues.append(f"Tables_with_Captions.docx: uncited table IDs in manuscript text: {uncited_tables}")
    if uncited_figs:
        issues.append(f"Figures_with_Captions.docx: uncited figure IDs in manuscript text: {uncited_figs}")

    cited_keys = sorted(parse_author_year_keys(man_text + "\n" + tab_text + "\n" + fig_text))
    expected_cites = baseline.get("citation_keys_all", [])
    if cited_keys != expected_cites:
        issues.append("Citation-key inventory drift outside baseline.")

    zpath = submission_dir / TRACKED_ZIP
    with zipfile.ZipFile(zpath, "r") as zf:
        cur_members = {name: sha256_bytes(zf.read(name)) for name in sorted(zf.namelist())}
        expected_members = baseline["zip"][TRACKED_ZIP]["members"]
        if set(cur_members.keys()) != set(expected_members.keys()):
            missing = sorted(set(expected_members.keys()) - set(cur_members.keys()))
            extra = sorted(set(cur_members.keys()) - set(expected_members.keys()))
            issues.append(f"{TRACKED_ZIP}: member set drift. missing={missing} extra={extra}")
        for name, meta in expected_members.items():
            if name in mutable_zip_members:
                continue
            if name in cur_members and cur_members[name] != meta["sha256"]:
                issues.append(f"{TRACKED_ZIP}: non-mutable member drift `{name}`.")

        for name in ["cross_validation_300.json", "cross_validation_600.json", "cross_validation_blocked_300.json"]:
            if name not in cur_members:
                issues.append(f"{TRACKED_ZIP}: missing required member `{name}`.")
        try:
            cv300 = json.loads(zf.read("cross_validation_300.json"))
            cv600 = json.loads(zf.read("cross_validation_600.json"))
            cvb = json.loads(zf.read("cross_validation_blocked_300.json"))
            if cv300.get("fold_mode") != "random":
                issues.append("Supplementary_Data_S2.zip: cross_validation_300.json must have fold_mode=random.")
            if cv600.get("fold_mode") != "random":
                issues.append("Supplementary_Data_S2.zip: cross_validation_600.json must have fold_mode=random.")
            if cvb.get("fold_mode") != "blocked":
                issues.append("Supplementary_Data_S2.zip: cross_validation_blocked_300.json must have fold_mode=blocked.")
            if cv300 == cvb or cv600 == cvb:
                issues.append("Supplementary_Data_S2.zip: random CV files are identical to blocked CV evidence.")
        except KeyError:
            pass

    notes.append("Policy-applied text replacements are allowed only where explicitly listed.")
    write_report(report_path=report_path, issues=issues, notes=notes)
    return 0 if not issues else 1


def main() -> None:
    parser = argparse.ArgumentParser(description="Strict no-drift guard for JAES submission regeneration.")
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_snapshot = sub.add_parser("snapshot")
    p_snapshot.add_argument("--submission-dir", default=str(DEFAULT_SUBMISSION_DIR))
    p_snapshot.add_argument("--manifest", default=str(DEFAULT_MANIFEST))

    p_audit = sub.add_parser("audit")
    p_audit.add_argument("--submission-dir", default=str(DEFAULT_SUBMISSION_DIR))
    p_audit.add_argument("--manifest", default=str(DEFAULT_MANIFEST))
    p_audit.add_argument("--policy", default=str(DEFAULT_POLICY))
    p_audit.add_argument("--report", default=str(ROOT / "review" / f"STRICT_DRIFT_REPORT_{datetime.now().strftime('%Y%m%d')}.md"))

    args = parser.parse_args()
    if args.cmd == "snapshot":
        snapshot(submission_dir=Path(args.submission_dir), manifest_path=Path(args.manifest))
        return
    if args.cmd == "audit":
        raise SystemExit(
            audit(
                submission_dir=Path(args.submission_dir),
                manifest_path=Path(args.manifest),
                policy_path=Path(args.policy),
                report_path=Path(args.report),
            )
        )


if __name__ == "__main__":
    main()
