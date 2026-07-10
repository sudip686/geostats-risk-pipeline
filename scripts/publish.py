from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
BUILD_SCRIPT = ROOT / "scripts" / "build_paper_from_meta.py"
RUN_DIR_FALLBACK = ROOT / "outputs_fit_tuning" / "refine" / "refine_n70_ng020"
SUBMISSION_READY = ROOT / "submission_ready"
INTERNAL_READY = ROOT / "internal_ready"
INTERNAL_DIR = ROOT / "internal"

SUBMISSION_DIR = ROOT / "submission"
PAPER_MIRROR = ROOT / "paper"
REPO_SUBMISSION = ROOT / "repo" / "submission"
ROOT_MANUSCRIPT = ROOT / "manuscript.md"
REPO_MANUSCRIPT = ROOT / "repo" / "manuscript.md"
EXTRA_DOCS = [
    "Authors Statement.docx",
    "Conflict of interest.docx",
    "Covering Letter-MS.docx",
    "Declaration of interest statement.docx",
    "Title page.docx",
    "Table.docx",
]


def resolve_run_dir() -> Path:
    best_summary = ROOT / "outputs_fit_tuning" / "best_summary.json"
    if best_summary.exists():
        try:
            data = json.loads(best_summary.read_text(encoding="utf-8"))
            out_dir = data.get("best_output_dir")
            if out_dir:
                p = Path(out_dir)
                if p.exists():
                    return p
        except Exception:
            pass
    return RUN_DIR_FALLBACK


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def clean_dir(target: Path) -> None:
    target.mkdir(parents=True, exist_ok=True)
    for p in sorted(target.rglob("*"), reverse=True):
        if p.is_file():
            try:
                p.unlink()
            except PermissionError:
                # Allow locked Office files to remain; they will be bypassed on copy.
                pass
        elif p.is_dir():
            try:
                p.rmdir()
            except OSError:
                pass


def build_profile(profile: str, run_dir: Path, out_dir: Path) -> None:
    base_manuscript = ROOT / "submission" / "paper_body.md"
    if not base_manuscript.exists():
        base_manuscript = ROOT / "submission" / "paper.md"
    if not base_manuscript.exists():
        base_manuscript = ROOT / "manuscript.md"
    cmd = [
        "python",
        str(BUILD_SCRIPT),
        "--profile",
        profile,
        "--run-dir",
        str(run_dir),
        "--out-dir",
        str(out_dir),
        "--project-yaml",
        str(ROOT / "config" / "project_best_fit.yaml"),
        "--base-manuscript",
        str(base_manuscript),
    ]
    subprocess.run(cmd, check=True)


def extract_figure_entries(paper_md: str) -> list[tuple[str, str]]:
    entries: list[tuple[str, str]] = []
    m = re.search(r"\n## FIGURES\n([\s\S]*)$", paper_md)
    if not m:
        return entries
    body = m.group(1)
    for line in body.splitlines():
        line = line.strip()
        mm = re.match(r"^!\[(Figure[^\]]+)\]\(([^)]+)\)\{ width=72% \}$", line)
        if mm:
            entries.append((mm.group(1), mm.group(2)))
    return entries


def ensure_paper_has_figures_section(out_dir: Path) -> None:
    paper = out_dir / "paper.md"
    text = paper.read_text(encoding="utf-8")
    if "\n## FIGURES\n" in text:
        return

    fig_map = [
        ("Figure 1. Regional geological map with project location.", "figures/geology_regional_map.png"),
        ("Figure 2. Directional variograms and model.", "figures/variogram.png"),
        ("Figure 3. Histogram validation.", "figures/histogram_validation.png"),
        ("Figure 4. Q-Q validation.", "figures/qq_plot.png"),
        ("Figure 5A. Swath X.", "figures/swath_x.png"),
        ("Figure 5B. Swath Y.", "figures/swath_y.png"),
        ("Figure 5C. Swath Z.", "figures/swath_z.png"),
        ("Figure 6. Variogram reproduction.", "figures/variogram_reproduction.png"),
    ]
    lines = ["", "## FIGURES", ""]
    for cap, path in fig_map:
        if (out_dir / path).exists():
            lines.append(f"![{cap}]({path}){{ width=72% }}")
            lines.append("")
    paper.write_text(text.rstrip() + "\n" + "\n".join(lines), encoding="utf-8")


def generate_docx_bundle(out_dir: Path) -> Path:
    ensure_paper_has_figures_section(out_dir)
    tmp = out_dir / "paper.__tmp__.docx"
    if tmp.exists():
        tmp.unlink()

    subprocess.run(["pandoc", "paper.md", "-o", tmp.name], cwd=out_dir, check=True)
    target = out_dir / "paper.docx"
    os.replace(tmp, target)
    shutil.copy2(target, out_dir / "SG Manuscript.docx")

    md = (out_dir / "paper.md").read_text(encoding="utf-8")
    entries = extract_figure_entries(md)

    for name in ("Fig.docx", "Figures_with_titles.docx"):
        doc = Document()
        doc.add_heading("Figures and Captions", level=1)
        for caption, rel in entries:
            fig = out_dir / rel
            if not fig.exists():
                continue
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ip = doc.add_paragraph()
            run = ip.add_run()
            run.add_picture(str(fig), width=Inches(6.2))
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")
        doc.save(out_dir / name)

    return target


def write_manifest(out_dir: Path) -> None:
    files = [p for p in out_dir.rglob("*") if p.is_file() and p.name not in {"submission_package_final.zip", "internal_package_final.zip"}]
    data = {
        "pipeline": "scripts/publish.py",
        "source_of_truth": str((ROOT / "build").resolve()),
        "out_dir": str(out_dir),
        "files": [
            {
                "path": str(p.relative_to(out_dir).as_posix()),
                "bytes": p.stat().st_size,
                "sha256": sha256(p),
            }
            for p in sorted(files)
        ],
    }
    (out_dir / "publish_manifest.json").write_text(json.dumps(data, indent=2), encoding="utf-8")


def build_zip(out_dir: Path, name: str) -> None:
    zip_path = out_dir / name
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in sorted(out_dir.rglob("*")):
            if not p.is_file() or p == zip_path:
                continue
            zf.write(p, p.relative_to(out_dir).as_posix())


def copy_extra_docs(out_dir: Path) -> None:
    for name in EXTRA_DOCS:
        src = ROOT / name
        if src.exists():
            shutil.copy2(src, out_dir / name)


def enforce_submission_sanitization(out_dir: Path) -> None:
    bad_patterns = [
        r"Serric_Data\.csv",
        r"validation_metrics_pre\.json",
        r"tuning was disabled",
        r"100 realizations",
        r"0\.48\s*km",
    ]
    for p in out_dir.rglob("*.md"):
        if p.name == "CONTRADICTION_FIX_REPORT.md":
            continue
        t = p.read_text(encoding="utf-8")
        for pat in bad_patterns:
            if re.search(pat, t, flags=re.I):
                raise RuntimeError(f"Submission sanitization failed: '{pat}' found in {p}")


def sync_submission_mirrors(out_dir: Path) -> None:
    # Non-destructive sync: overwrite generated files, preserve unrelated user files.
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    for p in sorted(out_dir.rglob("*")):
        if not p.is_file():
            continue
        rel = p.relative_to(out_dir)
        dst = SUBMISSION_DIR / rel
        dst.parent.mkdir(parents=True, exist_ok=True)
        try:
            shutil.copy2(p, dst)
        except PermissionError:
            if dst.suffix.lower() == ".docx":
                alt = dst.with_name(dst.stem + ".generated" + dst.suffix)
                shutil.copy2(p, alt)
            else:
                raise

    clean_dir(PAPER_MIRROR)
    clean_dir(REPO_SUBMISSION)
    for mirror in (PAPER_MIRROR, REPO_SUBMISSION):
        for p in sorted(out_dir.rglob("*")):
            if not p.is_file():
                continue
            rel = p.relative_to(out_dir)
            dst = mirror / rel
            dst.parent.mkdir(parents=True, exist_ok=True)
            try:
                shutil.copy2(p, dst)
            except PermissionError:
                if dst.suffix.lower() == ".docx":
                    alt = dst.with_name(dst.stem + ".generated" + dst.suffix)
                    shutil.copy2(p, alt)
                else:
                    raise

    paper_text = (out_dir / "paper.md").read_text(encoding="utf-8")
    ROOT_MANUSCRIPT.write_text(paper_text, encoding="utf-8")
    REPO_MANUSCRIPT.write_text(paper_text, encoding="utf-8")


def sync_internal_bundle(out_dir: Path) -> None:
    clean_dir(INTERNAL_DIR)
    for p in sorted(out_dir.rglob("*")):
        if not p.is_file():
            continue
        rel = p.relative_to(out_dir)
        dst = INTERNAL_DIR / rel
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(p, dst)


def main() -> None:
    parser = argparse.ArgumentParser(description="Publish manuscript package from one source of truth")
    parser.add_argument("--strict", action="store_true")
    parser.add_argument("--run-dir", default=None, help="Optional run directory override")
    args = parser.parse_args()

    run_dir = Path(args.run_dir) if args.run_dir else resolve_run_dir()
    profiles = ["submission", "internal"]

    for profile in profiles:
        out_dir = SUBMISSION_READY if profile == "submission" else INTERNAL_READY
        clean_dir(out_dir)
        build_profile(profile, run_dir, out_dir)
        copy_extra_docs(out_dir)
        generate_docx_bundle(out_dir)
        if profile == "submission" and args.strict:
            enforce_submission_sanitization(out_dir)
        write_manifest(out_dir)
        zip_name = "submission_package_final.zip" if profile == "submission" else "internal_package_final.zip"
        build_zip(out_dir, zip_name)

    sync_submission_mirrors(SUBMISSION_READY)
    sync_internal_bundle(INTERNAL_READY)
    # Temp build folders are not needed after sync.
    clean_dir(SUBMISSION_READY)
    clean_dir(INTERNAL_READY)

    # Keep script mirrored in repo folder.
    repo_script = ROOT / "repo" / "scripts" / "publish.py"
    repo_script.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(Path(__file__), repo_script)

    print("Publish pipeline complete.")
    print(f"- submission: {SUBMISSION_DIR}")
    print(f"- internal: {INTERNAL_DIR}")


if __name__ == "__main__":
    main()
